from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

def style_doc(doc):
    """Set base styles for the document."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_meta_line(doc, label, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run_val = p.add_run(value)
    run_val.font.size = Pt(10)
    run_val.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_after = Pt(1)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
    
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_rest = p.add_run(text)
        run_rest.font.size = Pt(10)
    else:
        p.runs[0].font.size = Pt(10) if p.runs else None
        if not p.runs:
            run = p.add_run(text)
            run.font.size = Pt(10)
    return p

# ============================================================
# DOCUMENT 1: Daily Progress Log
# ============================================================
def build_progress_log():
    doc = Document()
    style_doc(doc)
    
    # Title
    title = doc.add_heading('Daily Progress Log — LaaS Infrastructure', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    
    # Meta
    add_meta_line(doc, "Date", "11 March 2026")
    add_meta_line(doc, "Author", "Punith")
    add_meta_line(doc, "Sprint", "Infrastructure & Platform Core")
    
    doc.add_paragraph()  # spacer
    
    # --- Section 1: Focus ---
    add_styled_heading(doc, "Today's Focus: Full Infrastructure Stack Validation", 1)
    p = doc.add_paragraph(
        "Wrapped up the complete infrastructure setup and validation cycle. The core GPU fractional sharing stack "
        "is now functional end-to-end on the POC machine (ProArt X670E / Ryzen 9 7950X3D / RTX 4090 / 64GB DDR5)."
    )
    
    # --- Section 2: What Got Done ---
    add_styled_heading(doc, "What Got Done", 1)
    
    # 2.1 GPU Fractional
    add_styled_heading(doc, "1. GPU Fractional Sharing — Validated ✅", 2)
    items = [
        "HAMi-core (libvgpu.so) built from source, VRAM enforcement confirmed working on Ada architecture (sm_89)",
        "Containers correctly report capped VRAM (4GB / 8GB / 16GB) instead of the full 24GB",
        "Over-allocation attempts are rejected cleanly — PyTorch and raw CUDA both respect the limit",
        "Ran into two system crashes during early HAMi-core testing (incorrect LD_PRELOAD path + a cmake build flag issue). Both resolved. Documented the exact build flags that work.",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 2.2 MPS
    add_styled_heading(doc, "2. CUDA MPS Multi-Tenant Partitioning — Validated ✅", 2)
    items = [
        "MPS control daemon running as a systemd service, auto-starts on boot",
        "4 concurrent GPU containers running independently with individual VRAM caps (4+4+8+8 = 24GB total)",
        "SM partitioning via CUDA_MPS_ACTIVE_THREAD_PERCENTAGE confirmed functional — each container gets its compute slice",
        "Fault recovery tested: killed a CUDA process in container 1, MPS server auto-recovered, containers 2/3/4 continued unaffected",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 2.3 Selkies
    add_styled_heading(doc, "3. Selkies EGL Desktop Streaming — Validated ✅", 2)
    items = [
        "Built the custom Selkies EGL Desktop image (Ubuntu 22.04 + KDE Plasma + CUDA 12.8)",
        "Full GPU-accelerated KDE desktop renders inside a Docker container and streams to the browser via WebRTC",
        "NVENC hardware encoding active — confirmed nvh264enc in Selkies logs",
        "Ran 4 concurrent desktop sessions on a single 4090, each with its own VRAM limit. All 4 accessible simultaneously from different browser tabs.",
        "Desktop responsiveness is solid — feels like a local session, not a remote one",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 2.4 Storage
    add_styled_heading(doc, "4. Storage Layer — Simulated & Working ✅", 2)
    items = [
        "ZFS pool created with per-user datasets and 15GB quotas",
        "NFS export + mount working (loopback for POC, production will be TrueNAS over 10GbE)",
        "Persistent home directories survive container restarts — files created in session 1 persist in session 2",
        "Container stop → restart → NFS home intact. Container destroy → new container → same home directory, all files present.",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 2.5 Security
    add_styled_heading(doc, "5. Container Isolation & Security Hardening — Applied ✅", 2)
    items = [
        "--read-only base filesystem (no writes to image layers)",
        "--security-opt no-new-privileges on all containers",
        "PID namespace isolation, --pids-limit 512",
        "cgroups v2 enforced: --cpus and --memory caps confirmed working",
        "AppArmor default profile active",
        "No root/sudo access inside containers",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # 2.6 Monitoring
    add_styled_heading(doc, "6. Monitoring Stack — Deployed ✅", 2)
    items = [
        "DCGM Exporter collecting GPU metrics (utilization, VRAM usage, temperature, power draw)",
        "Prometheus scraping DCGM + node-exporter + Docker metrics",
        "Grafana dashboard live with per-container GPU utilization, VRAM allocation, thermal readings",
        "Alert rules configured: GPU temp > 80°C warning, > 90°C critical",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # --- Section 3: Known Issues ---
    add_styled_heading(doc, "Known Issues Encountered & Resolved", 1)
    
    add_table(doc,
        ["Issue", "What Happened", "Resolution"],
        [
            ["System crash #1", "Incorrect LD_PRELOAD path caused kernel panic during container startup", "Fixed: must use absolute path inside container, not host path"],
            ["System crash #2", "HAMi-core built with wrong cmake flags segfaulted on cuMemAlloc interception", "Fixed: -DCMAKE_BUILD_TYPE=Release is mandatory. Debug build has symbol conflicts with CUDA driver"],
            ["MPS single-user limitation", "MPS server initially only accepted one UID", "All containers must run under the same UID when sharing a single MPS server"],
            ["NVENC session limit", "Default limit is 3 concurrent NVENC sessions on GeForce", "Applied nvidia-patch (keylase). Verified 4+ concurrent NVENC streams working"],
            ["Selkies WebRTC NAT issue", "Browser couldn't connect to WebRTC stream on remote network", "Expected for POC (no TURN server). Production will have coturn. Local works fine."],
        ]
    )
    
    doc.add_paragraph()  # spacer
    
    # --- Section 4: Next Steps ---
    add_styled_heading(doc, "Next Steps", 1)
    
    add_styled_heading(doc, "Kubernetes Cluster Setup", 2)
    items = [
        "Transition from standalone Docker to a proper K8s cluster for container orchestration",
        "Deploy NVIDIA device plugin + HAMi scheduler for automated GPU VRAM accounting",
        "Implement pod-level resource quotas mirroring the tiered pricing model (Pro/Power/Max/Full)",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_styled_heading(doc, "Monitoring & Observability Expansion", 2)
    items = [
        "Add Loki for centralized container log aggregation",
        "Build GPU VRAM utilization heatmap dashboard (per-node, per-container view)",
        "Session lifecycle tracking: startup time, streaming quality metrics, idle detection",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_styled_heading(doc, "Security Enforcement — Next Layer", 2)
    items = [
        "Implement seccomp profiles restricting dangerous syscalls inside GPU containers",
        "Custom AppArmor profiles for the Selkies containers (restrict filesystem access to /home only)",
        "Network policies: container-to-container communication blocked at CNI level",
        "Egress filtering: containers can reach package repos and NFS only, no arbitrary internet access",
        "Session idle timeout: auto-pause after 30 min inactive, auto-terminate after 2 hours",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    # --- Section 5: Metrics ---
    add_styled_heading(doc, "Metrics Snapshot (from today's test runs)", 1)
    
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Concurrent GPU desktop sessions tested", "4"],
            ["VRAM enforcement accuracy", "✅ Exact (within 2MB of target)"],
            ["Container cold start (to usable desktop)", "~12 seconds"],
            ["WebRTC stream latency (local)", "~18ms frame-to-frame"],
            ["GPU utilization (4 idle desktops)", "~8%"],
            ["GPU utilization (4 desktops + 1 CUDA workload)", "~62%"],
            ["MPS fault recovery time", "< 3 seconds"],
            ["NFS home directory mount time", "< 1 second"],
        ]
    )
    
    doc.save('c:/Users/Punith/LaaS/Daily_Progress_Log.docx')
    print("✅ Daily_Progress_Log.docx created")


# ============================================================
# DOCUMENT 2: Architecture Decisions
# ============================================================
def build_architecture_decisions():
    doc = Document()
    style_doc(doc)
    
    # Title
    title = doc.add_heading('Architecture Decisions — LaaS Platform', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    
    # Meta
    add_meta_line(doc, "Date", "11 March 2026")
    add_meta_line(doc, "Author", "Punith")
    add_meta_line(doc, "Status", "Validated on POC hardware (RTX 4090), pending production deployment on RTX 5090 fleet")
    
    doc.add_paragraph()
    
    # --- 1. Containers vs VMs ---
    add_styled_heading(doc, "1. Why Containers Over VMs for GPU Tiers", 1)
    
    p = doc.add_paragraph(
        "The original architecture (v1) used KVM VMs with full GPU passthrough for every tier. "
        "This meant one GPU per user, period. That made Tier 1 (Full Machine) work, but completely killed the economics "
        "of Tier 2 — there's no way to charge ₹X/hour for 4GB VRAM when you're tying up a 32GB card per user."
    )
    
    p = doc.add_paragraph()
    run = p.add_run("The pivot: ")
    run.bold = True
    p.add_run(
        "After deep research into how university HPC clusters (NRP Nautilus, GPUnion) handle multi-tenant GPU access, "
        "I landed on a container-first architecture using Selkies EGL Desktop containers."
    )
    
    add_table(doc,
        ["", "VM Passthrough (v1)", "Selkies Containers (current)"],
        [
            ["GPU per user", "1:1 (exclusive)", "N:1 (fractional)"],
            ["Concurrent GPU users/node", "1", "4-8"],
            ["VRAM tiers", "Not possible", "4 / 8 / 16 / 32 GB (enforced)"],
            ["Desktop streaming", "xrdp (laggy)", "WebRTC via NVENC (smooth, 60fps)"],
            ["GPU driver mode", "vfio-pci switching", "Always nvidia-driver (stable)"],
        ]
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph(
        "This also eliminated the D3cold power state bug on RTX GPUs, which was causing the GPU to become "
        "unreachable after VM shutdown. Spent a good amount of time chasing that one before realizing "
        "the entire VM passthrough path was architecturally wrong for what we needed."
    )
    
    # --- 2. Dual-Layer VRAM ---
    add_styled_heading(doc, "2. Dual-Layer VRAM Enforcement (HAMi-core + MPS)", 1)
    
    p = doc.add_paragraph(
        "This was the hardest decision and the one with the least documentation to lean on. "
        "We needed a way to tell a container \"you have 4GB of VRAM\" when the physical GPU has 32GB — "
        "and actually enforce it, not just suggest it."
    )
    
    p = doc.add_paragraph()
    run = p.add_run("Why two layers, not one:")
    run.bold = True
    
    items = [
        ("HAMi-core (libvgpu.so): ", "Intercepts CUDA API calls at the driver level. When PyTorch calls cuMemAlloc, HAMi-core checks against the configured limit and returns OOM if exceeded. This is what makes nvidia-smi inside the container show 4GB instead of 32GB."),
        ("CUDA MPS: ", "NVIDIA's official multi-process service. Provides per-client GPU address space isolation (Volta+), compute partitioning via ACTIVE_THREAD_PERCENTAGE, and its own memory limit via PINNED_DEVICE_MEM_LIMIT."),
    ]
    for bold_part, rest in items:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(rest)

    p = doc.add_paragraph(
        "Neither one alone is sufficient. HAMi-core handles the VRAM reporting trick (critical for PyTorch/TensorFlow "
        "auto-detection) but doesn't partition compute. MPS partitions compute but doesn't fake the VRAM total. "
        "Together, they cover both gaps."
    )
    
    p = doc.add_paragraph()
    run = p.add_run("What I couldn't find documented anywhere: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    p.add_run(
        "Nobody seems to have combined these two in the same container stack. The HAMi project uses HAMi-core within "
        "their Kubernetes device plugin, and MPS is used in SLURM HPC clusters, but the dual-layer combination is — "
        "as far as I could find — novel to this deployment. Proceeded carefully because of that. The two system crashes "
        "during early testing were both related to getting the interception order right."
    )
    
    # --- 3. Selkies ---
    add_styled_heading(doc, "3. Selkies EGL Desktop — Why This and Not VNC/xrdp", 1)
    
    p = doc.add_paragraph(
        "Previous iterations used xrdp for remote desktop. It worked, but the experience was noticeably laggy "
        "and GPU acceleration wasn't available (xrdp renders in software by default)."
    )
    p = doc.add_paragraph(
        "Selkies EGL Desktop was discovered while researching how NRP Nautilus serves GPU desktops to researchers "
        "across 50+ universities. It solves three problems simultaneously:"
    )
    
    items = [
        ("GPU rendering without X11: ", "Uses EGL backend (direct GPU access via DRI device), no host-side X server needed"),
        ("Multi-container GPU sharing: ", 'Explicitly designed for it — "the EGL variant supports sharing one GPU with many containers"'),
        ("Browser delivery: ", "WebRTC streaming with NVENC encoding — no client install required, works on any modern browser"),
    ]
    for bold_part, rest in items:
        p = doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(bold_part)
        run_b.bold = True
        p.add_run(rest)
    
    p = doc.add_paragraph(
        "The tradeoff is that we need to apply nvidia-patch to remove the NVENC concurrent session limit "
        "(GeForce cards default to 3). This is a well-maintained community patch with years of track record "
        "in the media server community, but it is technically modifying the NVIDIA driver."
    )
    
    # --- 4. Storage ---
    add_styled_heading(doc, "4. Storage: ZFS + NFS (Not Ceph, Not Local)", 1)
    
    add_table(doc,
        ["Option", "Verdict", "Why"],
        [
            ["Local disk per node", "❌ Rejected", "User data tied to specific node. Can't migrate sessions."],
            ["Ceph distributed", "❌ Overkill", "Needs 3+ nodes minimum, complex for this scale"],
            ["ZFS + NFS (TrueNAS)", "✅ Selected", "Simple, proven, per-user quotas built into ZFS, NFS over 10GbE is fast enough"],
        ]
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph(
        "Each user gets a 15GB ZFS dataset with a hard quota. Mounted into every container at /home/<uid> via NFS. "
        "Base image is read-only — system software is immutable. User data persists across sessions, container restarts, "
        "even container destruction."
    )
    
    # --- 5. Why not K8s ---
    add_styled_heading(doc, "5. Why Not Kubernetes on Day 1", 1)
    
    p = doc.add_paragraph(
        "The current POC runs on standalone Docker with a FastAPI orchestrator managing container lifecycle via "
        "the Python Docker SDK. The argument for Kubernetes is obvious (scheduling, scaling, self-healing), "
        "but for Phase 0/1 validation:"
    )
    items = [
        "K8s adds a week of setup complexity that doesn't help validate the GPU sharing stack",
        "The HAMi Kubernetes device plugin exists but adds another variable to debug",
        "Standalone Docker lets us isolate HAMi-core + MPS issues from orchestration issues",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph(
        "K8s is the next step (already on the roadmap), but the validation had to happen with minimal moving parts first. "
        "Now that the core stack is proven, the migration is straightforward."
    )
    
    # --- 6. Known Limitations ---
    add_styled_heading(doc, "6. Known Limitations & Honest Tradeoffs", 1)
    
    p = doc.add_paragraph()
    run = p.add_run("GPU L2 Cache / Memory Bandwidth — Shared, Not Partitioned")
    run.bold = True
    doc.add_paragraph(
        "Consumer RTX GPUs don't have MIG (Multi-Instance GPU). L2 cache and DRAM bandwidth are shared across "
        "all containers. In practice: if one user runs a bandwidth-heavy workload, others on the same GPU might "
        "see slightly degraded throughput. This is the same limitation every university HPC cluster with MPS faces. "
        "Mitigation: Tier 1 (Full Machine) gives exclusive GPU access for users who need guaranteed performance."
    )
    
    p = doc.add_paragraph()
    run = p.add_run("GPU Fatal Fault Propagation")
    run.bold = True
    doc.add_paragraph(
        "If a user writes a buggy CUDA kernel that crashes the GPU, MPS auto-recovers (Volta+ feature), but "
        "all co-resident containers get interrupted. It's a ~30 second disruption, not data loss (NFS home is safe). "
        "For the target audience (students running MATLAB, PyTorch, Blender), the probability is very low."
    )
    
    p = doc.add_paragraph()
    run = p.add_run("nvidia-patch Dependency")
    run.bold = True
    doc.add_paragraph(
        "The NVENC session limit removal requires patching the driver after every update. There's an active community "
        "maintaining it, but it's a dependency worth noting."
    )
    
    # --- 7. Architecture Evolution ---
    add_styled_heading(doc, "7. Architecture Evolution Timeline", 1)
    
    add_table(doc,
        ["Version", "Architecture", "Problem Encountered"],
        [
            ["v1 (Jan 2026)", "KVM VMs + GPU passthrough per user", "1 GPU per user, D3cold bug, no fractional tiers"],
            ["v2 (Feb 2026)", "Hybrid: VMs for CPU-only, containers for GPU", "vfio-pci mode switching still needed for Tier 1"],
            ["v3 (Current)", "Unified container architecture\nHAMi-core + MPS + Selkies EGL\nAll nodes in nvidia-driver mode", "None — this is the validated architecture"],
        ]
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph(
        "Each iteration was driven by hitting a wall in the previous one. The D3cold bug forced the move from v1 to v2, "
        "and the economics of 1:1 GPU allocation forced the move from v2 to v3. "
        "The current architecture is the one that actually makes the pricing model work."
    )
    
    # --- 8. Novelty ---
    add_styled_heading(doc, "8. What's Genuinely Novel Here", 1)
    
    p = doc.add_paragraph(
        "After extensive research (Selkies project, HAMi-core docs, NVIDIA MPS documentation, NRP Nautilus deployment, "
        "GPUnion academic papers, and every commercial GPU cloud I could find), I'm reasonably confident that:"
    )
    
    items = [
        "The individual technologies are all proven in production separately",
        "The dual-layer HAMi-core + MPS enforcement combination doesn't appear to be documented anywhere publicly",
        "No commercial GPU platform (RunPod, Vast.ai, Lambda, Google Colab) offers fractional VRAM desktop tiers",
        "The closest reference is NRP Nautilus (Selkies + GPU sharing for universities), but they don't enforce VRAM tiers",
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph(
        "We're not inventing new technology — we're combining existing proven components in a configuration "
        "that hasn't been assembled before. Which means there are no reference implementations to copy from, "
        "and the validation had to be done from scratch. The two system crashes (and a few hours of head-scratching "
        "over MPS UID constraints) were the cost of treading new ground."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The POC has now validated that this combination works. The remaining risk is Blackwell-specific behavior "
        "on RTX 5090 (sm_120 vs the 4090's sm_89) — that's the first thing to verify when the production hardware arrives."
    )
    
    doc.save('c:/Users/Punith/LaaS/Architecture_Decisions.docx')
    print("✅ Architecture_Decisions.docx created")


# ============================================================
if __name__ == '__main__':
    build_progress_log()
    build_architecture_decisions()
    print("\nBoth .docx files generated in c:\\Users\\Punith\\LaaS\\")
