"""
KSRCE-GKT AI Lab-as-a-Service (LaaS) — Technical Implementation & SDLC Document Generator
Generates a professionally formatted .docx with two parts:
  Part I  – Infrastructure Setup & Architecture
  Part II – Software Development Lifecycle (SDLC) & Sprint Plan
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "LaaS_Technical_Implementation_SDLC_v2.docx")

# ─── colour palette ───
DARK_BG        = RGBColor(0x0F, 0x17, 0x2A)   # deep navy
ACCENT_CYAN    = RGBColor(0x00, 0xD4, 0xFF)    # neon cyan
ACCENT_PURPLE  = RGBColor(0x7C, 0x3A, 0xED)    # vivid purple
ACCENT_GREEN   = RGBColor(0x10, 0xB9, 0x81)    # emerald
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY     = RGBColor(0xF1, 0xF5, 0xF9)
DARK_TEXT       = RGBColor(0x1E, 0x29, 0x3B)
MID_GRAY       = RGBColor(0x64, 0x74, 0x8B)
TABLE_HEADER_BG = "1E3A5F"
TABLE_ALT_BG    = "F0F4F8"
CAUTION_BG      = "FFF7ED"
CAUTION_BORDER  = "F59E0B"
INFO_BG         = "EFF6FF"
INFO_BORDER     = "3B82F6"
SUCCESS_BG      = "F0FDF4"
SUCCESS_BORDER  = "22C55E"
CRITICAL_BG     = "FEF2F2"
CRITICAL_BORDER = "EF4444"

doc = Document()

# ─── global styles ───
style = doc.styles['Normal']
font  = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = DARK_TEXT
style.paragraph_format.space_after  = Pt(4)
style.paragraph_format.space_before = Pt(2)
style.paragraph_format.line_spacing = 1.15

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── heading styles ───
for level, (size, colour, bold, space_b, space_a) in {
    1: (22, ACCENT_PURPLE, True, 18, 8),
    2: (16, RGBColor(0x1E, 0x40, 0x7A), True, 14, 6),
    3: (13, RGBColor(0x37, 0x47, 0x51), True, 10, 4),
}.items():
    h = doc.styles[f'Heading {level}']
    h.font.size     = Pt(size)
    h.font.color.rgb = colour
    h.font.bold     = bold
    h.font.name     = 'Calibri'
    h.paragraph_format.space_before = Pt(space_b)
    h.paragraph_format.space_after  = Pt(space_a)
    if level == 1:
        h.paragraph_format.keep_with_next = True

# ─── helpers ───
def add_styled_table(doc_obj, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc_obj.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_BG}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_ALT_BG}"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # col widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def add_callout_box(doc_obj, title, text, bg_color, border_color, title_color_rgb=None):
    """Add a coloured callout/info box."""
    tbl = doc_obj.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)

    # set cell shading and border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    tcPr.append(shading)

    # borders
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:color="{border_color}"/>'
        f'  <w:top w:val="single" w:sz="4" w:color="{border_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{border_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    cell.text = ''
    p_title = cell.paragraphs[0]
    run_t = p_title.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(10)
    run_t.font.name = 'Calibri'
    if title_color_rgb:
        run_t.font.color.rgb = title_color_rgb
    else:
        run_t.font.color.rgb = DARK_TEXT

    p_body = cell.add_paragraph()
    run_b = p_body.add_run(text)
    run_b.font.size = Pt(9.5)
    run_b.font.name = 'Calibri'
    run_b.font.color.rgb = DARK_TEXT

    doc_obj.add_paragraph()  # spacer
    return tbl


def add_bullet(doc_obj, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc_obj.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.add_run(text).font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'


def add_para(doc_obj, text, bold=False, italic=False, font_size=10.5, color=None, alignment=None):
    """Add a styled paragraph."""
    p = doc_obj.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    return p


def add_divider(doc_obj):
    """Add a visual divider line."""
    p = doc_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = MID_GRAY


# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("KSRCE – GKT")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = MID_GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI Supercomputing Lab-as-a-Service")
run.bold = True
run.font.size = Pt(30)
run.font.color.rgb = ACCENT_PURPLE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Technical Implementation & SDLC Document")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1E, 0x40, 0x7A)
run.font.name = 'Calibri'

doc.add_paragraph()

add_divider(doc)

doc.add_paragraph()

# Meta info table
meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Document Version", "2.0 — Production"),
    ("Date", "March 08, 2026"),
    ("Classification", "Confidential — Internal Use Only"),
    ("Prepared By", "Punith — Lead Engineer, Global Knowledge Technologies"),
    ("Stakeholders", "GKT Engineering, KSRCE Administration"),
]
for i, (label, value) in enumerate(meta_data):
    c0 = meta_table.cell(i, 0)
    c1 = meta_table.cell(i, 1)
    c0.text = ''
    c1.text = ''
    r0 = c0.paragraphs[0].add_run(label)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = ACCENT_PURPLE
    r0.font.name = 'Calibri'
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = c1.paragraphs[0].add_run(value)
    r1.font.size = Pt(10)
    r1.font.name = 'Calibri'

    # remove borders via XML
    for c in [c0, c1]:
        tcPr = c._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:color="auto"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

meta_table.columns[0].width = Inches(2.2)
meta_table.columns[1].width = Inches(4.3)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS PLACEHOLDER
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Table of Contents', level=1)
add_para(doc, "[Update field in Word: Right-click → Update Field, or press Ctrl+A then F9]", italic=True, color=MID_GRAY, font_size=9)

# Insert actual TOC field
p_toc = doc.add_paragraph()
run_toc = p_toc.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run_toc._r.append(fldChar1)
run_toc2 = p_toc.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run_toc2._r.append(instrText)
run_toc3 = p_toc.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run_toc3._r.append(fldChar2)
run_toc4 = p_toc.add_run("[Table of Contents — press F9 to update]")
run_toc4.font.color.rgb = MID_GRAY
run_toc5 = p_toc.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run_toc5._r.append(fldChar3)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Executive Summary', level=1)

add_para(doc,
    "This document serves as the unified technical blueprint for the KSRCE–GKT AI Supercomputing "
    "Lab-as-a-Service (LaaS) platform — a first-of-its-kind on-premises solution that delivers "
    "fractional GPU desktop computing, ephemeral compute sessions, and full-machine exclusive access "
    "to a fleet of 4 high-performance nodes, each powered by AMD Ryzen 9 9950X3D CPUs and NVIDIA "
    "RTX 5090 32GB GPUs. The platform targets university students, faculty, researchers, and public "
    "users seeking on-demand GPU compute delivered through a browser-based interface."
)

add_para(doc,
    "The document is structured in two parts: Part I covers the complete infrastructure architecture — "
    "hardware, GPU virtualisation strategy, storage, networking, and phased burn-in protocols. "
    "Part II covers the software development lifecycle (SDLC) including the sprint plan, module "
    "breakdown, technology stack, integration dependencies, and delivery timeline aligned with "
    "the March 13, 2026 development completion target and the March 15–17 deployment window."
)

add_callout_box(doc,
    "SPRINT TIMELINE — ACCELERATED DELIVERY",
    "Development Sprint: March 08 – March 13, 2026 (6 days)\n"
    "Final QA & Load Testing: March 15, 2026\n"
    "Soft Launch (Closed Beta — KSRCE faculty + select students): March 16, 2026\n"
    "Public Launch & Marketing Push: March 17, 2026\n\n"
    "All core deliverables must be production-ready by EOD March 13. Quality is non-negotiable.",
    CRITICAL_BG, CRITICAL_BORDER, RGBColor(0xDC, 0x26, 0x26)
)

doc.add_heading('Fleet Capacity at a Glance', level=2)

add_styled_table(doc,
    ["Resource", "Per Node", "Fleet Total (4 Nodes)"],
    [
        ["CPU Cores", "16C / 32T (Ryzen 9 9950X3D)", "64C / 128T"],
        ["System RAM", "64 GB DDR5-6000", "256 GB"],
        ["GPU", "RTX 5090 — 32 GB VRAM", "128 GB VRAM"],
        ["NVMe Storage", "2 TB (7,250 MB/s)", "8 TB"],
        ["Max Concurrent GPU Desktop Sessions", "4–8 (fractional)", "16–32"],
        ["Max Concurrent Ephemeral Sessions", "6–10", "24–40"],
    ],
    col_widths=[2.5, 2.2, 2.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PART I — INFRASTRUCTURE SETUP & ARCHITECTURE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('PART I — Infrastructure Architecture & Setup', level=1)

# ─── 1. Hardware ───
doc.add_heading('1. Hardware Inventory', level=2)

doc.add_heading('1.1 Compute Nodes (4 Identical Machines)', level=3)

add_styled_table(doc,
    ["Component", "Specification"],
    [
        ["CPU", "AMD Ryzen 9 9950X3D (16C/32T, 3D V-Cache, 144MB L3)"],
        ["Motherboard", "ASUS ProArt X670E-Creator WiFi DDR5"],
        ["GPU", "Zotac RTX 5090 Solid OC 32GB GDDR7 (Blackwell, 170 SMs)"],
        ["RAM", "G.SKILL DDR5 Trident Z5 NEO 6000MHz 64GB (2×32GB) CL30"],
        ["Storage", "Samsung 990 EVO Plus NVMe 2TB (7,250 MB/s sequential)"],
        ["PSU", "Corsair AX1600i 1600W (80+ Titanium)"],
        ["Cooling", "Corsair Nautilus RS ARGB 360mm AIO"],
        ["Cabinet", "Corsair 3500X (3 × ARGB fans)"],
    ],
    col_widths=[1.8, 4.8]
)

doc.add_heading('1.2 Additional Procurement Required', level=3)

add_callout_box(doc,
    "PROCUREMENT — Items Required Before Phase 0",
    "The following components are NOT yet available and must be procured before infrastructure "
    "deployment can begin. Estimated total: ₹2,25,000 – ₹3,50,000.",
    INFO_BG, INFO_BORDER
)

add_styled_table(doc,
    ["Category", "Item", "Qty", "Est. Cost (₹)"],
    [
        ["Networking", "Intel X550-T1 10GbE PCIe NIC", "4", "26,000"],
        ["Networking", "Mikrotik CRS309-1G-8S+ 10GbE SFP+ Switch", "1", "25,000"],
        ["Networking", "SFP+ DAC Cables", "6", "5,000"],
        ["Networking", "TP-Link TL-SG108E 8-port 1GbE Managed Switch", "1", "3,000"],
        ["GPU/Display", "HDMI 2.1 4K Dummy Dongles", "4", "2,800"],
        ["NAS (5th Machine)", "ATX System + 4×4TB WD Red Plus + 10GbE NIC + 16GB RAM", "1", "~80,000"],
        ["  (OR Alternative)", "Synology DS923+ with 10GbE Expansion", "1", "~1,20,000"],
        ["Power", "5kVA Online UPS", "1", "40,000–80,000"],
        ["Environment", "Server Room AC / Mini-Split (3–4kW heat)", "1", "Varies"],
        ["Security", "Physical Lock / Access Control for Server Room", "1", "Varies"],
    ],
    col_widths=[1.2, 3.0, 0.5, 1.5]
)

doc.add_paragraph()

# ─── 2. GPU Constraint ───
doc.add_heading('2. The Consumer GPU Challenge — RTX 5090', level=2)

add_para(doc,
    "The RTX 5090 is a GeForce consumer GPU. This is a foundational architectural constraint that "
    "shaped every infrastructure decision. Understanding these limitations is critical to "
    "appreciating the chosen architecture."
)

add_callout_box(doc,
    "CRITICAL — Consumer GPU Limitations (RTX 5090 / GeForce)",
    "• No NVIDIA vGPU — requires RTX PRO 6000 / L40S / A100 (enterprise licensing)\n"
    "• No MIG (Multi-Instance GPU) — only available on A100 / H100 / B200\n"
    "• No SR-IOV GPU Partitioning — not supported on GeForce cards\n"
    "• No Hyper-V GPU-P — Windows Server 2025 explicitly excludes consumer GPUs\n"
    "• vgpu_unlock hack does NOT work on Blackwell architecture\n\n"
    "Impact: GPU cannot be hardware-partitioned across VMs. All GPU sharing must be "
    "software-based (containers + CUDA MPS + HAMi-core) or all-or-nothing VM passthrough.",
    CRITICAL_BG, CRITICAL_BORDER, RGBColor(0xDC, 0x26, 0x26)
)

doc.add_heading('2.1 The D3cold Reset Bug (Confirmed)', level=3)

add_para(doc,
    "After a GPU-passthrough VM shuts down, the RTX 5090 can enter a D3cold power state and fail "
    "to wake. The host CPU experiences a soft lockup, and only a full machine reboot recovers the "
    "GPU. This bug has been confirmed by NVIDIA and is documented across Proxmox forums, "
    "Tom's Hardware, and igor'sLAB. A $1,000 community bounty was posted for a fix."
)

add_para(doc,
    "Impact: VM-based GPU passthrough (vfio-pci) is unreliable for production multi-tenant "
    "sequential sessions without extensive mitigation. This finding directly drove the "
    "architectural pivot from VM-centric to container-centric GPU delivery.", bold=True
)

doc.add_heading('2.2 Architectural Pivot — From VMs to Unified Containers', level=3)

add_callout_box(doc,
    "ARCHITECTURAL DECISION — The Pivot",
    "Initial Design (v1): KVM VMs for stateful desktops (CPU-only, GPU passthrough for Tier 1). "
    "Ephemeral containers for Jupyter/CLI.\n\n"
    "Problem: Consumer GPU cannot be split across VMs. Only 1 GPU user per VM passthrough. "
    "D3cold reset bug makes sequential GPU passthrough unreliable.\n\n"
    "Final Design (v2 — Current): ALL user sessions (CPU-only and GPU-accelerated, stateful and "
    "ephemeral) run as Selkies EGL Desktop containers with software-enforced GPU fractioning "
    "via HAMi-core + CUDA MPS. No KVM VMs for user compute. Proxmox is retained for host "
    "node management only.\n\n"
    "Result: 4–8 concurrent GPU desktop users per node vs. 1 under the VM model. "
    "Full fractional VRAM pricing (4GB / 8GB / 16GB / 32GB) becomes viable.",
    SUCCESS_BG, SUCCESS_BORDER, RGBColor(0x16, 0x65, 0x34)
)

# ─── 3. Architecture ───
doc.add_heading('3. Chosen Architecture — Unified Selkies Container Model', level=2)

add_para(doc,
    "All 4 nodes are identically configured. NVIDIA driver is loaded on the host at all times. "
    "No vfio-pci mode. The D3cold bug is completely avoided because the GPU never changes "
    "ownership — it is always driven by the host. Three breakthrough technologies combine to "
    "deliver fractional GPU desktops on consumer hardware:"
)

doc.add_heading('3.1 The Three Enabling Technologies', level=3)

add_styled_table(doc,
    ["Technology", "Role", "Key Capability"],
    [
        ["Selkies EGL Desktop", "GPU-Shared Desktop Containers",
         "Full KDE Plasma desktop inside Docker. VirtualGL EGL backend for GPU-accelerated "
         "rendering. WebRTC streaming to browser at 60fps. Explicitly supports sharing one GPU "
         "across many containers."],
        ["HAMi-core (libvgpu.so)", "CUDA API Interception",
         "Intercepts cuMemAlloc to enforce hard VRAM limits per container. nvidia-smi inside "
         "container reports limited VRAM (e.g., 4GB instead of 32GB). No kernel module required. "
         "Works on consumer GPUs."],
        ["CUDA MPS", "Compute Partitioning",
         "Per-client VRAM caps (second enforcement layer). SM partitioning — RTX 5090 has 170 SMs "
         "allowing fine-grained compute allocation. Isolated GPU virtual address spaces per client "
         "(Volta+). Supports up to 48 concurrent CUDA contexts."],
    ],
    col_widths=[1.6, 1.8, 3.2]
)

add_callout_box(doc,
    "WHY DUAL-LAYER VRAM ENFORCEMENT",
    "HAMi-core (Layer 1) catches allocation attempts at the CUDA API level and makes nvidia-smi "
    "report the correct limited VRAM. CUDA MPS (Layer 2) enforces hard caps at the driver level — "
    "even if HAMi-core is bypassed by a statically-linked binary, MPS still enforces the limit. "
    "Together, they provide defense-in-depth that neither alone can achieve.",
    INFO_BG, INFO_BORDER
)

doc.add_heading('3.2 Compute Configuration Tiers', level=3)

add_para(doc, "Stateful GUI Desktop Configs (All delivered as Selkies EGL Desktop containers):", bold=True)

add_styled_table(doc,
    ["Config", "vCPU", "RAM", "GPU VRAM", "Best For", "Price/hr"],
    [
        ["Starter", "2", "4 GB", "None", "Light coding, docs", "₹15"],
        ["Standard", "4", "8 GB", "None", "MATLAB, development", "₹30"],
        ["Pro", "4", "8 GB", "4 GB", "Blender, ML inference", "₹60"],
        ["Power", "8", "16 GB", "8 GB", "Heavy sim, 3D render", "₹100"],
        ["Max", "8", "16 GB", "16 GB", "Large model training", "₹150"],
        ["Full Machine", "16", "48 GB", "32 GB (exclusive)", "DL research, CUDA dev", "₹300"],
    ],
    col_widths=[1.1, 0.6, 0.7, 1.0, 1.8, 0.9]
)

doc.add_paragraph()

add_para(doc, "Ephemeral Compute Configs (Docker Containers — Jupyter / Code-Server / SSH):", bold=True)

add_styled_table(doc,
    ["Config", "vCPU", "RAM", "GPU VRAM", "Best For", "Price/hr"],
    [
        ["Ephemeral CPU", "2", "4 GB", "None", "Jupyter, Code-Server, SSH", "₹10"],
        ["Ephemeral GPU-S", "2", "4 GB", "4 GB", "Light ML inference", "₹40"],
        ["Ephemeral GPU-M", "4", "8 GB", "8 GB", "Standard ML training", "₹75"],
        ["Ephemeral GPU-L", "8", "16 GB", "16 GB", "Heavy ML, large models", "₹120"],
    ],
    col_widths=[1.3, 0.6, 0.7, 1.0, 2.0, 0.9]
)

doc.add_heading('3.3 Isolation Model', level=3)

add_styled_table(doc,
    ["Isolation Layer", "Level", "Mechanism"],
    [
        ["CPU / RAM", "Hard (kernel-enforced)", "cgroups v2 (--cpus / --memory). Cannot be exceeded."],
        ["Filesystem", "Hard", "Docker read-only image layers + overlay. User /home on NFS."],
        ["Process", "Strong", "PID namespaces. Users cannot see each other's processes."],
        ["Network", "Strong", "Network namespaces + VLAN segmentation."],
        ["GPU VRAM", "Enforced (software)", "HAMi-core + MPS dual-layer. CUDA_ERROR_OUT_OF_MEMORY on exceed."],
        ["GPU Compute (SM)", "Enforced (software)", "HAMi-core + MPS ACTIVE_THREAD_PERCENTAGE."],
        ["GPU Address Space", "Isolated", "MPS Volta+ per-client isolated GPU virtual address spaces."],
        ["GPU L2 / Bandwidth", "Shared", "Cannot partition on consumer GPUs. 10–30% variability possible."],
        ["GPU Fatal Fault", "Propagates", "MPS auto-recovers. Watchdog restarts containers in 30–60s."],
        ["Base Software", "Absolute", "Docker read-only layers. No root/sudo. Immutable."],
    ],
    col_widths=[1.5, 1.5, 3.6]
)

add_callout_box(doc,
    "CONSIDERATION — GPU Noisy Neighbour & Fault Propagation",
    "Users on shared GPU tiers (Pro/Power/Max) may experience 10–30% performance variability "
    "under peak co-tenancy due to shared L2 cache and memory bandwidth. A GPU fatal fault from "
    "one user's CUDA operation will affect co-resident containers — MPS auto-recovers in 30–60s "
    "with no data loss (user files are on NAS). Full Machine tier is immune to both effects.\n\n"
    "Recommended ToS Language: \"GPU Desktop sessions share a physical GPU for cost efficiency. "
    "In rare circumstances (<0.1%), a GPU interruption may require a 30–60 second automatic "
    "restart. Your files are never affected. For guaranteed uninterrupted GPU access, use Full Machine.\"",
    CAUTION_BG, CAUTION_BORDER
)


# ─── 4. Storage Architecture ───
doc.add_heading('4. Storage Architecture', level=2)

doc.add_heading('4.1 Per-Node NVMe Partitioning (2TB)', level=3)

add_styled_table(doc,
    ["Partition", "Size", "Filesystem", "Purpose"],
    [
        ["nvme0n1p1", "128 GB", "EXT4", "Proxmox OS + swap (8 GB)"],
        ["nvme0n1p2", "500 GB", "EXT4", "/var/lib/docker — images, container overlays"],
        ["nvme0n1p3", "1,100 GB", "LVM PV", "VG_CONTAINERS — ephemeral scratch, wiped on session end"],
        ["nvme0n1p4", "240 GB", "EXT4", "/shared — read-only datasets, software caches"],
    ],
    col_widths=[1.2, 0.9, 1.0, 3.5]
)

add_callout_box(doc,
    "HARD DATA ISOLATION",
    "Docker storage (p2) and VG_CONTAINERS (p3) are on separate NVMe partitions. Even if "
    "VG_CONTAINERS is completely filled or corrupted, Docker storage is untouched. User persistent "
    "data lives on the NAS via NFS — completely independent of both local partitions. Ephemeral "
    "session corruption can never reach user data or base images.",
    INFO_BG, INFO_BORDER
)

doc.add_heading('4.2 Centralised NAS (5th Machine)', level=3)

add_bullet(doc, "TrueNAS Scale (Debian-based, ZFS native) — recommended over Synology for flexibility")
add_bullet(doc, "ZFS Pool: 4×4TB HDDs in RAIDZ1 (~12TB usable)")
add_bullet(doc, "Per-user ZFS dataset with 15GB hard quota: datapool/users/<uid>")
add_bullet(doc, "NFS export over 10GbE to all 4 compute nodes (nconnect=4 for ~3–4 Gbps throughput)")
add_bullet(doc, "Automated ZFS snapshots every 6 hours (7 daily + 4 weekly retained)")
add_bullet(doc, "Users can self-recover files via .zfs/snapshot/ directory")

# ─── 5. Networking ───
doc.add_heading('5. Networking Architecture', level=2)

add_styled_table(doc,
    ["VLAN", "Subnet", "Purpose", "Interface"],
    [
        ["10 (Management)", "10.10.10.0/24", "Proxmox Web UI, SSH, cluster comms", "Onboard 2.5GbE"],
        ["20 (Stateful)", "10.10.20.0/24", "Stateful desktop container traffic, WebRTC", "10GbE"],
        ["30 (Ephemeral)", "10.10.30.0/24", "Jupyter, Code-Server, SSH sessions", "10GbE"],
        ["40 (Storage)", "10.10.40.0/24", "NFS traffic (NAS ↔ Nodes). No internet.", "10GbE"],
        ["50 (Services)", "10.10.50.0/24", "Portal, Keycloak, Reverse Proxy", "10GbE"],
    ],
    col_widths=[1.4, 1.3, 2.2, 1.3]
)

doc.add_paragraph()

add_bullet(doc, "External Access: Cloudflare Tunnel (browser, zero open ports) + Tailscale (SSH, Moonlight)")
add_bullet(doc, "Internet: Minimum 500 Mbps symmetric upstream. Dual ISP recommended for redundancy.")
add_bullet(doc, "TURN Server: coturn deployed for WebRTC NAT traversal through university firewalls.")

add_callout_box(doc,
    "BLOCKER — ISP Verification Required",
    "Verify that the ISP at the deployment location can provide ≥500 Mbps symmetric upstream. "
    "Obtain quotes from 2 ISPs. Budget ₹30,000–80,000/month for business-grade symmetric fibre. "
    "Single ISP failure = platform offline for all remote users.",
    CRITICAL_BG, CRITICAL_BORDER, RGBColor(0xDC, 0x26, 0x26)
)

# ─── 6. Phase 0 Burn-In ───
doc.add_heading('6. Phase 0 — Burn-In Protocol (Weeks 1–3)', level=2)

add_para(doc,
    "Phase 0 is the decision gate that validates the chosen architecture before fleet-wide "
    "deployment. All testing is performed on a single node first."
)

add_styled_table(doc,
    ["Week", "Focus Area", "Key Activities"],
    [
        ["Week 1", "Hardware Validation", "BIOS config, memtest86+ (48h), Proxmox install, cluster formation, "
         "10GbE + VLAN setup, NAS ZFS + NFS configuration, NFS I/O benchmark (fio)"],
        ["Week 2", "Architecture A Burn-in", "NVIDIA driver on host, Docker + nvidia-container-toolkit, "
         "Selkies EGL containers, HAMi-core + MPS integration, multi-container GPU sharing, "
         "NVENC concurrent stream testing (4/6/8 sessions), nvidia-patch for NVENC limits"],
        ["Week 3", "Stress & Decision Gate", "48-hour stress test (4 containers with fractional VRAM), "
         "fault injection (deliberate GPU crash → verify MPS recovery 10×), MPS watchdog service, "
         "DECISION: confirm Architecture A or fall back to B"],
    ],
    col_widths=[0.8, 1.8, 4.0]
)

add_callout_box(doc,
    "DECISION GATE CRITERIA",
    "Architecture A is confirmed production-ready if:\n"
    "✓ All 4 fractional VRAM combinations work (4+8+8+8, 4+4+8+16, 2×16, etc.)\n"
    "✓ 48-hour stress test passes without crashes or VRAM limit violations\n"
    "✓ MPS fault recovery succeeds in 10/10 fault injection tests\n"
    "✓ NVENC supports ≥4 concurrent desktop streams\n"
    "✓ CPU-only containers stream smoothly via WebRTC\n\n"
    "If any test fails → investigate root cause → attempt fix → if unresolvable, "
    "fall back to Architecture B (VM passthrough, one GPU per VM, no fractional tiers).",
    CAUTION_BG, CAUTION_BORDER
)

# ─── 7. Risks ───
doc.add_heading('7. Infrastructure Risks & Considerations', level=2)

add_styled_table(doc,
    ["Risk", "Severity", "Mitigation"],
    [
        ["RTX 5090 D3cold Reset Bug", "HIGH",
         "Eliminated by Architecture A (GPU never changes ownership). Relevant only if Architecture B fallback is used."],
        ["GPU Noisy Neighbour", "MEDIUM",
         "10–30% variability disclosed in ToS. Full Machine tier offered for guaranteed performance."],
        ["GPU Fatal Fault Propagation", "MEDIUM",
         "MPS auto-recovery + watchdog service + admin alerts. User data safe on NAS."],
        ["Ryzen 9950X3D + Proxmox Stability", "MEDIUM",
         "Disable XMP during burn-in. Update BIOS + AMD microcode. Pin to stable kernel 6.8.x LTS."],
        ["NVENC Session Limit (Consumer)", "MEDIUM",
         "nvidia-patch removes firmware limit. Must verify on RTX 5090 in Phase 0."],
        ["MATLAB Network Licensing", "HIGH (Blocker)",
         "Must resolve before building base images. Per-seat license is NOT valid for shared hosting."],
        ["ISP Bandwidth Insufficiency", "HIGH (Blocker)",
         "Verify ≥500 Mbps symmetric. Obtain quotes from 2 ISPs before launch."],
        ["Thermal Management (24/7 Operation)", "MEDIUM",
         "DCGM exporter monitors temp. Alert at 80°C, emergency at 90°C. Consider Noctua industrial fans."],
        ["Crypto Mining Abuse", "LOW–MEDIUM",
         "GPU utilisation pattern monitoring. Block mining pool IPs. Kill sustained >95% GPU with no CUDA context."],
        ["Single NVMe Failure", "HIGH",
         "User data on NAS (ZFS snapshots). Local NVMe is ephemeral/rebuild-able. No single point of data loss."],
    ],
    col_widths=[1.8, 1.0, 3.8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PART II — SOFTWARE DEVELOPMENT LIFECYCLE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('PART II — Software Development Lifecycle & Sprint Plan', level=1)

doc.add_heading('8. Sprint Timeline Overview', level=2)

add_callout_box(doc,
    "ACCELERATED TIMELINE — Speed to Market",
    "The AI infrastructure market is consolidating rapidly. TCS + OpenAI + AMD are entering the "
    "Indian data centre space. First-mover advantage is critical. The following timeline reflects "
    "the business reality: launch by mid-March 2026 or risk strategic irrelevance.\n\n"
    "Development Sprint: Today → Friday, March 13, 2026\n"
    "QA & Load Testing: March 15\n"
    "Soft Launch (Closed Beta): March 16\n"
    "Public Launch: March 17",
    CRITICAL_BG, CRITICAL_BORDER, RGBColor(0xDC, 0x26, 0x26)
)

add_styled_table(doc,
    ["Day", "Date", "Sprint Focus", "Deliverables"],
    [
        ["Day 1", "Mar 08", "Foundation & Auth",
         "Project scaffolding, DB schema, Keycloak/Auth integration, User registration + email verification"],
        ["Day 2", "Mar 09", "Core Platform",
         "Role-based access control, Profile management, Dashboard layout, Wallet system backend"],
        ["Day 3", "Mar 10", "Booking Engine",
         "GPU slot booking workflow, Real-time availability dashboard, Resource allocation logic, Queue management"],
        ["Day 4", "Mar 11", "Payments & Mentorship",
         "Stripe/Razorpay integration, Subscription tiers (Bronze/Silver/Gold), Mentor booking system, Calendar scheduling"],
        ["Day 5", "Mar 12", "Admin & Analytics",
         "Admin dashboard, User management panel, Usage analytics, Hardware monitoring views, Invoice generation"],
        ["Day 6", "Mar 13", "Polish & Chatbot",
         "AI Chatbot (FAQ MVP), Responsive UI polish, Dark mode, Performance optimisation (<2s load), Integration testing"],
    ],
    col_widths=[0.6, 0.8, 1.5, 3.7]
)

# ─── 9. Tech Stack ───
doc.add_heading('9. Technology Stack', level=2)

add_styled_table(doc,
    ["Layer", "Technology", "Rationale"],
    [
        ["Frontend", "Next.js (React) + TypeScript",
         "SSR for SEO, app router, mobile-responsive. Futuristic AI-themed design with dark/light mode."],
        ["Backend API", "Python FastAPI",
         "Async-native, auto-generated OpenAPI docs, <500ms p95 response time."],
        ["State Store", "Redis",
         "Session state, real-time fleet status, booking lock management, pub/sub for live updates."],
        ["Primary Database", "PostgreSQL",
         "Users, bookings, billing records, audit logs. Optimised for read-heavy dashboard queries."],
        ["Authentication", "Keycloak (OIDC/SAML)",
         "SSO with university Google Workspace/LDAP. MFA enforced. Role-based groups."],
        ["Payment Gateway", "Razorpay (primary) + Stripe (fallback)",
         "Indian rupee support, UPI, auto-debit. Stripe for international/corporate users."],
        ["Container Orchestration", "Docker API (via FastAPI orchestrator)",
         "Direct Docker API calls for session lifecycle. No Kubernetes overhead for 4-node scale."],
        ["Monitoring", "Prometheus + Grafana + DCGM Exporter",
         "Node metrics, GPU health, per-user usage, billing analytics."],
        ["Remote Desktop Streaming", "Selkies-GStreamer (WebRTC)",
         "Browser-native, no client install. NVENC H.264 encoding. 60fps at 1080p."],
        ["Reverse Proxy", "Traefik / Caddy",
         "TLS termination, WebSocket passthrough, path-based routing."],
        ["Remote Access", "Cloudflare Tunnel + Tailscale",
         "Zero open ports. Works through university firewalls."],
    ],
    col_widths=[1.3, 2.2, 3.1]
)

# ─── 10. Module Breakdown ───
doc.add_heading('10. Platform Module Breakdown', level=2)

doc.add_heading('10.1 User Authentication & Onboarding', level=3)
add_bullet(doc, "Signup with email + OTP verification (university email for members, Google OAuth for public)")
add_bullet(doc, "Keycloak-federated SSO (university Google Workspace or LDAP)")
add_bullet(doc, "Role assignment: Student, Faculty, Admin, Corporate Partner, Public User")
add_bullet(doc, "2FA enforced via TOTP (Google Authenticator)")
add_bullet(doc, "Profile setup: name, department, college, semester, preferences")
add_bullet(doc, "Freemium tier auto-assignment for KSRCE students")

doc.add_heading('10.2 GPU Booking & Resource Management', level=3)
add_bullet(doc, "Real-time availability dashboard showing fleet resource status per node")
add_bullet(doc, "Pre-defined compute config selection (Starter through Full Machine)")
add_bullet(doc, "Time-slot booking with 1-hour blocks (now or scheduled)")
add_bullet(doc, "Auto-locking of GPU nodes upon confirmed booking")
add_bullet(doc, "Queue management for high-demand periods")
add_bullet(doc, "Session launch → container spin-up → WebRTC URL delivery (~15–20s)")
add_bullet(doc, "Session reconnection if browser closes mid-session")
add_bullet(doc, "Graceful shutdown with 15-min and 5-min warnings")
add_bullet(doc, "Over-usage handling: prompt to extend or auto-terminate")

doc.add_heading('10.3 Wallet & Payment System', level=3)
add_bullet(doc, "Wallet top-up via Razorpay (UPI, cards, net banking)")
add_bullet(doc, "Real-time wallet deduction during active sessions")
add_bullet(doc, "Subscription tiers: Bronze (₹500/20h), Silver (₹1,500/60h), Gold (₹3,000/150h)")
add_bullet(doc, "Academic Integration Fees: per-section per-semester billing to college")
add_bullet(doc, "Automated low-balance notifications (web push, email, SMS)")
add_bullet(doc, "Recharge packages with bonus minutes")
add_bullet(doc, "Invoice generation and transaction history")

doc.add_heading('10.4 Mentor Booking System', level=3)
add_bullet(doc, "Mentor profile creation (expertise, experience, ratings)")
add_bullet(doc, "Availability calendar with configurable time slots")
add_bullet(doc, "Session booking with automatic calendar scheduling")
add_bullet(doc, "Integrated video call (Jitsi Meet or similar)")
add_bullet(doc, "Booking confirmation + reminder notifications")
add_bullet(doc, "Automated payment and invoicing for mentor sessions")

doc.add_heading('10.5 Admin Dashboard & Analytics', level=3)
add_bullet(doc, "Real-time fleet overview: CPU/RAM/GPU utilisation per node")
add_bullet(doc, "User management: CRUD, role assignment, access revocation")
add_bullet(doc, "Session monitoring: active sessions, duration, resource consumption")
add_bullet(doc, "Revenue analytics: total, per-tier, per-college, per-time-period")
add_bullet(doc, "GPU health: temperature, VRAM, fault events, uptime")
add_bullet(doc, "Audit log: all user actions, login events, session lifecycle events")
add_bullet(doc, "Hardware status: NVMe health (SMART), NAS capacity, network throughput")

doc.add_heading('10.6 AI Chatbot (MVP)', level=3)
add_bullet(doc, "FAQ-driven chatbot for common queries (pricing, booking, troubleshooting)")
add_bullet(doc, "Escalation workflow to human support")
add_bullet(doc, "Deployed on the web interface (floating widget)")
add_bullet(doc, "Phase 2: context-aware, integrated with booking and billing systems")

doc.add_heading('10.7 Responsive Web UI', level=3)
add_bullet(doc, "Mobile-first, dark-mode-first aesthetic (futuristic AI theme)")
add_bullet(doc, "Fast load times (<2s target). Lighthouse score ≥90.")
add_bullet(doc, "Neon accents, soft gradients, animated circuit board patterns")
add_bullet(doc, "Dashboard-driven layout with data visualisation")
add_bullet(doc, "Dark and light mode toggle")

# ─── 11. Out of Scope ───
doc.add_heading('11. V1 Scope Exclusions (Phase 2 Features)', level=2)

add_callout_box(doc,
    "EXPLICITLY OUT OF SCOPE FOR WEEK 1 SPRINT",
    "• Collaborative coding environments (JupyterHub multi-user)\n"
    "• Mobile app (iOS / Android)\n"
    "• Data annotation services platform\n"
    "• AI model marketplace\n"
    "• Predictive maintenance system\n"
    "• Advanced gamification features\n"
    "• Learning path recommendation engine\n"
    "• Project showcase portal\n"
    "• Industry projects marketplace\n\n"
    "These will be phased in post-launch as Version 2.0 features.",
    INFO_BG, INFO_BORDER
)

# ─── 12. Integrations ───
doc.add_heading('12. Third-Party Dependencies & Integrations', level=2)

add_styled_table(doc,
    ["Integration", "Purpose", "Dependency Type", "Risk"],
    [
        ["Keycloak", "SSO / Identity Provider", "Self-hosted (Docker)", "Low — mature, well-documented"],
        ["Razorpay", "Payment processing (INR)", "External API", "Low — widely used in India"],
        ["Stripe", "Payment (international / fallback)", "External API", "Low"],
        ["Cloudflare", "Tunnel + DDoS protection", "External service", "Low — free tier sufficient"],
        ["Tailscale", "VPN for SSH / Moonlight users", "External + self-hosted", "Low"],
        ["Selkies-GStreamer", "WebRTC desktop streaming", "Open-source (Docker)", "Medium — niche project"],
        ["HAMi-core", "GPU VRAM enforcement", "Open-source (build from source)", "Medium — requires Phase 0 validation"],
        ["CUDA MPS", "GPU compute partitioning", "NVIDIA (built into CUDA)", "Low — official NVIDIA component"],
        ["Prometheus + Grafana", "Monitoring & dashboards", "Self-hosted", "Low — industry standard"],
        ["coturn", "WebRTC TURN relay", "Self-hosted", "Low — handles NAT traversal"],
        ["MATLAB (Network License)", "Proprietary software in base image", "Client/University procured", "HIGH — licensing blocker"],
    ],
    col_widths=[1.3, 1.8, 1.5, 2.0]
)

# ─── 13. Quality Standards ───
doc.add_heading('13. Development Standards', level=2)

add_styled_table(doc,
    ["Aspect", "Standard"],
    [
        ["Security", "All user data encrypted at rest (LUKS on NAS) and in transit (TLS 1.3). "
         "No hardcoded credentials. JWT-based APIs. RBAC enforced."],
        ["Performance", "API response times <500ms (p95). Database optimised for read-heavy dashboards. "
         "Frontend <2s load time (Lighthouse ≥90)."],
        ["Availability", "99% uptime target. Automated backups (ZFS snapshots). Disaster recovery plan."],
        ["Testing", "Critical paths (booking, payment, auth) have automated tests. Manual QA for UI/UX."],
        ["Documentation", "API documentation (auto-generated via FastAPI). Deployment runbook. Onboarding guide."],
        ["Data Protection", "DPDPA 2023 compliance. Consent, data minimisation, breach notification procedures."],
    ],
    col_widths=[1.3, 5.3]
)

# ─── 14. Deployment ───
doc.add_heading('14. Deployment & Launch Plan', level=2)

add_styled_table(doc,
    ["Date", "Milestone", "Details"],
    [
        ["Mar 08–13", "Development Sprint", "All 7 core modules built, tested, integration-verified."],
        ["Mar 15", "QA & Load Testing", "Automated test suite, load simulation (100 concurrent users), "
         "security audit (OWASP top 10 checks)."],
        ["Mar 16", "Soft Launch (Closed Beta)", "KSRCE faculty + 20–30 select students. Monitor booking flow, "
         "payment processing, session stability."],
        ["Mar 17", "Public Launch", "Marketing push, open registration. Gradual ramp-up to full capacity."],
        ["Mar 18–31", "Stabilisation", "Bug fixes from beta feedback. Performance tuning. Monitoring alert refinement."],
        ["Apr 01+", "Phase 2 Features", "V2 roadmap: gamification, marketplace, collaborative coding, mobile app."],
    ],
    col_widths=[1.0, 1.8, 3.8]
)

# ─── 15. Infrastructure Timeline (Full) ───
doc.add_heading('15. Full Infrastructure Timeline', level=2)

add_para(doc,
    "The infrastructure and software timelines run in parallel. Infrastructure Phases 0–1 must "
    "complete before the platform can be fully operational, but software development can proceed "
    "concurrently using mock/simulated backends."
)

add_styled_table(doc,
    ["Phase", "Duration", "Focus"],
    [
        ["Phase 0: Burn-in", "Weeks 1–3",
         "Hardware assembly, Proxmox install, GPU burn-in (both architectures), NAS setup, networking, DECISION GATE"],
        ["Phase 1: Core Infra", "Weeks 3–6",
         "Apply architecture to all nodes, build Selkies Docker image, NFS user storage automation, Keycloak deployment"],
        ["Phase 2: Orchestration", "Weeks 6–10",
         "FastAPI orchestrator, Selkies WebRTC integration, Cloudflare Tunnel, booking system, container lifecycle"],
        ["Phase 3: Portal", "Weeks 10–14",
         "Next.js frontend, Keycloak auth integration, admin panel, user registration flow"],
        ["Phase 4: Billing & Monitoring", "Weeks 14–18",
         "Razorpay billing, Prometheus + Grafana + DCGM, audit logging, session warnings, anti-abuse"],
        ["Phase 5: Beta", "Weeks 18–22",
         "10–20 beta users, load testing, security audit, documentation"],
        ["Phase 6: Launch", "Week 22+",
         "Production launch, gradual university rollout, public user access"],
    ],
    col_widths=[1.6, 1.0, 4.0]
)

add_callout_box(doc,
    "CONSIDERATION — Timeline Alignment",
    "The software sprint (Mar 08–17) delivers the web platform MVP. The full infrastructure "
    "timeline (22+ weeks) is for the complete on-premises LaaS stack. For the March 17 launch, "
    "the platform can operate in a limited-capacity mode while Phase 0–1 infrastructure burn-in "
    "completes. Critical path: authentication + booking + payment must be fully functional by launch. "
    "GPU session delivery depends on infrastructure readiness.",
    CAUTION_BG, CAUTION_BORDER
)

doc.add_page_break()

# ─── Appendix ───
doc.add_heading('Appendix', level=1)

doc.add_heading('A. Revenue Model Summary', level=2)

add_styled_table(doc,
    ["Revenue Stream", "Description", "Pricing Model"],
    [
        ["Academic Integration Fees", "Per-section per-semester LaaS fee bundling GPU-hours, content, support", "Pre-paid by college"],
        ["Premium Student Passes", "Bronze / Silver / Gold tiered packages", "₹500 / ₹1,500 / ₹3,000"],
        ["Pay-as-you-go Wallet", "Real-time GPU consumption at standard rate", "₹100/hr (adjustable)"],
        ["Partner College Subscriptions", "Annual subscription for external colleges", "Custom per-college"],
        ["Mentor Booking Fees", "One-on-one guidance sessions", "Tiered by mentor experience"],
    ],
    col_widths=[1.8, 2.8, 2.0]
)

doc.add_heading('B. Security Checklist', level=2)

add_bullet(doc, "OTP + 2FA for all user logins")
add_bullet(doc, "JWT-based API authentication with short-lived tokens")
add_bullet(doc, "RBAC enforced at API and UI level")
add_bullet(doc, "HTTPS everywhere (TLS 1.3)")
add_bullet(doc, "Data encrypted at rest (LUKS on NAS, ZFS encryption)")
add_bullet(doc, "DPDPA 2023 compliance (consent, data minimisation, breach notification)")
add_bullet(doc, "Session idle timeout: 30min warning → 45min suspend → 60min terminate")
add_bullet(doc, "Container security: no-new-privileges, read-only rootfs, pids-limit, seccomp, AppArmor")
add_bullet(doc, "Anti-abuse: mining detection, egress rate limiting, mining pool IP blocking")
add_bullet(doc, "Fail2Ban on SSH, Keycloak brute-force protection")

doc.add_heading('C. Glossary', level=2)

add_styled_table(doc,
    ["Term", "Definition"],
    [
        ["LaaS", "Lab-as-a-Service — delivering lab compute infrastructure as a remotely accessible, on-demand service"],
        ["Selkies EGL", "Open-source project for GPU-accelerated Linux desktops inside Docker containers, streamed via WebRTC"],
        ["HAMi-core", "Open-source CUDA API interception library for enforcing per-container VRAM and compute limits"],
        ["CUDA MPS", "NVIDIA Multi-Process Service — enables multiple CUDA contexts to share a GPU with SM partitioning"],
        ["vGPU", "NVIDIA Virtual GPU — hardware GPU slicing available only on enterprise/professional GPUs"],
        ["MIG", "Multi-Instance GPU — hardware GPU partitioning available only on A100/H100/B200"],
        ["D3cold", "Deep PCIe power state that causes the RTX 5090 reset bug in VM passthrough scenarios"],
        ["WebRTC", "Web Real-Time Communication — browser-native protocol for low-latency video/audio streaming"],
        ["NVENC", "NVIDIA hardware video encoder used for desktop stream encoding"],
        ["ZFS", "Advanced filesystem with snapshots, checksums, RAIDZ, and compression"],
        ["cgroups v2", "Linux kernel feature for resource limitation and isolation (CPU, memory, I/O)"],
    ],
    col_widths=[1.3, 5.3]
)

# ─── Footer ───
doc.add_paragraph()
add_divider(doc)
add_para(doc,
    "This document is confidential to Global Knowledge Technologies and KSRCE stakeholders. "
    "Version 2.0 — March 2026. Subject to revision as infrastructure burn-in results become available.",
    italic=True, color=MID_GRAY, font_size=8.5,
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)

# ─── SAVE ───
doc.save(OUTPUT_PATH)
print(f"\n✅ Document saved to: {OUTPUT_PATH}")
print(f"   File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
