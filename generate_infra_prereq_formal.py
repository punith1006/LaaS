"""
Generate LaaS Infrastructure Prerequisites & Requirements document (.docx)
for KSRCE Infrastructure Team — formal technical document style.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"c:\Users\Punith\Downloads\LaaS_Infrastructure_Prerequisites_KSRCE - Copy.docx"

# ── Colour palette ──────────────────────────────────────────────────────────
HEADING_BLUE  = RGBColor(0x2E, 0x50, 0x90)
DARK_TEXT     = RGBColor(0x1A, 0x1A, 0x2E)
MID_TEXT      = RGBColor(0x3A, 0x3A, 0x3A)
GREY_TEXT     = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY    = RGBColor(0x88, 0x88, 0x88)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)


# ── Low-level XML helpers ────────────────────────────────────────────────────

def set_para_shading(paragraph, hex_color: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_para_border(paragraph, color: str = "2E5090", sz: str = "6",
                    space: str = "4", val: str = "single", side: str = "left"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), val)
    bdr.set(qn("w:sz"), sz)
    bdr.set(qn("w:space"), space)
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)
    pPr.append(pBdr)


def add_page_numbers(doc: Document):
    """Add 'Page X of Y' footer to every section."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("Page ")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = LIGHT_GREY

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText1 = OxmlElement("w:instrText")
        instrText1.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "begin")
        instrText2 = OxmlElement("w:instrText")
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")

        r = para.add_run()
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = LIGHT_GREY
        r._r.append(fldChar1)
        r._r.append(instrText1)
        r._r.append(fldChar2)

        r2 = para.add_run(" of ")
        r2.font.name = "Calibri"
        r2.font.size = Pt(9)
        r2.font.color.rgb = LIGHT_GREY

        r3 = para.add_run()
        r3.font.name = "Calibri"
        r3.font.size = Pt(9)
        r3.font.color.rgb = LIGHT_GREY
        r3._r.append(fldChar3)
        r3._r.append(instrText2)
        r3._r.append(fldChar4)

        r4 = para.add_run("  |  LaaS Platform — Infrastructure Prerequisites & Requirements")
        r4.font.name = "Calibri"
        r4.font.size = Pt(9)
        r4.font.color.rgb = LIGHT_GREY


# ── Style helpers ────────────────────────────────────────────────────────────

def set_run_font(run, bold=False, italic=False, size=11, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    sizes = {1: 15, 2: 12, 3: 11}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = HEADING_BLUE
    if level == 1:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bdr = OxmlElement("w:bottom")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "6")
        bdr.set(qn("w:space"), "1")
        bdr.set(qn("w:color"), "2E5090")
        pBdr.append(bdr)
        pPr.append(pBdr)


def add_body(doc: Document, text: str, bold: bool = False, italic: bool = False,
             indent: bool = False, color=None) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(3)
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(text)
    set_run_font(run, bold=bold, italic=italic, color=color)


def add_bullet(doc: Document, text: str, level: int = 0,
               bold_prefix: str = None, italic_suffix: bool = False) -> None:
    """
    Renders a bullet as a plain paragraph with a dash prefix
    to avoid List Bullet style indent unpredictability.
    level 0 → single bullet  (•  prefix, slight indent)
    level 1 → sub-bullet     (–  prefix, more indent)
    """
    bullet_char = "•" if level == 0 else "–"
    indent_val  = Inches(0.3 + level * 0.25)

    para = doc.add_paragraph()
    para.paragraph_format.space_before    = Pt(2)
    para.paragraph_format.space_after     = Pt(2)
    para.paragraph_format.left_indent     = indent_val
    para.paragraph_format.first_line_indent = Inches(-0.22)

    # bullet glyph
    r0 = para.add_run(f"{bullet_char}  ")
    set_run_font(r0, color=HEADING_BLUE if level == 0 else MID_TEXT)

    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = para.add_run(text)
        set_run_font(r2, italic=italic_suffix)
    else:
        r = para.add_run(text)
        set_run_font(r, italic=italic_suffix)


def add_checklist_item(doc: Document, text: str) -> None:
    """Renders a checklist bullet with a dash (–) prefix."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before    = Pt(2)
    para.paragraph_format.space_after     = Pt(2)
    para.paragraph_format.left_indent     = Inches(0.3)
    para.paragraph_format.first_line_indent = Inches(-0.22)

    r0 = para.add_run("–  ")
    set_run_font(r0, color=HEADING_BLUE)
    r = para.add_run(text)
    set_run_font(r)


def add_callout(doc: Document, label: str, text: str,
                callout_type: str = "consideration") -> None:
    """Shaded callout box with left border stripe."""
    bg_color     = "E8EFF8" if callout_type == "consideration" else "FFF0E8"
    border_color = "2E5090" if callout_type == "consideration" else "C0392B"
    label_color  = HEADING_BLUE if callout_type == "consideration" \
                   else RGBColor(0xC0, 0x39, 0x2B)

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(10)
    para.paragraph_format.left_indent  = Inches(0.2)
    para.paragraph_format.right_indent = Inches(0.2)
    set_para_shading(para, bg_color)
    set_para_border(para, color=border_color, sz="12", space="8")

    r1 = para.add_run(f"  {label}:  ")
    set_run_font(r1, bold=True, size=10, color=label_color)
    r2 = para.add_run(text)
    set_run_font(r2, size=10)


def add_spacer(doc: Document, size: int = 4) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run("")
    run.font.size = Pt(size)


def add_numbered_item(doc: Document, text: str, number: int) -> None:
    """Renders a numbered list item."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.3)

    r0 = para.add_run(f"{number}. ")
    set_run_font(r0, bold=True, color=HEADING_BLUE)
    r = para.add_run(text)
    set_run_font(r)


# ── Main document builder ────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── Default paragraph style ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════════════════════
    add_spacer(doc, 28)

    # Decorative top bar
    bar = doc.add_paragraph()
    set_para_shading(bar, "2E5090")
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after  = Pt(0)
    r = bar.add_run("   ")
    r.font.size = Pt(18)

    add_spacer(doc, 18)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after  = Pt(4)
    title_run = title_para.add_run("LaaS Platform")
    title_run.font.name  = "Calibri"
    title_run.font.size  = Pt(28)
    title_run.bold       = True
    title_run.font.color.rgb = HEADING_BLUE

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub1.paragraph_format.space_before = Pt(0)
    sub1.paragraph_format.space_after  = Pt(4)
    r = sub1.add_run("Infrastructure Prerequisites & Requirements")
    r.font.name  = "Calibri"
    r.font.size  = Pt(20)
    r.bold       = True
    r.font.color.rgb = MID_TEXT

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub2.paragraph_format.space_before = Pt(4)
    sub2.paragraph_format.space_after  = Pt(2)
    r = sub2.add_run("For KSRCE Infrastructure Team")
    r.font.name  = "Calibri"
    r.font.size  = Pt(14)
    r.font.color.rgb = GREY_TEXT

    add_spacer(doc, 10)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_para.paragraph_format.space_before = Pt(0)
    meta_para.paragraph_format.space_after  = Pt(0)
    r = meta_para.add_run("Version 1.0  |  April 2026  |  Prepared by: LaaS Platform Team")
    r.font.name   = "Calibri"
    r.font.size   = Pt(10)
    r.font.color.rgb = GREY_TEXT

    add_spacer(doc, 26)

    bar2 = doc.add_paragraph()
    set_para_shading(bar2, "2E5090")
    bar2.paragraph_format.space_before = Pt(0)
    bar2.paragraph_format.space_after  = Pt(0)
    r = bar2.add_run("   ")
    r.font.size = Pt(8)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "EXECUTIVE SUMMARY")

    add_body(doc, "This document defines the infrastructure prerequisites and environmental requirements for deploying the LaaS (Lab-as-a-Service) GPU computing platform at KSRCE. Compliance with all mandatory requirements is necessary before the LaaS engineering team can commence remote software deployment and configuration.")

    add_spacer(doc, 6)

    # Critical Requirements
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(4)
    r = para.add_run("CRITICAL REQUIREMENTS (Must be completed before deployment):")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    critical_requirements = [
        "Ubuntu 22.04.5 LTS installed on all compute nodes with SSH access configured",
        "Network-attached storage (NAS) with minimum 4TB capacity, NFS-enabled",
        "Static IP addressing for all infrastructure components",
        "Minimum 1 Gbps symmetric internet connectivity",
        "VPN access provisioned for the LaaS engineering team",
        "Dedicated application server (minimum 8 vCPU, 16 GB RAM, 256 GB SSD)",
        "UPS protection for all critical infrastructure (minimum 5 kVA)",
        "Adequate cooling maintaining 18-24°C ambient temperature",
        "Backup storage allocation (minimum 500 GB) for database and snapshot retention",
        "Designated infrastructure support contact with <4 hour response SLA"
    ]

    for i, req in enumerate(critical_requirements, 1):
        add_numbered_item(doc, req, i)

    add_spacer(doc, 6)

    # Key Recommendations
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(4)
    r = para.add_run("KEY RECOMMENDATIONS:")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = HEADING_BLUE

    recommendations = [
        "10 Gbps internal network links between compute nodes and NAS",
        "16 vCPU / 32 GB RAM / 500 GB SSD for the application server",
        "Redundant internet connectivity via secondary ISP",
        "Offsite backup replication to cloud storage",
        "IPMI/BMC access for remote hardware management",
        "Server room compliance with ODC standards"
    ]

    for rec in recommendations:
        add_bullet(doc, rec)

    add_spacer(doc, 6)

    # Handover Note
    add_body(doc, "All requirements in this document must be fulfilled and verified prior to handover. The LaaS engineering team will perform all software installation, service configuration, and platform deployment remotely via SSH.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION — Document Purpose
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "Document Purpose & Scope")
    add_bullet(doc, "This document enumerates all infrastructure prerequisites required from the KSRCE infrastructure team prior to deployment of the LaaS platform software stack.")
    add_bullet(doc, "Upon fulfillment of all prerequisites and provision of remote access credentials, the LaaS engineering team shall assume responsibility for all software installation, configuration, and deployment activities, which will be executed entirely via remote SSH access.")
    add_bullet(doc, "This document serves as a formal pre-deployment checklist and verification framework for both parties.")

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1 — OS & Machine Setup
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Operating System & Machine Setup")

    add_bullet(doc, "All compute nodes shall have Ubuntu 22.04.5 LTS (Jammy Jellyfish) installed as the base operating system.",
               bold_prefix="Mandatory: ")
    add_bullet(doc, "The installation configuration shall be one of the following:")
    add_bullet(doc, "Standalone Linux installation — designated exclusively for LaaS workloads.", level=1)
    add_bullet(doc, "Dual-boot configuration (Windows + Linux) — for shared utilization per institutional requirements.", level=1)

    add_bullet(doc, "For dual-boot configurations, the following requirements apply:")
    add_bullet(doc, "The system BIOS shall be configured to boot into Linux by default.", level=1)
    add_bullet(doc, "A minimum of 80% of available disk space shall be allocated to the Linux partition to accommodate workload storage requirements.", level=1)

    add_bullet(doc, "All compute nodes shall be configured for continuous 24/7 operation without automatic sleep, hibernation, or shutdown behavior.")
    add_bullet(doc, "The BIOS shall be configured with 'AC Power Recovery → Power On' to enable automatic system boot following power restoration without requiring manual intervention.")

    add_bullet(doc, "SSH server configuration requirements:",
               bold_prefix="")
    add_bullet(doc, "OpenSSH server shall be installed and enabled on all compute nodes.", level=1)
    add_bullet(doc, "The SSH daemon (sshd) must be configured to start automatically on system boot via systemd.", level=1)
    add_bullet(doc, "A dedicated administrative user account (recommended: laas-admin) with sudo privileges shall be created for remote administrative access.", level=1)

    add_bullet(doc,
               "Each machine shall be assigned a descriptive hostname following a consistent naming convention — "
               "for example: laas-node-01, laas-node-02, laas-node-03, laas-node-04 for compute nodes, and laas-nas for the storage server.")

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2 — Network Storage (NAS)
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Network Storage (NAS)")

    add_bullet(doc, "A minimum of 4TB of network-attached storage (NAS) shall be provisioned and made accessible to all compute nodes via the local network.",
               bold_prefix="Mandatory: ")
    add_bullet(doc, "The NAS system must support NFS (Network File System) protocol for read/write operations from compute nodes.")
    add_bullet(doc, "All compute nodes shall have read/write access to the NAS without authentication or permission impediments — this shall be verified post-configuration.")
    add_bullet(doc, "The NAS shall reside on the same local network segment as the compute nodes without routing through external networks.")
    add_bullet(doc, "The storage architecture shall support future capacity expansion without requiring complete system reconstruction.",
               bold_prefix="Important: ")
    add_bullet(doc, "A RAID-based or ZFS-based storage configuration is strongly recommended for data redundancy and resilience.", level=1)

    add_spacer(doc, 4)
    add_callout(doc,
        "Critical Infrastructure Component",
        "The NAS serves as the primary storage tier for all user data. NAS unavailability or inaccessibility will result in "
        "complete loss of user data access. Reliable network connectivity between compute nodes and the NAS is therefore a critical operational requirement.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3 — Network & Connectivity
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Network & Connectivity")

    add_bullet(doc, "Static IP addresses shall be assigned to all infrastructure components (compute nodes, NAS, network equipment) — dynamic IP addressing (DHCP) is not permitted for production infrastructure.")
    add_bullet(doc, "The following network documentation shall be provided:")
    add_bullet(doc, "Hostname-to-IP address mapping for each machine.", level=1)
    add_bullet(doc, "Static IP address assignments.", level=1)
    add_bullet(doc, "MAC address registry for all network interfaces.", level=1)

    add_bullet(doc, "Internet connectivity requirements:", bold_prefix="")
    add_bullet(doc, "Minimum 1 Gbps symmetric internet connectivity (equal upload and download throughput) is required.", level=1)
    add_bullet(doc, "The platform architecture involves streaming remote GPU desktops to end users; consequently, upstream bandwidth is critical for user experience.", level=1)

    add_bullet(doc, "Internal network requirements:", bold_prefix="")
    add_bullet(doc, "Network links between compute nodes and the NAS shall be minimum 1 Gbps.", level=1)
    add_bullet(doc, "10 Gbps internal network links are strongly recommended for optimal performance, particularly for data-intensive workloads.", level=1)
    add_bullet(doc, "Implementing 10 Gbps networking during initial deployment is significantly more cost-effective than retrofitting post-deployment.", level=1)

    add_bullet(doc, "All infrastructure components shall reside on the same network segment to minimize inter-node latency and simplify network architecture.")

    add_spacer(doc, 4)
    add_callout(doc,
        "Scalability Consideration",
        "If KSRCE anticipates future platform expansion with increased user capacity, deployment of 10 Gbps internal networking "
        "is strongly recommended at initial build. Retrofitting internal networking after production deployment involves "
        "significantly greater complexity, cost, and operational disruption.",
        callout_type="consideration"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4 — Remote Access — VPN & SSH
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Remote Access — VPN & SSH")

    add_bullet(doc, "VPN access to the KSRCE network segment hosting the LaaS infrastructure.",
               bold_prefix="Mandatory: ")
    add_bullet(doc, "The LaaS engineering team utilizes Fortinet VPN (FortiClient) for client network connectivity.", level=1)
    add_bullet(doc, "The following VPN configuration data shall be provided: VPN server address, authentication credentials, and any required configuration files.", level=1)
    add_bullet(doc, "A network topology diagram illustrating the infrastructure layout shall be provided to facilitate navigation upon VPN connection establishment.", level=1)

    add_bullet(doc, "Upon VPN connection, direct SSH access to each compute node via static IP address shall be available.")
    add_bullet(doc, "The laas-admin user account (with sudo privileges) shall be created on all machines, with credentials provided to the LaaS engineering team.")
    add_bullet(doc,
               "All software installation, configuration, troubleshooting, and ongoing maintenance shall be performed via SSH. "
               "Reliable VPN and SSH connectivity is therefore essential for both initial deployment and continued operations.",
               bold_prefix="Critical: ")

    add_spacer(doc, 4)
    add_callout(doc,
        "Operational Dependency",
        "The LaaS engineering team shall perform all software installation, service configuration, and platform "
        "deployment exclusively via remote SSH access. Any instability or interruption in VPN or SSH connectivity "
        "will directly impact deployment timelines and ongoing maintenance capabilities.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 5 — Firewall & Security
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Firewall & Security")

    add_bullet(doc, "Where institutional firewall policies exist, the following shall be provided:")
    add_bullet(doc, "Documentation of current firewall rules affecting LaaS infrastructure.", level=1)
    add_bullet(doc, "A defined procedure for requesting port openings or rule modifications.", level=1)
    add_bullet(doc,
               "Required internal ports: SSH (22), HTTP/HTTPS (80, 443, 8080), "
               "NFS (2049), and application ports within the range 3000–9999.", level=1)

    add_bullet(doc, "Outbound internet access from all compute nodes is required for software package downloads, container image retrieval, and system updates.")
    add_bullet(doc, "During initial deployment, unrestricted outbound internet access is strongly preferred.", level=1)
    add_bullet(doc, "Following successful deployment, firewall policies may be collaboratively tightened per institutional security requirements.", level=1)

    add_spacer(doc, 4)
    add_callout(doc,
        "Deployment Phase Consideration",
        "During initial deployment, unrestricted outbound internet connectivity from compute nodes is "
        "strongly preferred to facilitate efficient software installation and configuration. "
        "Upon platform stabilization, the LaaS and KSRCE teams may collaboratively define and apply "
        "appropriate outbound traffic restrictions.",
        callout_type="consideration"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 6 — Power & Environment
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Power & Environment")

    add_bullet(doc, "All compute infrastructure is designed for continuous 24/7 operation under sustained high-performance computing workloads and must be provisioned accordingly.")
    add_bullet(doc, "Power consumption parameters:")
    add_bullet(doc, "Individual compute nodes may draw up to 800–900 watts at peak load.", level=1)
    add_bullet(doc, "Aggregate power requirement for the complete infrastructure (4 compute nodes + NAS + networking) is approximately 3.5–4 kW.", level=1)
    add_bullet(doc, "The power circuit shall be rated to accommodate this load — a dedicated 32A circuit is recommended.", level=1)

    add_bullet(doc, "UPS (Uninterruptible Power Supply) requirements:", bold_prefix="")
    add_bullet(doc, "A minimum 5 kVA online UPS is strongly recommended for the complete infrastructure.", level=1)
    add_bullet(doc, "This provides approximately 40–80 minutes of backup runtime, sufficient for graceful system shutdown during extended outages.", level=1)
    add_bullet(doc, "Power interruption without UPS protection may result in data corruption, hardware damage, and unplanned service disruption.", level=1)
    add_bullet(doc, "UPS installation is the single most critical environmental consideration.", level=1)

    add_bullet(doc, "Cooling requirements:", bold_prefix="")
    add_bullet(doc, "GPU compute nodes generate significant thermal output, particularly under sustained workload.", level=1)
    add_bullet(doc, "The facility shall provide adequate air conditioning or ventilation.", level=1)
    add_bullet(doc, "Recommended ambient temperature range: 18–24°C.", level=1)
    add_bullet(doc, "For dedicated server rooms, the cooling system shall be capable of dissipating 3–4 kW of continuous thermal output.", level=1)

    add_bullet(doc, "Physical security requirements:", bold_prefix="")
    add_bullet(doc, "All infrastructure shall be housed in a locked, access-controlled room.", level=1)
    add_bullet(doc, "The facility shall not be a general-purpose area where unauthorized personnel may inadvertently disrupt operations.", level=1)

    add_spacer(doc, 4)
    add_callout(doc,
        "Critical Risk Factor",
        "Power outages without UPS protection are the leading cause of data corruption "
        "and hardware damage in GPU computing environments. Sudden power loss during active "
        "workloads can corrupt user data, damage storage drives, and cause irreversible hardware "
        "failures. UPS installation is strongly recommended prior to platform deployment.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 7 — Technical Support
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Technical Support")

    add_bullet(doc, "A designated point of contact from the KSRCE infrastructure team for incident response:",
               bold_prefix="Required: ")
    add_bullet(doc, "Network infrastructure issues (VPN connectivity failures, firewall policy conflicts).", level=1)
    add_bullet(doc, "Hardware issues (unresponsive systems requiring physical intervention).", level=1)
    add_bullet(doc, "Power and environmental issues (UPS alerts, cooling system failures).", level=1)

    add_bullet(doc, "Response time requirement: within 4 hours for critical incidents (system outages, network disruptions).")
    add_bullet(doc, "Communication protocol: email for non-urgent matters; telephone or instant messaging for urgent incidents.")
    add_bullet(doc, "While frequent incidents are not anticipated, responsive support contact significantly reduces mean time to resolution when issues arise.")

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 8 — Application Server
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. Application Server")

    add_bullet(doc, "In addition to GPU compute nodes, one dedicated machine (or virtual machine) shall be provisioned to host the LaaS application stack — designated as the 'application server'.",
               bold_prefix="Mandatory: ")
    add_bullet(doc, "This component is distinct from GPU compute nodes: compute nodes execute user workloads, while the application server hosts the platform infrastructure.")

    add_bullet(doc, "Minimum specifications:", bold_prefix="")
    add_bullet(doc, "8 vCPU", level=1)
    add_bullet(doc, "16 GB RAM", level=1)
    add_bullet(doc, "256 GB SSD storage", level=1)

    add_bullet(doc, "Recommended specifications:", bold_prefix="")
    add_bullet(doc, "16 vCPU", level=1)
    add_bullet(doc, "32 GB RAM", level=1)
    add_bullet(doc, "500 GB SSD (providing additional headroom for logs, monitoring data, and database growth)", level=1)

    add_bullet(doc, "The application server hosts the following services:", bold_prefix="")
    add_bullet(doc, "Web frontend (Next.js)", level=1)
    add_bullet(doc, "Backend API (NestJS)", level=1)
    add_bullet(doc, "PostgreSQL database", level=1)
    add_bullet(doc, "Keycloak (SSO/identity management)", level=1)
    add_bullet(doc, "Nginx reverse proxy", level=1)
    add_bullet(doc, "Monitoring stack (Prometheus, Grafana, Loki, Alertmanager)", level=1)

    add_bullet(doc, "Operating system: Ubuntu 22.04.5 LTS (consistent with compute node specifications)")
    add_bullet(doc, "The application server shall be accessible via the same VPN and SSH configuration as compute nodes.")
    add_bullet(doc, "A static IP address shall be assigned.")

    add_bullet(doc, "If KSRCE maintains existing virtualization infrastructure (VMware, Proxmox, etc.), a virtual machine is an acceptable deployment option.")
    add_bullet(doc, "Cloud hosting (AWS, Azure, etc.) is also permissible, provided the instance can establish network connectivity to the internal network segment where compute nodes reside.")

    add_spacer(doc, 4)
    add_callout(doc,
        "Critical Infrastructure Component",
        "The application server serves as the primary interface for end users accessing the LaaS platform. "
        "Application server unavailability will result in complete platform inaccessibility, even if compute nodes "
        "remain operational. This component represents a single point of failure for the platform UI.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 9 — Database Backup & Snapshot Plans
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. Database Backup & Snapshot Plans")

    add_bullet(doc, "The LaaS platform maintains critical operational data in PostgreSQL, including: user accounts, billing records, session history, compute resource allocation, and system configuration.")

    add_bullet(doc, "Automated daily database backups (pg_dump) shall be scheduled, recommended execution time: 02:00 local time.",
               bold_prefix="Configuration: ")

    add_bullet(doc, "Backup storage shall be located on a separate device or storage tier from the database server — ideally on the NAS or a dedicated backup storage system.")

    add_bullet(doc, "A minimum of 30 days of rolling daily backups shall be retained.",
               bold_prefix="Retention Policy: ")

    add_bullet(doc, "A weekly full backup supplemented by daily incremental backups is the recommended backup strategy.")

    add_bullet(doc, "Periodic backup replication to an offsite location (cloud storage bucket, alternate physical site) is recommended for disaster recovery purposes.")

    add_bullet(doc, "For NAS user data storage: ZFS snapshots shall be configured on a daily schedule — ZFS provides near-zero-cost snapshot capabilities.")
    add_bullet(doc, "NAS snapshots shall be retained for a minimum of 14 days to enable user file recovery from accidental deletion.")

    add_bullet(doc, "A documented restore procedure shall be maintained and validated to ensure recovery steps are clear and tested.")

    add_bullet(doc, "KSRCE shall provide or allocate a dedicated backup storage device or partition (minimum 500 GB) for database backups and NAS snapshots.")

    add_spacer(doc, 4)
    add_callout(doc,
        "Critical Operational Requirement",
        "In the absence of a backup strategy, a single database corruption event, accidental deletion, or hardware failure "
        "could result in permanent loss of all user data, billing history, and platform configuration. "
        "Implementation of a comprehensive backup strategy is mandatory.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 10 — Contingency & Disaster Recovery Plans
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Contingency & Disaster Recovery Plans")
    add_body(doc, "The following section defines response procedures for various failure scenarios:")
    add_spacer(doc, 4)

    # -- Network/VPN Failure --
    add_heading(doc, "Network / VPN Failure", level=2)
    add_bullet(doc, "In the event of primary VPN unavailability, an alternative remote access method shall be available.")
    add_bullet(doc, "Acceptable alternatives include: secondary VPN profile, direct SSH via router port forwarding, or Tailscale/WireGuard as backup VPN solution.")
    add_bullet(doc, "KSRCE shall designate on-site personnel capable of physically verifying system status and restarting network equipment if remote access is completely unavailable.")
    add_bullet(doc, "A 4G/LTE USB modem or mobile hotspot is recommended as emergency internet connectivity if the primary ISP experiences outage.")

    # -- Power Failure --
    add_heading(doc, "Power Failure", level=2)
    add_bullet(doc, "UPS systems provide 40–80 minutes of operational buffer — the following procedures shall be defined:")
    add_bullet(doc, "If power is restored within the buffer window: systems continue normal operation without intervention.", level=1)
    add_bullet(doc, "If power is not restored: a documented graceful shutdown procedure shall be implemented, ideally automated via UPS monitoring software (NUT/apcupsd).", level=1)
    add_bullet(doc, "Following power restoration: systems shall auto-boot (BIOS AC Power Recovery) and services shall auto-start (systemd); however, verification of successful recovery is required.", level=1)
    add_bullet(doc, "A documented boot sequence shall be maintained: NAS initialization → compute node boot → NFS mount verification → platform service startup.", level=1)

    # -- Hardware Failure --
    add_heading(doc, "Hardware Failure", level=2)
    add_bullet(doc, "Compute node failure: the platform continues operation on remaining nodes with reduced capacity (degraded mode, not complete outage).")
    add_bullet(doc, "NAS failure: this is a critical event — all user data resides on the NAS. RAID/ZFS mirroring provides protection against single-drive failure, but not against controller failure.")
    add_bullet(doc, "Application server failure: complete platform UI unavailability — this component is a single point of failure.")
    add_bullet(doc, "Recommendation: documented response procedures for 'failure scenario X' shall be maintained for each critical component.")
    add_bullet(doc, "An inventory of spare components (drives, cables, power supplies) shall be maintained to minimize mean time to recovery (MTTR).")

    # -- Data Centre Standards --
    add_heading(doc, "Data Centre Standards (ODC Best Practices)", level=2)
    add_bullet(doc, "The server room shall comply with basic ODC/data centre standards, regardless of facility scale:")
    add_bullet(doc, "Raised flooring or cable management trays to prevent cable congestion.", level=1)
    add_bullet(doc, "Proper rack mounting for all equipment (not desktop or shelf placement).", level=1)
    add_bullet(doc, "Fire suppression capability (minimum: CO₂ fire extinguisher in the room).", level=1)
    add_bullet(doc, "Access control logging (personnel entry/exit records).", level=1)
    add_bullet(doc, "Temperature and humidity monitoring with alert notifications.", level=1)
    add_bullet(doc, "Clear labeling of all cables, ports, and equipment.", level=1)
    add_bullet(doc, "Dust-controlled environment with regular cleaning schedule for room and equipment air filters.", level=1)

    add_spacer(doc, 4)
    add_callout(doc,
        "Operational Best Practice",
        "A well-organized server room directly impacts uptime, troubleshooting efficiency, "
        "and hardware longevity. Preventable incidents such as accidental cable disconnection, "
        "thermal events due to blocked airflow, or component failure from dust accumulation "
        "can be mitigated through adherence to basic data centre operational standards.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 11 — Handover Checklist
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. Handover Checklist")
    add_body(doc, "Prior to handover of remote access credentials, the following items shall be verified as complete:")
    add_spacer(doc, 4)

    checklist_items = [
        "Ubuntu 22.04.5 LTS installed on all compute nodes",
        "SSH server installed, enabled, and configured for automatic start on boot",
        "laas-admin user account with sudo privileges created on all systems",
        "Static IP addresses assigned to all machines and NAS",
        "NAS configured and accessible from all compute nodes via NFS",
        "VPN access credentials provided to the LaaS engineering team",
        "Firewall rules documented and required ports opened",
        "UPS installed and connected to critical infrastructure",
        "Cooling/ventilation verified for the server room",
        "Hostname-to-IP mapping document provided",
        "Internet connectivity verified on all systems (ping and package update tests successful)",
        "KSRCE technical support contact details provided",
        "Application server (or VM) provisioned with minimum specifications",
        "Backup storage device/partition allocated (≥500 GB)",
        "UPS configured with automated shutdown capability (NUT/apcupsd)",
        "Emergency contact procedure documented for after-hours incidents",
        "Boot sequence documented (NAS → compute nodes → services)",
        "Server room compliant with basic ODC standards (rack, cooling, fire safety, access control)",
        "Backup/restore procedure documented and validated through testing",
    ]
    for item in checklist_items:
        add_checklist_item(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 12 — Additional Considerations
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. Additional Considerations")
    add_body(doc,
             "(The following items are not mandatory requirements but are recommended for enhanced operational resilience)",
             italic=True, color=GREY_TEXT)
    add_spacer(doc, 4)

    optional_items = [
        (
            "Redundant internet connectivity — ",
            "reliance on a single ISP creates a single point of failure for remote platform access. "
            "A secondary ISP connection (even at reduced bandwidth) maintains platform availability during primary ISP outages."
        ),
        (
            "Power monitoring and alerting — ",
            "UPS activation events shall trigger automated notifications to designated infrastructure personnel "
            "enabling proactive response prior to battery depletion."
        ),
        (
            "Physical KVM or IPMI access — ",
            "in scenarios where a system is completely unresponsive to SSH, remote reboot capability via IPMI/BMC "
            "or rapid on-site physical access significantly reduces downtime."
        ),
        (
            "Scheduled maintenance windows — ",
            "a regular maintenance schedule (e.g., first Sunday of each month, 02:00–06:00) shall be established "
            "for updates, system reboots, and hardware inspections. Users shall receive advance notification."
        ),
        (
            "Spare components inventory — ",
            "maintaining an inventory of critical spare components (power supplies, RAM modules, network cables) "
            "is essential to minimize mean time to recovery (MTTR) and avoid extended downtime awaiting replacement parts."
        ),
        (
            "Network documentation — ",
            "a network topology diagram illustrating the interconnection of compute nodes, NAS, network switch, "
            "and internet gateway significantly accelerates troubleshooting efforts."
        ),
        (
            "Backup power for network infrastructure — ",
            "the UPS shall protect network switches and routers in addition to compute infrastructure. "
            "Switch power loss results in complete network connectivity loss regardless of compute node status."
        ),
        (
            "Environmental monitoring — ",
            "temperature sensors with alert capabilities (including consumer-grade smart sensors) "
            "provide early warning of thermal conditions before equipment damage occurs."
        ),
        (
            "Physical access for maintenance — ",
            "equipment shall be physically accessible (not obstructed by furniture or other items) "
            "to facilitate drive replacement, cable management, and component servicing."
        ),
        (
            "Database backup validation — ",
            "monthly test restoration of database backups shall be scheduled to verify backup integrity. "
            "Untested backups cannot be relied upon for disaster recovery."
        ),
        (
            "Offsite backup replication — ",
            "storing all backups at a single physical location creates vulnerability to site-wide disasters "
            "(fire, flood, theft). Cloud storage replication (S3, Azure Blob) for weekly backups provides significant data protection."
        ),
    ]

    for bold_part, rest_part in optional_items:
        add_bullet(doc, rest_part, bold_prefix=bold_part)

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 13 — Contact Information
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13. Contact Information")

    add_bullet(doc, "[To be provided]", bold_prefix="LaaS Platform Team:  ")
    add_bullet(doc, "[To be provided]", bold_prefix="KSRCE Infrastructure Team:  ")

    add_spacer(doc, 16)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("— End of Document —")
    r.font.name  = "Calibri"
    r.font.size  = Pt(10)
    r.font.color.rgb = LIGHT_GREY
    r.italic = True

    # ── Footer ──
    add_page_numbers(doc)

    doc.save(OUTPUT_PATH)
    print(f"Document generated successfully:\n  {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
