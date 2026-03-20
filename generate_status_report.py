"""
LaaS Project Status Report Generator
Clean, modern, high-contrast design with teal accents
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Color scheme
TEAL = RGBColor(0x00, 0x97, 0xA7)
CHARCOAL = RGBColor(0x2D, 0x2D, 0x2D)
BODY_TEXT = RGBColor(0x33, 0x33, 0x33)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xEF, 0x6C, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
BORDER_GRAY = RGBColor(0xDD, 0xDD, 0xDD)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, color="DDDDDD", width="4"):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def style_document(doc):
    """Set base document styles."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.space_after = Pt(6)
    
    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)


def add_title_page(doc):
    """Create clean title page with teal accent."""
    # Add some spacing at top
    for _ in range(4):
        doc.add_paragraph()
    
    # Teal accent bar (using a table trick)
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent_cell = accent_table.rows[0].cells[0]
    accent_cell.width = Inches(6)
    set_cell_shading(accent_cell, "0097A7")
    p = accent_cell.paragraphs[0]
    p.add_run(" ")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph()
    
    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LaaS — Lab-as-a-Service")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = CHARCOAL
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Project Status Report")
    run.font.size = Pt(24)
    run.font.color.rgb = TEAL
    
    doc.add_paragraph()
    
    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("20 March 2026")
    run.font.size = Pt(16)
    run.font.color.rgb = BODY_TEXT
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Organization
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KSRCE × GKT AI Supercomputing Lab")
    run.font.size = Pt(14)
    run.font.color.rgb = CHARCOAL
    run.font.italic = True
    
    # Page break
    doc.add_page_break()


def add_section_header(doc, text):
    """Add section header with teal left accent bar."""
    # Create a 2-column table for accent effect
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Teal accent bar (narrow)
    accent_cell = table.rows[0].cells[0]
    accent_cell.width = Inches(0.1)
    set_cell_shading(accent_cell, "0097A7")
    accent_cell.paragraphs[0].add_run(" ")
    
    # Header text
    text_cell = table.rows[0].cells[1]
    text_cell.width = Inches(6.4)
    p = text_cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = CHARCOAL
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph()


def add_styled_table(doc, headers, rows, col_widths=None, status_col=None):
    """Create a styled table with teal headers and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, "0097A7")
        set_cell_borders(cell, "0097A7", "4")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
        
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_shading(cell, bg_color)
            set_cell_borders(cell, "DDDDDD", "4")
            p = cell.paragraphs[0]
            
            # Handle status coloring
            if status_col is not None and c_idx == status_col:
                if "✅" in val or "Done" in val or "Complete" in val:
                    run = p.add_run(val)
                    run.font.color.rgb = GREEN
                    run.font.bold = True
                elif "❌" in val or "Pending" in val or "Not Started" in val:
                    run = p.add_run(val)
                    run.font.color.rgb = RED
                    run.font.bold = True
                elif "🔄" in val or "Partial" in val or "In Progress" in val:
                    run = p.add_run(val)
                    run.font.color.rgb = AMBER
                    run.font.bold = True
                else:
                    run = p.add_run(val)
                    run.font.color.rgb = BODY_TEXT
            else:
                run = p.add_run(val)
                run.font.color.rgb = BODY_TEXT
            
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    doc.add_paragraph()
    return table


def add_bullet_list(doc, items):
    """Add a bullet list."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT


def build_report():
    doc = Document()
    style_document(doc)
    
    # TITLE PAGE
    add_title_page(doc)
    
    # SECTION 1: PROJECT AT A GLANCE
    add_section_header(doc, "1. PROJECT AT A GLANCE")
    
    glance_table = doc.add_table(rows=6, cols=2)
    glance_data = [
        ("What", "On-premises GPU compute platform — remote access to GPU-accelerated desktops via browser"),
        ("Who", "University students, faculty, researchers, public users"),
        ("Where", "K.S.R. College of Engineering (KSRCE) × Global Knowledge Technologies (GKT)"),
        ("Hardware", "4 nodes — AMD Ryzen 9 9950X3D, RTX 5090 32GB, 64GB DDR5, 2TB NVMe each"),
        ("Innovation", "First platform to offer fractional GPU desktop streaming with per-user isolation"),
        ("Status", "Infrastructure ✅ | Database ✅ | Auth & Home UI ✅ | Core Modules 🔄 In Progress"),
    ]
    
    for r_idx, (label, value) in enumerate(glance_data):
        bg_color = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
        label_cell = glance_table.rows[r_idx].cells[0]
        value_cell = glance_table.rows[r_idx].cells[1]
        
        set_cell_shading(label_cell, "0097A7")
        set_cell_borders(label_cell, "0097A7", "4")
        p = label_cell.paragraphs[0]
        run = p.add_run(label)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        label_cell.width = Inches(1.2)
        
        set_cell_shading(value_cell, bg_color)
        set_cell_borders(value_cell, "DDDDDD", "4")
        p = value_cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT
        value_cell.width = Inches(5.3)
    
    doc.add_paragraph()
    
    # SECTION 2: PROJECT OBJECTIVES
    add_section_header(doc, "2. PROJECT OBJECTIVES")
    
    objectives = [
        "Democratize HPC access for academic institutions",
        "Monetize 4 on-premises GPU machines through multi-stream revenue",
        "Deliver seamless remote desktop experience with fractional GPU allocation",
        "Support compute booking, billing, mentorship, and academic integration",
        "Achieve >99.9% session uptime with sub-500ms API response times",
    ]
    add_bullet_list(doc, objectives)
    
    # SECTION 3: PLANNED MODULES
    add_section_header(doc, "3. PLANNED MODULES (Complete Feature Scope)")
    
    modules = [
        ("Authentication & SSO", "Email/OTP signup, university Keycloak SSO, OAuth (Google/GitHub)"),
        ("Home Dashboard", "User overview — storage, sessions, quick actions"),
        ("Billing Dashboard", "Credit balance, spend tracking, burn rate, usage charts"),
        ("Compute Sessions", "Launch/manage GPU-accelerated desktop sessions, tier selection"),
        ("Session Booking", "Time-slot based reservation system, availability checks"),
        ("Wallet & Payments", "Razorpay/Stripe integration, credit packages, auto-deduction"),
        ("Storage Management", "15GB per-user ZFS storage, quota management, OS switching"),
        ("Mentor Booking", "Mentor profiles, calendar scheduling, booking & reviews"),
        ("Academic/LMS", "Courses, labs, assignments, submissions, grading"),
        ("Organization Management", "Departments, user groups, RBAC, org-level quotas"),
        ("Community Hub", "Discussions, project showcase, achievements/gamification"),
        ("Admin Panel", "User management, node monitoring, billing admin, audit logs"),
        ("Support System", "Ticket system, user feedback"),
        ("Notifications", "Multi-channel (email, SMS, web push) notification system"),
        ("Landing Page", "Public-facing product page"),
    ]
    
    add_styled_table(doc, ["Module", "Description"], modules, col_widths=[2.0, 4.5])
    
    # SECTION 4: WHAT IS DONE
    add_section_header(doc, "4. CURRENT STATUS — WHAT IS DONE ✅")
    
    # 4A. Infrastructure
    p = doc.add_paragraph()
    run = p.add_run("A. Infrastructure (All Complete)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = CHARCOAL
    
    infra_items = [
        ("POC Server (Ryzen 9 7950X3D + RTX 4090)", "✅ Done"),
        ("GPU Driver + CUDA 12.8", "✅ Done"),
        ("Docker + NVIDIA Container Toolkit", "✅ Done"),
        ("HAMi-core GPU VRAM Enforcement", "✅ Done"),
        ("CUDA MPS (SM Partitioning)", "✅ Done"),
        ("ZFS Storage Pool + Per-User Datasets", "✅ Done"),
        ("NFS Configuration", "✅ Done"),
        ("Storage Provisioning Service", "✅ Done"),
        ("Keycloak SSO (2 realms configured)", "✅ Done"),
        ("Selkies Desktop + WebRTC Streaming", "✅ Done"),
        ("NVENC Hardware Encoding", "✅ Done"),
        ("lxcfs Resource Spoofing", "✅ Done"),
        ("Monitoring Stack (Prometheus/Grafana/Loki)", "✅ Done"),
        ("Multi-Container GPU Sharing (4 concurrent)", "✅ Done"),
    ]
    
    add_styled_table(doc, ["Component", "Status"], infra_items, col_widths=[4.5, 2.0], status_col=1)
    
    # 4B. Software Platform
    p = doc.add_paragraph()
    run = p.add_run("B. Software Platform")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = CHARCOAL
    
    software_items = [
        ("Database Schema (75+ tables, 16 domains)", "✅ Done", "PostgreSQL + Prisma, all migrations applied"),
        ("Email/OTP Signup", "✅ Done", "Full flow working"),
        ("Email/Password Login", "✅ Done", "JWT access + refresh tokens"),
        ("University SSO (Keycloak)", "✅ Done", "OIDC federation, auto-provisioning"),
        ("Home Dashboard", "✅ Done", "Storage overview, session stats, quick actions"),
        ("Billing Dashboard", "✅ Done", "Balance, burn rate, spend chart, resources"),
        ("App Shell & Navigation", "✅ Done", "Sidebar, header, theme toggle"),
        ("Storage Auto-Provisioning", "✅ Done", "ZFS dataset created on first SSO login"),
        ("Live Storage Usage Display", "✅ Done", "Real-time from Flask service"),
        ("Dark/Light Theme", "✅ Done", "CSS variables-based theming"),
    ]
    
    add_styled_table(doc, ["Module", "Status", "Notes"], software_items, col_widths=[2.8, 1.0, 2.7], status_col=1)
    
    # Page break before pending
    doc.add_page_break()
    
    # SECTION 5: WHAT IS PENDING
    add_section_header(doc, "5. WHAT IS PENDING ❌")
    
    pending_items = [
        ("Compute Session Management", "HIGH", "Backend API + Frontend UI for launching/managing GPU sessions"),
        ("Session Booking / Reservation", "HIGH", "Time-slot booking system with availability checks"),
        ("Wallet & Payment Integration", "HIGH", "Razorpay/Stripe integration, credit purchase flow"),
        ("Mentor Booking", "MEDIUM", "Profiles, calendar, booking flow, reviews"),
        ("Landing Page", "MEDIUM", "Public-facing product marketing page"),
        ("Organization Management", "MEDIUM", "Department/group admin, RBAC controls"),
        ("Academic/LMS Features", "MEDIUM", "Course/lab management, submissions, grading"),
        ("Admin Panels", "MEDIUM", "User, billing, and system administration"),
        ("Support Ticket System", "LOW", "Ticket creation, tracking, messaging"),
        ("Community Features", "LOW", "Discussions, project showcase, achievements"),
        ("Notification System", "LOW", "Multi-channel notification delivery"),
        ("Invoice Generation", "LOW", "Automated billing invoices"),
        ("Password Reset UI", "LOW", "Frontend form (backend endpoint exists)"),
        ("2FA Implementation", "LOW", "Two-factor auth (schema supports it)"),
    ]
    
    # Custom table with priority coloring
    table = doc.add_table(rows=1 + len(pending_items), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Module", "Priority", "What's Needed"]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, "0097A7")
        set_cell_borders(cell, "0097A7", "4")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
    
    for r_idx, (module, priority, needed) in enumerate(pending_items):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
        
        # Module cell
        cell = row.cells[0]
        set_cell_shading(cell, bg_color)
        set_cell_borders(cell, "DDDDDD", "4")
        run = cell.paragraphs[0].add_run(module)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT
        cell.width = Inches(2.5)
        
        # Priority cell with color
        cell = row.cells[1]
        set_cell_shading(cell, bg_color)
        set_cell_borders(cell, "DDDDDD", "4")
        run = cell.paragraphs[0].add_run(priority)
        run.font.size = Pt(10)
        run.font.bold = True
        if priority == "HIGH":
            run.font.color.rgb = RED
        elif priority == "MEDIUM":
            run.font.color.rgb = AMBER
        else:
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        cell.width = Inches(1.0)
        
        # Needed cell
        cell = row.cells[2]
        set_cell_shading(cell, bg_color)
        set_cell_borders(cell, "DDDDDD", "4")
        run = cell.paragraphs[0].add_run(needed)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT
        cell.width = Inches(3.0)
    
    doc.add_paragraph()
    
    # SECTION 6: TECHNOLOGY STACK
    add_section_header(doc, "6. TECHNOLOGY STACK")
    
    tech_stack = [
        ("Frontend", "Next.js 15, TypeScript, Tailwind CSS 4, Radix UI"),
        ("Backend", "NestJS 11, Fastify, TypeScript"),
        ("Database", "PostgreSQL, Prisma ORM"),
        ("Auth", "Keycloak (OIDC/SAML), JWT, Passport.js"),
        ("Payments", "Razorpay, Stripe"),
        ("GPU Sharing", "HAMi-core, CUDA MPS"),
        ("Remote Desktop", "Selkies-GStreamer (WebRTC), NVENC"),
        ("Containers", "Docker CE, NVIDIA Container Toolkit"),
        ("Storage", "ZFS, NFS v4"),
        ("Monitoring", "Prometheus, Grafana, Loki, DCGM Exporter"),
    ]
    
    add_styled_table(doc, ["Layer", "Technology"], tech_stack, col_widths=[1.8, 4.7])
    
    # SECTION 7: COMPUTE TIERS
    add_section_header(doc, "7. COMPUTE TIERS (Planned Pricing)")
    
    tiers = [
        ("Starter", "2", "4 GB", "—", "₹15"),
        ("Standard", "4", "8 GB", "—", "₹30"),
        ("Pro", "4", "8 GB", "4 GB", "₹60"),
        ("Power", "8", "16 GB", "8 GB", "₹100"),
        ("Max", "8", "16 GB", "16 GB", "₹150"),
        ("Full Machine", "16", "48 GB", "32 GB", "₹300"),
    ]
    
    add_styled_table(doc, ["Tier", "vCPU", "RAM", "VRAM", "Rate/hr"], tiers, col_widths=[1.5, 0.8, 1.0, 1.0, 1.0])
    
    # Save document
    output_path = r"c:\Users\Punith\LaaS\LaaS_Project_Status_Update_20-03-2026.docx"
    temp_path = r"c:\Users\Punith\LaaS\LaaS_Project_Status_Update_20-03-2026_NEW.docx"
    
    try:
        doc.save(output_path)
        print(f"✅ Document saved: {output_path}")
    except PermissionError:
        doc.save(temp_path)
        print(f"⚠️ Original file locked. Saved to: {temp_path}")
        print("Close Word and rename the file to replace the original.")


if __name__ == "__main__":
    build_report()
