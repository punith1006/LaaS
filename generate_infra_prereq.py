"""
Generate LaaS Infrastructure Prerequisites & Requirements document (.docx)
for KSRCE Infrastructure Team — conversational bullet-point style.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"c:\Users\Punith\Downloads\LaaS_Infrastructure_Prerequisites_KSRCE.docx"

# ── Colour palette ──────────────────────────────────────────────────────────
HEADING_BLUE  = RGBColor(0x2E, 0x50, 0x90)
DARK_TEXT     = RGBColor(0x1A, 0x1A, 0x2E)
MID_TEXT      = RGBColor(0x3A, 0x3A, 0x3A)
GREY_TEXT     = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREY    = RGBColor(0x88, 0x88, 0x88)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)


# ── Low-level XML helpers ────────────────────────────────────────────────────

def set_para_shading(paragraph, hex_color: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_para_border(paragraph, color: str = "2E5090", sz: str = "6",
                    space: str = "4", val: str = "single", side: str = "left"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), val)
    bdr.set(qn("w:sz"), sz)
    bdr.set(qn("w:space"), space)
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)
    pPr.append(pBdr)


def add_page_numbers(doc: Document):
    """Add 'Page X of Y' footer to every section."""
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("Page ")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = LIGHT_GREY

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText1 = OxmlElement("w:instrText")
        instrText1.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "begin")
        instrText2 = OxmlElement("w:instrText")
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")

        r = para.add_run()
        r.font.name = "Calibri"
        r.font.size = Pt(9)
        r.font.color.rgb = LIGHT_GREY
        r._r.append(fldChar1)
        r._r.append(instrText1)
        r._r.append(fldChar2)

        r2 = para.add_run(" of ")
        r2.font.name = "Calibri"
        r2.font.size = Pt(9)
        r2.font.color.rgb = LIGHT_GREY

        r3 = para.add_run()
        r3.font.name = "Calibri"
        r3.font.size = Pt(9)
        r3.font.color.rgb = LIGHT_GREY
        r3._r.append(fldChar3)
        r3._r.append(instrText2)
        r3._r.append(fldChar4)

        r4 = para.add_run("  |  LaaS Platform — Infrastructure Prerequisites & Requirements")
        r4.font.name = "Calibri"
        r4.font.size = Pt(9)
        r4.font.color.rgb = LIGHT_GREY


# ── Style helpers ────────────────────────────────────────────────────────────

def set_run_font(run, bold=False, italic=False, size=11, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    sizes = {1: 15, 2: 12, 3: 11}
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = HEADING_BLUE
    if level == 1:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bdr = OxmlElement("w:bottom")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "6")
        bdr.set(qn("w:space"), "1")
        bdr.set(qn("w:color"), "2E5090")
        pBdr.append(bdr)
        pPr.append(pBdr)


def add_body(doc: Document, text: str, bold: bool = False, italic: bool = False,
             indent: bool = False, color=None) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(3)
    if indent:
        para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(text)
    set_run_font(run, bold=bold, italic=italic, color=color)


def add_bullet(doc: Document, text: str, level: int = 0,
               bold_prefix: str = None, italic_suffix: bool = False) -> None:
    """
    Renders a bullet as a plain paragraph with a dash prefix
    to avoid List Bullet style indent unpredictability.
    level 0 → single bullet  (•  prefix, slight indent)
    level 1 → sub-bullet     (–  prefix, more indent)
    """
    bullet_char = "•" if level == 0 else "–"
    indent_val  = Inches(0.3 + level * 0.25)

    para = doc.add_paragraph()
    para.paragraph_format.space_before    = Pt(2)
    para.paragraph_format.space_after     = Pt(2)
    para.paragraph_format.left_indent     = indent_val
    para.paragraph_format.first_line_indent = Inches(-0.22)

    # bullet glyph
    r0 = para.add_run(f"{bullet_char}  ")
    set_run_font(r0, color=HEADING_BLUE if level == 0 else MID_TEXT)

    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = para.add_run(text)
        set_run_font(r2, italic=italic_suffix)
    else:
        r = para.add_run(text)
        set_run_font(r, italic=italic_suffix)


def add_checklist_item(doc: Document, text: str) -> None:
    """Renders a checklist bullet with a dash (–) prefix."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before    = Pt(2)
    para.paragraph_format.space_after     = Pt(2)
    para.paragraph_format.left_indent     = Inches(0.3)
    para.paragraph_format.first_line_indent = Inches(-0.22)

    r0 = para.add_run("–  ")
    set_run_font(r0, color=HEADING_BLUE)
    r = para.add_run(text)
    set_run_font(r)


def add_callout(doc: Document, label: str, text: str,
                callout_type: str = "consideration") -> None:
    """Shaded callout box with left border stripe."""
    bg_color     = "E8EFF8" if callout_type == "consideration" else "FFF0E8"
    border_color = "2E5090" if callout_type == "consideration" else "C0392B"
    label_color  = HEADING_BLUE if callout_type == "consideration" \
                   else RGBColor(0xC0, 0x39, 0x2B)

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(10)
    para.paragraph_format.left_indent  = Inches(0.2)
    para.paragraph_format.right_indent = Inches(0.2)
    set_para_shading(para, bg_color)
    set_para_border(para, color=border_color, sz="12", space="8")

    r1 = para.add_run(f"  {label}:  ")
    set_run_font(r1, bold=True, size=10, color=label_color)
    r2 = para.add_run(text)
    set_run_font(r2, size=10)


def add_spacer(doc: Document, size: int = 4) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run("")
    run.font.size = Pt(size)


# ── Main document builder ────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── Default paragraph style ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════════════════════
    add_spacer(doc, 28)

    # Decorative top bar
    bar = doc.add_paragraph()
    set_para_shading(bar, "2E5090")
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after  = Pt(0)
    r = bar.add_run("   ")
    r.font.size = Pt(18)

    add_spacer(doc, 18)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after  = Pt(4)
    title_run = title_para.add_run("LaaS Platform")
    title_run.font.name  = "Calibri"
    title_run.font.size  = Pt(28)
    title_run.bold       = True
    title_run.font.color.rgb = HEADING_BLUE

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub1.paragraph_format.space_before = Pt(0)
    sub1.paragraph_format.space_after  = Pt(4)
    r = sub1.add_run("Infrastructure Prerequisites & Requirements")
    r.font.name  = "Calibri"
    r.font.size  = Pt(20)
    r.bold       = True
    r.font.color.rgb = MID_TEXT

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub2.paragraph_format.space_before = Pt(4)
    sub2.paragraph_format.space_after  = Pt(2)
    r = sub2.add_run("For KSRCE Infrastructure Team")
    r.font.name  = "Calibri"
    r.font.size  = Pt(14)
    r.font.color.rgb = GREY_TEXT

    add_spacer(doc, 10)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_para.paragraph_format.space_before = Pt(0)
    meta_para.paragraph_format.space_after  = Pt(0)
    r = meta_para.add_run("Version 1.0  |  April 2026  |  Prepared by: LaaS Platform Team")
    r.font.name   = "Calibri"
    r.font.size   = Pt(10)
    r.font.color.rgb = GREY_TEXT

    add_spacer(doc, 26)

    bar2 = doc.add_paragraph()
    set_para_shading(bar2, "2E5090")
    bar2.paragraph_format.space_before = Pt(0)
    bar2.paragraph_format.space_after  = Pt(0)
    r = bar2.add_run("   ")
    r.font.size = Pt(8)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # SECTION — What This Document Is About
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "What This Document Is About")
    add_bullet(doc, "This document lists out everything we need from the KSRCE infrastructure team before we can begin setting up the LaaS platform software.")
    add_bullet(doc, "Once all these prerequisites are in place and we have remote access, we'll handle everything else — all the software installation, configuration, and deployment will be done by us remotely.")
    add_bullet(doc, "Think of this as a checklist of \"what we need ready before we start\".")

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 1 — OS & Machine Setup
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Operating System & Machine Setup")

    add_bullet(doc, "Ubuntu 22.04.5 LTS installed on all the compute machines — please ensure it's exactly this version.",
               bold_prefix="We need ")
    add_bullet(doc, "The installation can be either:")
    add_bullet(doc, "Standalone Linux — if these machines are being dedicated solely for LaaS.", level=1)
    add_bullet(doc, "Dual-boot (Windows + Linux) — if KSRCE has other uses for these machines as well.", level=1)

    add_bullet(doc, "If it's a dual-boot setup:")
    add_bullet(doc, "Please set the BIOS to boot into Linux by default.", level=1)
    add_bullet(doc, "Make sure at least 80% of the disk space is allocated to the Linux partition — we'll need the space.", level=1)

    add_bullet(doc, "All machines need to be configured to stay on and running 24/7 — they shouldn't go to sleep or shut down on their own.")
    add_bullet(doc, "In the BIOS, enable \"AC Power Recovery → Power On\" so that if there's a power cut and power comes back, the machines automatically turn on and boot into the OS without anyone having to physically press the power button.")

    add_bullet(doc, "SSH must be set up on every machine:",
               bold_prefix="")
    add_bullet(doc, "Install and enable OpenSSH server.", level=1)
    add_bullet(doc, "Make sure the SSH service starts automatically on boot.", level=1)
    add_bullet(doc, "Create a dedicated user account (something like laas-admin) with sudo (admin) privileges — this is what we'll use to log in remotely.", level=1)

    add_bullet(doc,
               "Give each machine a clear hostname so we can easily identify them — like "
               "laas-node-01, laas-node-02, laas-node-03, laas-node-04, and laas-nas for the storage machine.")

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 2 — Network Storage (NAS)
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Network Storage (NAS)")

    add_bullet(doc, "At least 4TB of network-attached storage (NAS) that all the compute machines can access over the local network.",
               bold_prefix="We need ")
    add_bullet(doc, "The NAS should support NFS (Network File System) — this is how our machines will read and write data to it.")
    add_bullet(doc, "All compute machines should be able to access the NAS without any permission issues — please test this after setup.")
    add_bullet(doc, "The NAS should be on the same local network as the compute machines — no routing through external networks.")
    add_bullet(doc, "Please set up the storage in a way that allows for future expansion — if we need more space later, we should be able to add more drives without rebuilding everything from scratch.",
               bold_prefix="Important: ")
    add_bullet(doc, "A RAID-based or ZFS-based setup is ideal for this.", level=1)

    add_spacer(doc, 4)
    add_callout(doc,
        "Why this matters",
        "The NAS stores all user data. If the NAS goes down or becomes inaccessible, users lose "
        "access to their files and work. So reliable connectivity between the NAS and compute "
        "machines is critical.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 3 — Network & Connectivity
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Network & Connectivity")

    add_bullet(doc, "Every machine (compute nodes + NAS) needs a static IP address — IPs that don't change.")
    add_bullet(doc, "Please provide us a simple document/sheet listing:")
    add_bullet(doc, "Each machine's hostname.", level=1)
    add_bullet(doc, "Its static IP address.", level=1)
    add_bullet(doc, "Its MAC address.", level=1)

    add_bullet(doc, "Internet bandwidth:", bold_prefix="")
    add_bullet(doc, "Minimum 1 Gbps symmetric (same upload and download speed) internet connection.", level=1)
    add_bullet(doc, "Our platform streams remote GPU desktops to users — so upload speed matters a lot.", level=1)

    add_bullet(doc, "Internal network (between machines):", bold_prefix="")
    add_bullet(doc, "The connection between compute machines and the NAS should be at least 1 Gbps.", level=1)
    add_bullet(doc, "We strongly recommend 10 Gbps internal links for best performance — especially for data-heavy workloads.", level=1)
    add_bullet(doc, "It's much easier and cheaper to set up 10 Gbps from the start than to upgrade later.", level=1)

    add_bullet(doc, "All machines should be on the same network segment so they can talk to each other freely with minimal latency.")

    add_spacer(doc, 4)
    add_callout(doc,
        "Consideration",
        "If KSRCE plans to grow the platform with more users later, a 10 Gbps internal network "
        "is strongly recommended from the outset. Upgrading internal networking after machines "
        "are in production is significantly more complex and disruptive than getting it right from the start.",
        callout_type="consideration"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 4 — Remote Access — VPN & SSH
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Remote Access — VPN & SSH")

    add_bullet(doc, "VPN access to the KSRCE network where these machines are hosted.",
               bold_prefix="We need ")
    add_bullet(doc, "We use Fortinet VPN (FortiClient) to connect to client networks.", level=1)
    add_bullet(doc, "Please provide us with: VPN server address, login credentials, and any configuration files needed.", level=1)
    add_bullet(doc, "Also share a simple network diagram showing how the machines are laid out on the network — so we know what to expect once we're connected.", level=1)

    add_bullet(doc, "Once we're on the VPN, we should be able to SSH directly into each machine using their static IPs.")
    add_bullet(doc, "Please create the laas-admin user account (with sudo privileges) on all machines and share the credentials with us.")
    add_bullet(doc,
               "This is how we'll do everything — all software setup, configuration, troubleshooting, and maintenance "
               "will happen through SSH. So reliable VPN + SSH access is essential, not just during setup but ongoing.",
               bold_prefix="Important: ")

    add_spacer(doc, 4)
    add_callout(doc,
        "Critical Information",
        "The LaaS team will perform all software installation, service configuration, and platform "
        "deployment entirely via remote SSH access. Any instability in VPN or SSH connectivity will "
        "directly delay the deployment timeline and impact ongoing maintenance.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 5 — Firewall & Security
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Firewall & Security")

    add_bullet(doc, "If there's a firewall at the KSRCE network level (which there likely is), we need:")
    add_bullet(doc, "Visibility — please share the current firewall rules that affect our machines.", level=1)
    add_bullet(doc, "A process for requesting port openings or rule changes when we need them.", level=1)
    add_bullet(doc,
               "The key ports we'll need open (internal): SSH (22), HTTP/HTTPS (80, 443, 8080), "
               "NFS (2049), and application ports in the range 3000–9999.", level=1)

    add_bullet(doc, "We'll also need outbound internet access from all machines — we need to download software packages, container images, and updates from the internet.")
    add_bullet(doc, "During initial setup, unrestricted outbound access would be ideal.", level=1)
    add_bullet(doc, "Once everything is deployed, we can work together to tighten things down.", level=1)

    add_spacer(doc, 4)
    add_callout(doc,
        "Consideration",
        "During initial setup, unrestricted outbound internet from the compute machines is "
        "strongly preferred. Once the platform is live, the LaaS and KSRCE teams can collaboratively "
        "define and apply appropriate outbound rules together.",
        callout_type="consideration"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 6 — Power & Environment
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Power & Environment")

    add_bullet(doc, "These machines will run 24/7 under heavy compute loads — they're not regular office PCs.")
    add_bullet(doc, "Power consumption:")
    add_bullet(doc, "Each machine can draw up to 800–900 watts at peak.", level=1)
    add_bullet(doc, "The full setup (4 compute nodes + NAS + networking) needs about 3.5–4 kW total.", level=1)
    add_bullet(doc, "Make sure the power circuit can handle this — a dedicated 32A circuit is recommended.", level=1)

    add_bullet(doc, "UPS (Uninterruptible Power Supply):", bold_prefix="")
    add_bullet(doc, "We strongly recommend a 5 kVA online UPS for the entire setup.", level=1)
    add_bullet(doc, "This gives roughly 40–80 minutes of backup time, which is enough for a graceful shutdown if power doesn't return.", level=1)
    add_bullet(doc, "Without a UPS, a sudden power cut can corrupt data, damage hardware, and bring the platform down unexpectedly.", level=1)
    add_bullet(doc, "This is the single most important environmental consideration.", level=1)

    add_bullet(doc, "Cooling:", bold_prefix="")
    add_bullet(doc, "GPU machines generate a lot of heat — especially under sustained load.", level=1)
    add_bullet(doc, "The room should be air-conditioned or well-ventilated.", level=1)
    add_bullet(doc, "Recommended ambient temperature: 18–24°C.", level=1)
    add_bullet(doc, "If it's a dedicated server room, make sure the AC can handle 3–4 kW of continuous heat output.", level=1)

    add_bullet(doc, "Physical security:", bold_prefix="")
    add_bullet(doc, "The machines should be in a locked, access-controlled room.", level=1)
    add_bullet(doc, "Not a general-purpose area where anyone can walk in and accidentally unplug something.", level=1)

    add_spacer(doc, 4)
    add_callout(doc,
        "Critical Information",
        "Power outages without UPS protection are the single most common cause of data corruption "
        "and hardware damage in GPU compute environments. A sudden loss of power during active "
        "workloads can corrupt user data, damage storage drives, and cause irreversible hardware "
        "failures. Installing a UPS is strongly recommended before the platform goes live.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 7 — Technical Support
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Technical Support")

    add_bullet(doc, "A designated point of contact from the KSRCE infra team for:",
               bold_prefix="We'd appreciate ")
    add_bullet(doc, "Network issues (VPN not connecting, firewall blocking something).", level=1)
    add_bullet(doc, "Hardware issues (a machine is unresponsive, needs a physical restart).", level=1)
    add_bullet(doc, "Power/environment issues (UPS alerts, AC failure).", level=1)

    add_bullet(doc, "Ideal response time: within 4 hours for critical issues (machine down, network outage).")
    add_bullet(doc, "Preferred communication: email for non-urgent, phone/WhatsApp for urgent matters.")
    add_bullet(doc, "We don't expect frequent issues, but when something does come up, having a responsive contact saves a lot of time.")

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 8 — Application Server
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. Application Server")

    add_bullet(doc, "one dedicated machine (or VM) to run the LaaS application stack - this is the \"application server\".",
               bold_prefix="Besides the GPU compute nodes, we also need ")
    add_bullet(doc, "This is separate from the GPU compute nodes - the compute nodes run user workloads, the app server runs the platform itself.")

    add_bullet(doc, "Minimum specs:", bold_prefix="")
    add_bullet(doc, "8 vCPU", level=1)
    add_bullet(doc, "16 GB RAM", level=1)
    add_bullet(doc, "256 GB SSD storage", level=1)

    add_bullet(doc, "Recommended specs:", bold_prefix="")
    add_bullet(doc, "16 vCPU", level=1)
    add_bullet(doc, "32 GB RAM", level=1)
    add_bullet(doc, "500 GB SSD (better headroom for logs, monitoring data, database growth)", level=1)

    add_bullet(doc, "This machine runs:", bold_prefix="")
    add_bullet(doc, "Web frontend (Next.js)", level=1)
    add_bullet(doc, "Backend API (NestJS)", level=1)
    add_bullet(doc, "PostgreSQL database", level=1)
    add_bullet(doc, "Keycloak (SSO/identity management)", level=1)
    add_bullet(doc, "Nginx reverse proxy", level=1)
    add_bullet(doc, "Monitoring stack (Prometheus, Grafana, Loki, Alertmanager)", level=1)

    add_bullet(doc, "OS: Ubuntu 22.04.5 LTS (same as compute nodes)")
    add_bullet(doc, "It should be accessible via the same VPN and SSH setup as the compute nodes.")
    add_bullet(doc, "It needs a static IP just like everything else.")

    add_bullet(doc, "If KSRCE has an existing VM infrastructure (VMware, Proxmox, etc.), a VM works perfectly fine for this.")
    add_bullet(doc, "If hosting on a cloud provider (AWS, Azure, etc.) is preferred, that's also fine — just ensure it can reach the internal network where compute nodes are.")

    add_spacer(doc, 4)
    add_callout(doc,
        "Important",
        "This is the machine users interact with when they open the LaaS platform in their browser. "
        "If it goes down, the entire platform UI is inaccessible even though compute nodes may still be running.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 9 — Database Backup & Snapshot Plans
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. Database Backup & Snapshot Plans")

    add_bullet(doc, "The LaaS platform stores critical data in PostgreSQL: user accounts, billing records, session history, compute resource allocation, and configuration.")

    add_bullet(doc, "automated daily database backups (pg_dump) that run at a scheduled time (e.g., 2 AM).",
               bold_prefix="We'll set up ")

    add_bullet(doc, "Backups should be stored on a separate device/location from the database server — ideally on the NAS or a separate backup drive.")

    add_bullet(doc, "at least 30 days of rolling backups (daily snapshots).",
               bold_prefix="We recommend keeping ")

    add_bullet(doc, "Weekly full backups + daily incremental backups is the ideal strategy.")

    add_bullet(doc, "If possible, also set up periodic backup replication to an offsite location (cloud storage bucket, separate physical site) for disaster recovery.")

    add_bullet(doc, "For the NAS (user data storage): ZFS snapshots should be taken daily — ZFS makes this nearly zero-cost.")
    add_bullet(doc, "NAS snapshots should be retained for at least 14 days so users can recover accidentally deleted files.")

    add_bullet(doc, "We'll also need a documented restore procedure — so if something goes wrong, the recovery steps are clear and tested.")

    add_bullet(doc, "KSRCE should provide or allocate a dedicated backup storage device/partition (at least 500 GB) for database backups and snapshots.")

    add_spacer(doc, 4)
    add_callout(doc,
        "Critical",
        "Without a backup strategy, a single database corruption, accidental deletion, or hardware failure "
        "could mean permanent loss of all user data, billing history, and platform configuration. "
        "Backups are not optional.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 10 — Contingency & Disaster Recovery Plans
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "10. Contingency & Disaster Recovery Plans")
    add_body(doc, "What happens when things go wrong? Here's a breakdown by failure type:")
    add_spacer(doc, 4)

    # -- Network/VPN Failure --
    add_heading(doc, "Network / VPN Failure", level=2)
    add_bullet(doc, "If the primary VPN goes down, we need a backup way to reach the machines.")
    add_bullet(doc, "Options: secondary VPN profile, direct SSH via port forwarding on the router, or Tailscale/WireGuard as a fallback VPN.")
    add_bullet(doc, "KSRCE should have someone on-site who can physically check machines and restart network equipment if remote access is completely lost.")
    add_bullet(doc, "Recommended: keep a 4G/LTE USB dongle or mobile hotspot as emergency internet if the primary ISP fails.")

    # -- Power Failure --
    add_heading(doc, "Power Failure", level=2)
    add_bullet(doc, "UPS gives us 40–80 minutes of buffer — during this time, we need a plan:")
    add_bullet(doc, "If power returns within the buffer: nothing to do, machines continue normally.", level=1)
    add_bullet(doc, "If power doesn't return: graceful shutdown procedure must be documented and ideally automated (UPS with NUT/apcupsd can trigger auto-shutdown).", level=1)
    add_bullet(doc, "After power restoration: machines should auto-boot (BIOS AC Power Recovery), services should auto-start (systemd), but someone should verify everything came back cleanly.")
    add_bullet(doc, "Document the boot-up order: NAS first → compute nodes → verify NFS mounts → start platform services.")

    # -- Hardware Failure --
    add_heading(doc, "Hardware Failure", level=2)
    add_bullet(doc, "If a compute node fails: the platform continues running on remaining nodes (reduced capacity, not full outage).")
    add_bullet(doc, "If the NAS fails: this is critical — all user data is on the NAS. RAID/ZFS mirroring protects against single-drive failures, but not against controller failure.")
    add_bullet(doc, "If the application server fails: entire platform UI goes down — this is the single point of failure.")
    add_bullet(doc, "Recommendation: have documented steps for \"what to do if X fails\" for each critical component.")
    add_bullet(doc, "Keep spare drives, cables, and at minimum one spare power supply.")

    # -- Data Centre Standards --
    add_heading(doc, "Data Centre Standards (ODC Best Practices)", level=2)
    add_bullet(doc, "The server room should follow basic ODC/data centre standards even if it's a small setup:")
    add_bullet(doc, "Raised flooring or cable management trays to avoid cable clutter.", level=1)
    add_bullet(doc, "Proper rack mounting for all machines (not stacked on a desk).", level=1)
    add_bullet(doc, "Fire suppression system (at minimum a CO₂ fire extinguisher in the room).", level=1)
    add_bullet(doc, "Access control logs (who entered the room and when).", level=1)
    add_bullet(doc, "Temperature and humidity monitoring with alerts.", level=1)
    add_bullet(doc, "Clear labeling of all cables, ports, and machines.", level=1)
    add_bullet(doc, "Dust-free environment — regular cleaning schedule for the room and equipment air filters.", level=1)

    add_spacer(doc, 4)
    add_callout(doc,
        "Critical",
        "A well-organized server room isn't just about aesthetics — it directly impacts uptime, troubleshooting speed, "
        "and hardware lifespan. A cable pulled accidentally, a machine overheating due to blocked airflow, or dust buildup "
        "causing component failure are all preventable with basic data centre hygiene.",
        callout_type="critical"
    )

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 11 — Handover Checklist
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "11. Handover Checklist")
    add_body(doc, "Before handing over access to us, please make sure the following are ready:")
    add_spacer(doc, 4)

    checklist_items = [
        "Ubuntu 22.04.5 LTS installed on all machines",
        "SSH server installed, enabled, and starts on boot",
        "laas-admin user with sudo privileges created on all machines",
        "Static IPs assigned to all machines and the NAS",
        "NAS configured and accessible from all machines via NFS",
        "VPN access credentials provided to the LaaS team",
        "Firewall rules documented / required ports opened",
        "UPS installed and connected",
        "Cooling/ventilation verified for the machine room",
        "Machine hostname-to-IP mapping document provided",
        "Internet connectivity verified on all machines (can ping google.com, can apt update)",
        "KSRCE technical support contact details provided",
        "Application server (or VM) provisioned with minimum specs",
        "Backup storage device/partition allocated (≥500 GB)",
        "UPS configured with auto-shutdown capability (NUT/apcupsd)",
        "Emergency contact procedure documented for after-hours incidents",
        "Boot-up order documented (NAS → compute → services)",
        "Server room meets basic ODC standards (rack, cooling, fire, access control)",
        "Backup/restore procedure documented and tested at least once",
    ]
    for item in checklist_items:
        add_checklist_item(doc, item)

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 12 — Things You Might Want to Consider
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "12. Things You Might Want to Consider")
    add_body(doc,
             "(These are not strict requirements, but things worth thinking about)",
             italic=True, color=GREY_TEXT)
    add_spacer(doc, 4)

    optional_items = [
        (
            "Redundant internet connection — ",
            "if you have only one ISP and it goes down, the entire platform becomes inaccessible "
            "to remote users. A backup ISP (even a slower one) can keep things running during outages."
        ),
        (
            "Monitoring/alerting for power — ",
            "if the UPS kicks in, someone should be notified immediately so they can take action "
            "before the battery runs out."
        ),
        (
            "Physical KVM or IPMI access — ",
            "if a machine completely locks up and SSH stops responding, having a way to remotely "
            "reboot it (via IPMI/BMC) or having someone physically access it quickly saves downtime."
        ),
        (
            "Scheduled maintenance windows — ",
            "plan a regular maintenance window (e.g., first Sunday of every month, 2 AM – 6 AM) "
            "for updates, reboots, and hardware checks. Communicate this to users in advance."
        ),
        (
            "Spare parts — ",
            "having a spare power supply, extra RAM stick, or spare network cable on hand can mean "
            "the difference between 30 minutes of downtime and 2 weeks waiting for a replacement."
        ),
        (
            "Network documentation — ",
            "a simple network diagram showing how the machines, NAS, switch, and internet gateway "
            "are connected makes troubleshooting 10× faster."
        ),
        (
            "Backup power for network equipment — ",
            "the UPS should also cover the network switch and router, not just the compute machines. "
            "If the switch loses power, all machines lose connectivity even if they're still running."
        ),
        (
            "Environmental monitoring — ",
            "a simple temperature sensor in the server room with alerts (even a smart home sensor) "
            "can warn you before overheating becomes a problem."
        ),
        (
            "Access for physical maintenance — ",
            "make sure the machines are physically accessible (not crammed behind furniture) so "
            "drives, cables, and components can be serviced when needed."
        ),
        (
            "Database backup testing — ",
            "schedule a monthly test restore of the database backup to make sure backups are actually working "
            "and not corrupted. A backup you've never tested is not a real backup."
        ),
        (
            "Offsite backup replication — ",
            "if all backups are stored on the same physical site, a fire, flood, or theft means you lose everything "
            "including backups. Even a simple cloud storage sync (S3, Azure Blob) for weekly backups adds significant protection."
        ),
    ]

    for bold_part, rest_part in optional_items:
        add_bullet(doc, rest_part, bold_prefix=bold_part)

    # ════════════════════════════════════════════════════════════════════════
    # SECTION 13 — Contact Information
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "13. Contact Information")

    add_bullet(doc, "[To be filled]", bold_prefix="LaaS Platform Team:  ")
    add_bullet(doc, "[To be filled]", bold_prefix="KSRCE Infrastructure Team:  ")

    add_spacer(doc, 16)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("— End of Document —")
    r.font.name  = "Calibri"
    r.font.size  = Pt(10)
    r.font.color.rgb = LIGHT_GREY
    r.italic = True

    # ── Footer ──
    add_page_numbers(doc)

    doc.save(OUTPUT_PATH)
    print(f"Document generated successfully:\n  {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
