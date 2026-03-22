"""
LaaS Technical Architecture Document Generator
Generates a professional .docx report using python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# ============================================================================
# COLOR CONSTANTS
# ============================================================================
COLORS = {
    'body_text': '2C2C2C',
    'section_header_bg': 'E8D5B7',  # warm tan
    'table_header_bg': 'E8D5B7',    # warm tan
    'table_alt_row': 'FDF8F0',       # light warm
    'callout_bg': 'F5E6D3',          # warm sand
    'critical_bg': 'F4D3D3',         # soft rose
    'inline_code_bg': 'E8D5B7',
    'section_number': '8B7355',      # warm brown
    'border': 'C4B8A8',              # muted tan border
}

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def set_cell_shading(cell, fill_color):
    """Set background shading for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_paragraph_shading(paragraph, fill_color):
    """Set background shading for a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, color='C4B8A8'):
    """Set borders for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for border_name, size in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if size is not None:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(size))
            border.set(qn('w:color'), color)
            tcBorders.append(border)
    
    tcPr.append(tcBorders)

def set_run_font(run, font_name='Segoe UI', size=11, bold=False, italic=False, color=None):
    """Configure font properties for a run."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_section_header(doc, number, title):
    """Add a styled section header with background band."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    set_paragraph_shading(p, COLORS['section_header_bg'])
    
    # Add section number
    run_num = p.add_run(f'  SECTION {number} — ')
    set_run_font(run_num, 'Consolas', 14, bold=True, color=COLORS['section_number'])
    
    # Add title
    run_title = p.add_run(title.upper() + '  ')
    set_run_font(run_title, 'Segoe UI', 14, bold=True, color=COLORS['body_text'])
    
    return p

def add_subsection(doc, title):
    """Add a bold subsection heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, 'Segoe UI', 12, bold=True, color=COLORS['body_text'])
    return p

def add_body(doc, text):
    """Add a standard body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, 'Segoe UI', 11, color=COLORS['body_text'])
    return p

def add_tree_bullet(doc, points):
    """Add tree-style bullet points with |_> prefix."""
    # Add the leading pipe
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('|')
    set_run_font(run, 'Consolas', 10, color=COLORS['section_number'])
    
    # Add each bullet point
    for point in points:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        
        run_prefix = p.add_run('|_> ')
        set_run_font(run_prefix, 'Consolas', 10, color=COLORS['section_number'])
        
        run_text = p.add_run(point)
        set_run_font(run_text, 'Segoe UI', 11, color=COLORS['body_text'])
    
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header shading and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        set_cell_shading(cell, COLORS['table_header_bg'])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, 'Segoe UI', 10, bold=True, color=COLORS['body_text'])
    
    # Style data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        # Alternate row shading
        bg_color = COLORS['table_alt_row'] if row_idx % 2 == 1 else None
        
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            if bg_color:
                set_cell_shading(cell, bg_color)
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_run_font(run, 'Segoe UI', 10, color=COLORS['body_text'])
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    # Add spacing after table
    doc.add_paragraph()
    
    return table

def add_callout(doc, title, text, callout_type='info'):
    """Add a bordered callout box."""
    color = COLORS['critical_bg'] if callout_type == 'critical' else COLORS['callout_bg']
    border_color = 'D9534F' if callout_type == 'critical' else COLORS['section_number']
    
    # Create a single-cell table for the callout
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, color)
    set_cell_borders(cell, top=4, bottom=4, left=24, right=4, color=border_color)
    
    # Add title
    if title:
        p_title = cell.paragraphs[0]
        run = p_title.add_run(title)
        set_run_font(run, 'Segoe UI', 11, bold=True, color=COLORS['body_text'])
        p_title.paragraph_format.space_after = Pt(4)
    
    # Add body text
    p_body = cell.add_paragraph()
    run = p_body.add_run(text)
    set_run_font(run, 'Segoe UI', 10, color=COLORS['body_text'])
    
    doc.add_paragraph()
    return table

def add_diagram_placeholder(doc, diagram_name):
    """Add a placeholder box for diagrams."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F5F5F5')
    set_cell_borders(cell, top=8, bottom=8, left=8, right=8, color='CCCCCC')
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(20)
    
    run = p.add_run(f'[DIAGRAM: {diagram_name} — Insert from docs/diagrams/]')
    set_run_font(run, 'Segoe UI', 11, italic=True, color='666666')
    
    doc.add_paragraph()
    return table

def add_code_block(doc, code):
    """Add a monospace code block."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'F8F8F8')
    set_cell_borders(cell, top=4, bottom=4, left=4, right=4, color='DDDDDD')
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    for line in code.split('\n'):
        if p.text:  # Add new paragraph for subsequent lines
            p = cell.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, 'Consolas', 9, color=COLORS['body_text'])
    
    doc.add_paragraph()
    return table

def add_inline_code(paragraph, text):
    """Add inline code styled text to a paragraph."""
    run = paragraph.add_run(f' {text} ')
    set_run_font(run, 'Consolas', 10, color=COLORS['body_text'])
    # Note: inline background shading is complex in docx, using font styling instead
    return run

def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()

def add_cover_page(doc):
    """Generate the cover page."""
    # Add some spacing at top
    for _ in range(4):
        doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('LAAS — LAB AS A SERVICE')
    set_run_font(run, 'Segoe UI', 28, bold=True, color=COLORS['body_text'])
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run('Technical Architecture Document')
    set_run_font(run, 'Segoe UI', 18, color=COLORS['section_number'])
    
    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run('━' * 40)
    set_run_font(run, 'Consolas', 12, color=COLORS['section_header_bg'])
    
    # Version info
    for _ in range(2):
        doc.add_paragraph()
    
    info_items = [
        ('Version', 'v1.0'),
        ('Date', 'March 2026'),
        ('Status', 'POC Validated | Production Ready'),
        ('Prepared for', 'KSRCE-GKT AI Lab'),
        ('Classification', 'CONFIDENTIAL — Internal Use Only'),
    ]
    
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f'{label}: ')
        set_run_font(run_label, 'Segoe UI', 11, bold=True, color=COLORS['section_number'])
        run_value = p.add_run(value)
        set_run_font(run_value, 'Segoe UI', 11, color=COLORS['body_text'])

def add_toc(doc):
    """Generate a manual table of contents."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TABLE OF CONTENTS')
    set_run_font(run, 'Segoe UI', 16, bold=True, color=COLORS['body_text'])
    
    doc.add_paragraph()
    
    toc_entries = [
        ('01', 'EXECUTIVE SUMMARY'),
        ('02', 'PLATFORM OVERVIEW & OBJECTIVES'),
        ('03', 'SYSTEM ARCHITECTURE (HLD)'),
        ('04', 'HARDWARE SPECIFICATIONS'),
        ('05', 'GPU FRACTIONAL DESKTOP ARCHITECTURE'),
        ('06', 'INFRASTRUCTURE ARCHITECTURE'),
        ('07', 'STORAGE ARCHITECTURE (ZFS/NFS)'),
        ('08', 'AUTHENTICATION & AUTHORIZATION'),
        ('09', 'MONITORING & OBSERVABILITY'),
        ('10', 'APPLICATION ARCHITECTURE'),
        ('11', 'BILLING & PAYMENT SYSTEM'),
        ('12', 'SESSION LIFECYCLE'),
        ('13', 'ISOLATION & MULTI-TENANCY'),
        ('14', 'DATA FLOW & INTEGRATION MODEL'),
        ('15', 'FAILURE ANALYSIS & MITIGATIONS'),
        ('16', 'ARCHITECTURE RECONSIDERATIONS & LESSONS'),
        ('17', 'RISK ASSESSMENT'),
        ('18', 'TECHNOLOGY STACK SUMMARY'),
        ('A', 'APPENDIX A — API ENDPOINT REFERENCE'),
        ('B', 'APPENDIX B — ALERT RULES REFERENCE'),
    ]
    
    for num, title in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)
        
        run_num = p.add_run(f'{num}  ')
        set_run_font(run_num, 'Consolas', 11, color=COLORS['section_number'])
        
        run_title = p.add_run(title)
        set_run_font(run_title, 'Segoe UI', 11, color=COLORS['body_text'])

# ============================================================================
# SECTION CONTENT FUNCTIONS
# ============================================================================

def add_section_01(doc):
    """Section 01 — Executive Summary"""
    add_section_header(doc, '01', 'Executive Summary')
    
    add_body(doc, 
        'LaaS (Lab as a Service) delivers GPU-accelerated remote desktops and ephemeral compute '
        'environments to university students via fractional GPU sharing on RTX 5090 hardware. '
        'The platform combines Selkies EGL Desktop, WebRTC streaming, ZFS/NFS persistent storage, '
        'and enterprise-grade monitoring to create a unique educational computing platform.')
    
    add_callout(doc, 'KEY DIFFERENTIATOR', 
        '4-8 concurrent GPU-accelerated full desktop sessions per consumer GPU — no commercial '
        'platform offers this capability. Achieved through HAMi-core CUDA API interception and '
        'NVIDIA MPS (Multi-Process Service) for driver-level VRAM and compute partitioning.',
        'info')
    
    add_subsection(doc, 'Technology Stack Overview')
    add_tree_bullet(doc, [
        'Frontend: Next.js 15, React 19, TypeScript 5.7, Tailwind CSS 4, Zustand 5',
        'Backend: NestJS 11 + Fastify, Prisma 6, PostgreSQL, Redis',
        'Host Services: Python Flask for storage provisioning (ZFS/NFS)',
        'Infrastructure: Docker CE, nvidia-container-toolkit, HAMi-core, CUDA MPS',
        'Monitoring: Prometheus 3.4, Grafana 11.5, Loki 3.4, Alertmanager 0.28',
    ])

def add_section_02(doc):
    """Section 02 — Platform Overview & Objectives"""
    add_section_header(doc, '02', 'Platform Overview & Objectives')
    
    add_subsection(doc, 'Business Model')
    add_body(doc,
        '4 on-premises machines monetized for multi-tenant use. Revenue streams include university '
        'semester contracts, per-user storage allocation, per-session compute billing, and mentorship '
        'services planned for Phase 3.')
    
    add_subsection(doc, 'Target Users')
    add_tree_bullet(doc, [
        'University students — Stateful desktop sessions with persistent storage + ephemeral compute',
        'Faculty — Teaching labs, course management, assignment grading',
        'Researchers — High-performance GPU workloads, exclusive machine access',
        'Public users — Ephemeral compute only (Jupyter, Code-Server, SSH CLI)',
    ])
    
    add_subsection(doc, 'Operational Modes')
    add_body(doc, 'Mode 1: Stateful GUI Desktop — Full KDE Plasma desktop via Selkies EGL, '
        '15GB persistent ZFS storage, CPU+GPU compute tiers, session booking required.')
    add_body(doc, 'Mode 2: Ephemeral Compute — Jupyter notebooks, Code-Server, SSH CLI access, '
        'no persistent storage, no booking required, pay-per-use.')
    
    add_subsection(doc, 'Compute Tier Configuration')
    add_table(doc, 
        ['Tier', 'vCPU', 'RAM', 'GPU VRAM', 'Best For'],
        [
            ['Starter', '2', '4GB', 'None', 'Light coding, documentation'],
            ['Standard', '4', '8GB', 'None', 'MATLAB, general development'],
            ['Pro', '4', '8GB', '4GB', 'Blender, ML inference'],
            ['Power', '8', '16GB', '8GB', 'Heavy simulation, 3D rendering'],
            ['Max', '8', '16GB', '16GB', 'Large model training'],
            ['Full Machine', '16', '64GB', '32GB', 'Exclusive research/DL'],
        ],
        col_widths=[1.2, 0.8, 0.8, 1.0, 2.2])

def add_section_03(doc):
    """Section 03 — System Architecture (HLD)"""
    add_section_header(doc, '03', 'System Architecture (HLD)')
    
    add_diagram_placeholder(doc, '01-hld-application-architecture')
    
    add_subsection(doc, 'Layered Architecture')
    
    add_body(doc, 'Access Layer — Cloudflare Tunnel for public HTTPS ingress, Tailscale for '
        'administrative access. No public IP addresses exposed on compute nodes.')
    
    add_body(doc, 'Frontend Layer — Next.js 15 with React 19, server-side rendering, Zustand for '
        'state management, Tailwind CSS 4 for styling. Lambda.ai Utilitarian Minimalism design system.')
    
    add_body(doc, 'Backend Layer — NestJS 11 running on Fastify for maximum performance. JWT '
        'authentication with Keycloak SSO integration. RESTful API with OpenAPI documentation.')
    
    add_body(doc, 'Data Layer — PostgreSQL via Prisma 6 ORM, Redis for caching and session state, '
        'ZFS datasets for user storage with NFS export.')
    
    add_body(doc, 'Infrastructure Layer — 4× compute nodes (Proxmox VE 8.x), TrueNAS Scale NAS, '
        'Docker CE with nvidia-container-toolkit, HAMi-core for GPU fractioning, CUDA MPS daemon.')

def add_section_04(doc):
    """Section 04 — Hardware Specifications"""
    add_section_header(doc, '04', 'Hardware Specifications')
    
    add_subsection(doc, 'Compute Node Fleet')
    add_table(doc,
        ['Component', 'Model', 'Qty'],
        [
            ['CPU', 'AMD Ryzen 9 9950X3D (16C/32T)', '4'],
            ['GPU', 'Zotac RTX 5090 Solid OC 32GB', '4'],
            ['RAM', 'G.SKILL DDR5 6000MHz 64GB', '4'],
            ['Storage', 'Samsung 990 EVO Plus 2TB', '4'],
            ['Motherboard', 'ASUS ProArt X670E-Creator', '4'],
            ['PSU', 'Corsair AX1600i 1600W', '4'],
        ],
        col_widths=[1.5, 3.0, 0.8])
    
    add_callout(doc, 'FLEET TOTAL CAPACITY',
        '64 CPU cores • 256GB system RAM • 128GB GPU VRAM • 8TB NVMe storage',
        'info')
    
    add_subsection(doc, 'Network Attached Storage')
    add_body(doc, 'TrueNAS Scale with 4×4TB drives in RAIDZ1 configuration providing approximately '
        '12TB usable capacity. 10GbE connectivity to compute nodes.')
    
    add_subsection(doc, 'NVMe Partition Layout (Per Node)')
    add_table(doc,
        ['Partition', 'Size', 'Filesystem', 'Purpose'],
        [
            ['p1', '128GB', 'EXT4', 'Proxmox OS + swap'],
            ['p2', '500GB', 'EXT4', 'Docker images + layers'],
            ['p3', '1100GB', 'LVM', 'Container scratch space'],
            ['p4', '240GB', 'EXT4', 'Shared read-only datasets'],
        ],
        col_widths=[1.0, 1.0, 1.2, 3.0])
    
    add_subsection(doc, 'Network Infrastructure')
    add_body(doc, 'Mikrotik CRS309 10GbE switch with 4 VLANs:')
    add_tree_bullet(doc, [
        'VLAN 10: Management — Proxmox UI, SSH, monitoring endpoints',
        'VLAN 20: Stateful — Desktop session traffic, NFS mounts',
        'VLAN 30: Ephemeral — Jupyter, Code-Server, SSH sessions',
        'VLAN 40: Proxy — Cloudflare Tunnel, external ingress',
    ])

def add_section_05(doc):
    """Section 05 — GPU Fractional Desktop Architecture"""
    add_section_header(doc, '05', 'GPU Fractional Desktop Architecture')
    
    add_callout(doc, 'BREAKTHROUGH CAPABILITY',
        'This section describes the core innovation of LaaS: running 4-8 concurrent GPU-accelerated '
        'full desktop sessions on a single consumer RTX 5090 GPU. No other commercial platform offers '
        'this capability.',
        'critical')
    
    add_diagram_placeholder(doc, '06-gpu-fractional-architecture')
    
    add_subsection(doc, 'Technology Stack')
    
    add_body(doc, '1. Selkies EGL Desktop — Full KDE Plasma desktop environment running inside Docker '
        'containers. GPU rendering via VirtualGL EGL backend. WebRTC streaming with NVENC hardware '
        'encoding for sub-50ms latency.')
    
    add_body(doc, '2. HAMi-core (libvgpu.so) — CUDA API interception library providing software-level '
        'GPU partitioning:')
    add_tree_bullet(doc, [
        'cuMemAlloc interception — Enforces per-container VRAM allocation limits',
        'cuDeviceGetAttribute override — Reports fake VRAM visibility to applications',
        'cuLaunchKernel throttling — Limits compute kernel execution rate',
    ])
    
    add_body(doc, '3. CUDA MPS (Multi-Process Service) — NVIDIA driver-level GPU sharing:')
    add_tree_bullet(doc, [
        'CUDA_MPS_PINNED_DEVICE_MEM_LIMIT — Secondary VRAM enforcement',
        'CUDA_MPS_ACTIVE_THREAD_PERCENTAGE — SM (Streaming Multiprocessor) partitioning',
        'Isolated VA (Virtual Address) spaces — Memory isolation between clients',
    ])
    
    add_subsection(doc, 'Competitive Advantage')
    add_table(doc,
        ['Platform', 'Concurrent/GPU', 'VRAM Tiers', 'Desktop Stream'],
        [
            ['RunPod', '1 (exclusive)', 'All-or-nothing', 'No'],
            ['Vast.ai', '1 (exclusive)', 'Exclusive', 'No'],
            ['Lambda Labs', '1 per user', 'Exclusive', 'No'],
            ['Google Colab', 'Shared (preemption)', 'Fractional', 'No (notebook only)'],
            ['LaaS', '4-8', '4/8/16/32 GB', 'Yes (WebRTC)'],
        ],
        col_widths=[1.5, 1.5, 1.5, 1.8])

def add_section_06(doc):
    """Section 06 — Infrastructure Architecture"""
    add_section_header(doc, '06', 'Infrastructure Architecture')
    
    add_diagram_placeholder(doc, '03-infrastructure-architecture')
    
    add_subsection(doc, 'Node Configuration')
    add_body(doc, 'Each compute node runs:')
    add_tree_bullet(doc, [
        'Proxmox VE 8.x — Hypervisor and container management',
        'NVIDIA Driver 555.x — GPU driver with CUDA 12.x support',
        'nvidia-container-toolkit — Docker GPU integration',
        'CUDA MPS Daemon — Multi-process GPU sharing service',
        'Docker CE — Container runtime',
        'HAMi-core (libvgpu.so) — Preloaded via LD_PRELOAD for CUDA interception',
    ])
    
    add_subsection(doc, 'Dynamic Scheduling')
    add_body(doc, 'No static node roles — all nodes capable of running any workload type. Session '
        'scheduler evaluates available capacity (CPU, RAM, VRAM) across fleet and assigns sessions '
        'to optimal node.')
    
    add_callout(doc, 'OPERATIONAL REQUIREMENT',
        'HDMI dummy dongles required on all GPUs for headless rendering. Without a connected '
        'display (or dummy plug), the GPU will not initialize display outputs required for EGL.',
        'info')

def add_section_07(doc):
    """Section 07 — Storage Architecture (ZFS/NFS)"""
    add_section_header(doc, '07', 'Storage Architecture (ZFS/NFS)')
    
    add_diagram_placeholder(doc, '05-file-store-architecture')
    
    add_subsection(doc, 'Design Principles')
    add_body(doc, 'Hard partition between stateful user data and ephemeral compute data. User '
        'files persist across sessions; container scratch space is destroyed on session end.')
    
    add_subsection(doc, 'Per-User ZFS Dataset')
    add_code_block(doc, '''# Dataset path
datapool/users/<storage_uid>

# Quota enforcement
zfs set quota=15G datapool/users/<storage_uid>

# Automatic snapshots every 6 hours
zfs-auto-snapshot --label=hourly --keep=4 datapool/users/<storage_uid>''')
    
    add_subsection(doc, 'NFS Export Configuration')
    add_body(doc, 'Individual NFS export per user dataset, mounted on compute nodes at '
        '/mnt/nfs/users/<uid>. Bind mounted into containers:')
    add_code_block(doc, 'docker run -v /mnt/nfs/users/<uid>:/home/ubuntu ...')
    
    add_callout(doc, 'LESSON LEARNED',
        'NFS gotcha discovered during POC: Must export individual child datasets, not the parent '
        'dataset. Exporting the parent causes permission and mount issues.',
        'info')
    
    add_subsection(doc, 'File Operations Pipeline')
    add_body(doc, 'Browser → NestJS Backend → Python Flask Host Service → ZFS')
    add_body(doc, 'Supported operations: List, Create folder, Upload (500MB limit), Download, Delete.')

def add_section_08(doc):
    """Section 08 — Authentication & Authorization"""
    add_section_header(doc, '08', 'Authentication & Authorization')
    
    add_diagram_placeholder(doc, '04-auth-architecture')
    
    add_subsection(doc, 'Authentication Paths')
    add_tree_bullet(doc, [
        'University SSO — Keycloak brokering to institution IdP (Google Workspace or LDAP)',
        'Public OAuth — Direct Google/GitHub authentication',
        'Local Authentication — Email + password with OTP verification',
    ])
    
    add_subsection(doc, 'Storage Provisioning')
    add_body(doc, 'SSO users from registered universities automatically receive 15GB ZFS dataset '
        'provisioned on first login. Public OAuth users do not receive persistent storage (ephemeral only).')
    
    add_subsection(doc, 'Token Lifecycle')
    add_body(doc, 'JWT access tokens (15min expiry) + refresh tokens (7 day expiry). Tokens stored '
        'in localStorage. Automatic refresh on 401 response.')
    
    add_subsection(doc, 'RBAC Model')
    add_body(doc, 'Users → UserOrgRoles → Organizations → Roles → Permissions')
    add_body(doc, 'Permissions are action-based (e.g., session:create, storage:upload, billing:view).')

def add_section_09(doc):
    """Section 09 — Monitoring & Observability"""
    add_section_header(doc, '09', 'Monitoring & Observability')
    
    add_diagram_placeholder(doc, '07-monitoring-architecture')
    
    add_subsection(doc, 'Monitoring Stack')
    add_table(doc,
        ['Component', 'Version', 'Purpose'],
        [
            ['Prometheus', '3.4.0', 'Metrics collection and alerting'],
            ['Grafana', '11.5.0', 'Visualization and dashboards'],
            ['Loki', '3.4.2', 'Log aggregation'],
            ['Alertmanager', '0.28.1', 'Alert routing and deduplication'],
            ['Uptime Kuma', '1.23.x', 'Synthetic monitoring'],
        ],
        col_widths=[1.5, 1.0, 3.5])
    
    add_subsection(doc, 'Metrics Exporters')
    add_table(doc,
        ['Exporter', 'Port', 'Metrics'],
        [
            ['DCGM', '9400', 'GPU temp, power, VRAM, ECC errors'],
            ['node-exporter', '9100', 'CPU, RAM, disk, network, NFS'],
            ['cAdvisor', '8999', 'Container CPU%, RAM%, I/O'],
            ['mps-exporter', '9500', 'MPS state, faults, NVENC sessions'],
            ['session-exporter', '9501', 'Active sessions, VRAM allocation'],
            ['Blackbox', '9115', 'HTTP/TCP endpoint probes'],
        ],
        col_widths=[1.5, 0.8, 3.7])
    
    add_subsection(doc, 'Alert Categories')
    add_tree_bullet(doc, [
        'GPU Alerts — Temperature, power draw, VRAM overcommit, MPS faults, ECC errors',
        'Session Alerts — Startup timeout, concurrent limit, WebRTC FPS degradation',
        'Host Alerts — Disk space, NFS mount health, CPU/RAM saturation',
        'Security Alerts — Failed login attempts, suspicious activity patterns',
    ])
    
    add_subsection(doc, 'Log Aggregation')
    add_body(doc, 'Promtail agents scrape container logs, systemd journal, and CUDA MPS logs. '
        'Forwarded to Loki for storage. Grafana LogQL for querying and visualization.')

def add_section_10(doc):
    """Section 10 — Application Architecture"""
    add_section_header(doc, '10', 'Application Architecture')
    
    add_subsection(doc, '10.1 Backend (NestJS 11 + Fastify)')
    add_body(doc, 'Modular architecture with clear separation of concerns:')
    add_table(doc,
        ['Module', 'Responsibility'],
        [
            ['AuthModule', 'JWT issuance, Keycloak integration, SSO flows'],
            ['StorageModule', 'File operations, ZFS provisioning orchestration'],
            ['DashboardModule', 'User metrics, session history, billing summary'],
            ['PaymentModule', 'Razorpay integration, wallet management'],
            ['MailModule', 'Transactional emails, OTP delivery'],
            ['PrismaModule', 'Database connection, query building'],
        ],
        col_widths=[1.8, 4.2])
    
    add_body(doc, 'Key dependencies: Prisma 6.0, Razorpay 2.9.6, Keycloak Admin 26.5.5, PDFKit 0.18, '
        'bcrypt, Zod for validation.')
    
    add_subsection(doc, '10.2 Frontend (Next.js 15)')
    add_body(doc, 'Technology stack: React 19, TypeScript 5.7, Tailwind CSS 4, Zustand 5 for state, '
        'Recharts 3.8 for charts, Radix UI for accessible components.')
    add_body(doc, 'Design system: Lambda.ai Utilitarian Minimalism — Clean typography, subtle '
        'animations, dark/light mode support with warm color palette.')
    
    add_subsection(doc, '10.3 Database (PostgreSQL + Prisma)')
    add_diagram_placeholder(doc, '10-database-erd')
    
    add_table(doc,
        ['Domain', 'Tables', 'Purpose'],
        [
            ['Auth & Identity', 'users, user_profiles, refresh_tokens, login_history, otp_verification', 'User management'],
            ['Organizations', 'universities, university_idp_configs, departments, user_groups', 'Multi-tenancy'],
            ['Storage', 'user_storage_volumes, storage_extensions, os_switch_history', 'Persistent storage'],
            ['Infrastructure', 'nodes, base_images, node_base_images', 'Fleet management'],
            ['Compute', 'compute_configs, compute_config_access', 'Tier definitions'],
            ['Sessions', 'bookings, sessions, session_events', 'Session lifecycle'],
            ['Billing', 'wallets, wallet_holds, wallet_transactions, billing_charges, payment_transactions, invoices', 'Payments'],
            ['Academic', 'courses, labs, lab_assignments, lab_submissions, lab_grades', 'LMS integration'],
        ],
        col_widths=[1.3, 3.5, 1.5])
    
    add_subsection(doc, '10.4 Python Host Service')
    add_body(doc, 'Flask application running on port 9999 on each compute node. Provides ZFS '
        'provisioning and file operations. Secured with X-Provision-Secret header, path traversal '
        'protection, and storageUid validation.')

def add_section_11(doc):
    """Section 11 — Billing & Payment System"""
    add_section_header(doc, '11', 'Billing & Payment System')
    
    add_diagram_placeholder(doc, '08-billing-payment-flow')
    
    add_subsection(doc, 'Payment Integration')
    add_body(doc, 'Razorpay payment gateway (test mode during POC, INR currency). Flow:')
    add_tree_bullet(doc, [
        'User initiates wallet recharge → Backend creates Razorpay order',
        'User completes payment in Razorpay UI',
        'Razorpay webhook → Backend verifies signature → Credits wallet',
        'Invoice generated with PDFKit → Emailed to user',
    ])
    
    add_subsection(doc, 'Wallet Model')
    add_body(doc, 'Prepaid balance stored in paise (1/100 of INR). Wallet holds placed for active '
        'sessions to reserve funds. Spend limits configurable per user/organization.')
    
    add_subsection(doc, 'Per-Session Billing')
    add_body(doc, 'Billing formula: duration (minutes) × tier rate = charge amount')
    add_body(doc, 'Charge deducted from wallet on session end. Hold released after successful capture.')

def add_section_12(doc):
    """Section 12 — Session Lifecycle"""
    add_section_header(doc, '12', 'Session Lifecycle')
    
    add_diagram_placeholder(doc, '09-session-lifecycle')
    
    add_subsection(doc, 'Session Launch Flow')
    add_tree_bullet(doc, [
        'User selects compute tier → Backend checks fleet capacity',
        'Backend creates session record → Calls compute node API',
        'Node executes docker run with HAMi+MPS environment variables',
        'Container starts → Health poll until ready (30s timeout)',
        'WebRTC URL returned to user → Desktop stream begins',
    ])
    
    add_subsection(doc, 'Docker Run Configuration')
    add_code_block(doc, '''docker run -d \\
  --name session-<session_id> \\
  --gpus all \\
  -e CUDA_MPS_PINNED_DEVICE_MEM_LIMIT=0=8G \\
  -e CUDA_MPS_ACTIVE_THREAD_PERCENTAGE=25 \\
  -e LD_PRELOAD=/usr/lib/libvgpu.so \\
  -e VGPU_MEMORY_LIMIT=8589934592 \\
  -v /mnt/nfs/users/<uid>:/home/ubuntu \\
  selkies-gstreamer:latest''')
    
    add_subsection(doc, 'Session States')
    add_body(doc, 'pending → starting → running → reconnecting → stopping → ended/failed')
    
    add_subsection(doc, 'Crash Recovery')
    add_body(doc, 'MPS fault detection → Watchdog triggers auto-recovery → Container restart → '
        'Session reconnection prompt to user.')

def add_section_13(doc):
    """Section 13 — Isolation & Multi-Tenancy"""
    add_section_header(doc, '13', 'Isolation & Multi-Tenancy')
    
    add_table(doc,
        ['Layer', 'Isolation Level', 'Mechanism'],
        [
            ['CPU/RAM', 'Hard', 'cgroups v2'],
            ['Filesystem', 'Hard', 'Docker RO layers + NFS mount'],
            ['Process', 'Strong', 'PID namespaces'],
            ['Network', 'Strong', 'Network namespaces + VLANs'],
            ['GPU VRAM', 'Enforced', 'HAMi-core + MPS'],
            ['GPU Compute', 'Enforced', 'HAMi-core + MPS'],
            ['GPU VA Space', 'Isolated', 'MPS (Volta+ GPUs)'],
            ['GPU L2/Bandwidth', 'Shared', 'Cannot partition consumer GPUs'],
        ],
        col_widths=[1.5, 1.2, 3.3])
    
    add_callout(doc, 'PERFORMANCE CONSIDERATION',
        'Honest assessment: 10-30% performance variability under peak co-tenancy conditions due to '
        'shared GPU L2 cache and memory bandwidth. This is acceptable for educational workloads and '
        'significantly better than preemption-based alternatives.',
        'info')

def add_section_14(doc):
    """Section 14 — Data Flow & Integration Model"""
    add_section_header(doc, '14', 'Data Flow & Integration Model')
    
    add_diagram_placeholder(doc, '02-data-flow-diagram')
    
    add_subsection(doc, 'Primary Data Flows')
    add_tree_bullet(doc, [
        'User Authentication — Browser → Frontend → Backend → Keycloak/Database',
        'Session Launch — Frontend → Backend → Node Scheduler → Docker → WebRTC',
        'File Operations — Frontend → Backend → Flask Host Service → ZFS/NFS',
        'Metrics Collection — Exporters → Prometheus → Grafana → Alertmanager',
        'Payments — Frontend → Backend → Razorpay → Webhook → Wallet',
    ])
    
    add_subsection(doc, 'External Integrations')
    add_table(doc,
        ['System', 'Integration Type', 'Data Flow'],
        [
            ['Keycloak', 'OIDC/OAuth2', 'SSO authentication, token validation'],
            ['Razorpay', 'REST API + Webhooks', 'Payment orders, verification'],
            ['Cloudflare', 'Tunnel', 'HTTPS ingress, DDoS protection'],
            ['Tailscale', 'WireGuard mesh', 'Admin VPN access'],
        ],
        col_widths=[1.5, 1.5, 3.0])

def add_section_15(doc):
    """Section 15 — Failure Analysis & Mitigations"""
    add_section_header(doc, '15', 'Failure Analysis & Mitigations')
    
    add_subsection(doc, 'Architecture Evolution')
    add_table(doc,
        ['Version', 'Date', 'Model', 'Issues Encountered'],
        [
            ['v1', 'Jan 2026', 'KVM VMs + GPU passthrough', 'D3cold power bug, no fractional GPU'],
            ['v2', 'Feb 2026', 'Hybrid VMs + Containers', 'Mode switching complexity'],
            ['v3', 'Mar 2026', 'Unified Containers', 'Current — all major issues resolved'],
        ],
        col_widths=[0.8, 1.0, 2.0, 2.7])
    
    add_subsection(doc, 'D3cold Power State Bug')
    add_body(doc, 'Initial VM approach caused GPU to enter D3cold power state after VM shutdown, '
        'requiring host reboot to recover. Eliminated by moving to container-only architecture '
        'where GPU remains attached to host.')
    
    add_subsection(doc, 'NVENC Concurrent Session Limit')
    add_body(doc, 'Consumer GPUs have 3-5 concurrent NVENC session limit. Resolved using nvidia-patch '
        'to remove artificial driver limit.')
    
    add_subsection(doc, 'RTX 5090 Constraints')
    add_callout(doc, 'HARDWARE LIMITATION',
        'RTX 5090 (consumer GPU) does not support vGPU, MIG, or SR-IOV hardware partitioning. '
        'Software partitioning via HAMi-core + MPS is the only viable approach.',
        'critical')

def add_section_16(doc):
    """Section 16 — Architecture Reconsiderations & Lessons"""
    add_section_header(doc, '16', 'Architecture Reconsiderations & Lessons')
    
    add_subsection(doc, 'Why Containers Over VMs')
    add_tree_bullet(doc, [
        'Faster startup time (seconds vs minutes)',
        'Lower overhead (no hypervisor layer)',
        'GPU remains attached to host — no D3cold issues',
        'Easier resource management with cgroups',
    ])
    
    add_subsection(doc, 'Why Not Kubernetes on Day 1')
    add_tree_bullet(doc, [
        'Complexity overhead for 4-node cluster',
        'GPU scheduling in K8s requires additional tooling',
        'Docker Compose sufficient for POC validation',
        'Planned migration to K8s for production scale-out',
    ])
    
    add_subsection(doc, 'Why ZFS Over Ceph')
    add_tree_bullet(doc, [
        'Native quotas and snapshots without additional tooling',
        'Lower operational complexity for small cluster',
        'Excellent NFS integration',
        'Per-user datasets provide strong isolation',
    ])

def add_section_17(doc):
    """Section 17 — Risk Assessment"""
    add_section_header(doc, '17', 'Risk Assessment')
    
    add_subsection(doc, 'Technical Risks')
    add_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['GPU driver crash', 'Low', 'High', 'Watchdog + auto-recovery'],
            ['NFS mount failure', 'Medium', 'High', 'Health checks + alerting'],
            ['MPS fault', 'Low', 'Medium', 'Per-client isolation + restart'],
            ['VRAM overcommit', 'Medium', 'Medium', 'HAMi-core enforcement'],
        ],
        col_widths=[1.5, 1.0, 1.0, 2.5])
    
    add_subsection(doc, 'Operational Risks')
    add_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['Power failure', 'Low', 'High', 'UPS + graceful shutdown'],
            ['Network partition', 'Low', 'High', 'Redundant uplinks'],
            ['Storage full', 'Medium', 'Medium', 'Quota enforcement + alerts'],
        ],
        col_widths=[1.5, 1.0, 1.0, 2.5])
    
    add_subsection(doc, 'Licensing Risks')
    add_table(doc,
        ['Risk', 'Likelihood', 'Impact', 'Mitigation'],
        [
            ['NVIDIA EULA violation', 'Low', 'High', 'Legal review of MPS usage'],
            ['NVENC patch legality', 'Medium', 'Medium', 'Educational use exemption'],
        ],
        col_widths=[1.8, 1.0, 1.0, 2.2])

def add_section_18(doc):
    """Section 18 — Technology Stack Summary"""
    add_section_header(doc, '18', 'Technology Stack Summary')
    
    add_table(doc,
        ['Category', 'Technology', 'Version'],
        [
            ['Frontend Framework', 'Next.js', '15.x'],
            ['UI Library', 'React', '19.x'],
            ['Language', 'TypeScript', '5.7'],
            ['Styling', 'Tailwind CSS', '4.x'],
            ['State Management', 'Zustand', '5.x'],
            ['Backend Framework', 'NestJS', '11.x'],
            ['HTTP Server', 'Fastify', '5.x'],
            ['ORM', 'Prisma', '6.0'],
            ['Database', 'PostgreSQL', '16.x'],
            ['Cache', 'Redis', '7.x'],
            ['Container Runtime', 'Docker CE', '26.x'],
            ['GPU Toolkit', 'nvidia-container-toolkit', '1.16'],
            ['GPU Sharing', 'HAMi-core + CUDA MPS', 'Custom'],
            ['Desktop Streaming', 'Selkies EGL + WebRTC', 'Latest'],
            ['Monitoring', 'Prometheus', '3.4.0'],
            ['Visualization', 'Grafana', '11.5.0'],
            ['Logging', 'Loki', '3.4.2'],
            ['Payment Gateway', 'Razorpay', '2.9.6'],
        ],
        col_widths=[2.0, 2.5, 1.5])

def add_appendix_a(doc):
    """Appendix A — API Endpoint Reference"""
    add_section_header(doc, 'A', 'API Endpoint Reference')
    
    add_subsection(doc, 'Authentication Endpoints')
    add_table(doc,
        ['Method', 'Path', 'Description'],
        [
            ['POST', '/auth/login', 'Email/password login'],
            ['POST', '/auth/register', 'New user registration'],
            ['POST', '/auth/refresh', 'Refresh access token'],
            ['POST', '/auth/logout', 'Revoke refresh token'],
            ['GET', '/auth/sso/keycloak', 'Initiate Keycloak SSO'],
            ['GET', '/auth/sso/callback', 'SSO callback handler'],
        ],
        col_widths=[1.0, 2.5, 2.5])
    
    add_subsection(doc, 'Storage Endpoints')
    add_table(doc,
        ['Method', 'Path', 'Description'],
        [
            ['GET', '/storage/files', 'List user files'],
            ['POST', '/storage/files', 'Upload file (500MB limit)'],
            ['GET', '/storage/files/:path', 'Download file'],
            ['DELETE', '/storage/files/:path', 'Delete file'],
            ['POST', '/storage/folders', 'Create folder'],
            ['GET', '/storage/usage', 'Get storage usage'],
        ],
        col_widths=[1.0, 2.5, 2.5])
    
    add_subsection(doc, 'Session Endpoints')
    add_table(doc,
        ['Method', 'Path', 'Description'],
        [
            ['POST', '/sessions', 'Launch new session'],
            ['GET', '/sessions/:id', 'Get session status'],
            ['POST', '/sessions/:id/stop', 'Stop session'],
            ['GET', '/sessions/:id/connect', 'Get WebRTC URL'],
        ],
        col_widths=[1.0, 2.5, 2.5])
    
    add_subsection(doc, 'Billing Endpoints')
    add_table(doc,
        ['Method', 'Path', 'Description'],
        [
            ['GET', '/wallet', 'Get wallet balance'],
            ['POST', '/wallet/recharge', 'Create recharge order'],
            ['POST', '/wallet/webhook', 'Razorpay webhook'],
            ['GET', '/invoices', 'List invoices'],
            ['GET', '/invoices/:id/pdf', 'Download invoice PDF'],
        ],
        col_widths=[1.0, 2.5, 2.5])

def add_appendix_b(doc):
    """Appendix B — Alert Rules Reference"""
    add_section_header(doc, 'B', 'Alert Rules Reference')
    
    add_subsection(doc, 'GPU Alerts')
    add_table(doc,
        ['Alert Name', 'Condition', 'Severity'],
        [
            ['GPUTemperatureHigh', 'temp > 85°C for 5m', 'warning'],
            ['GPUTemperatureCritical', 'temp > 95°C for 1m', 'critical'],
            ['GPUPowerDraw', 'power > 450W for 10m', 'warning'],
            ['GPUVRAMOvercommit', 'allocated > physical', 'critical'],
            ['GPUMPSFault', 'mps_fault_count > 0', 'critical'],
            ['GPUECCError', 'ecc_errors > 10', 'warning'],
        ],
        col_widths=[2.0, 2.5, 1.5])
    
    add_subsection(doc, 'Session Alerts')
    add_table(doc,
        ['Alert Name', 'Condition', 'Severity'],
        [
            ['SessionStartupTimeout', 'starting > 60s', 'warning'],
            ['SessionConcurrentLimit', 'sessions > max_per_node', 'warning'],
            ['WebRTCFPSDrop', 'fps < 20 for 2m', 'warning'],
            ['SessionOrphaned', 'no heartbeat for 10m', 'critical'],
        ],
        col_widths=[2.0, 2.5, 1.5])
    
    add_subsection(doc, 'Host Alerts')
    add_table(doc,
        ['Alert Name', 'Condition', 'Severity'],
        [
            ['DiskSpaceLow', 'free < 10%', 'warning'],
            ['DiskSpaceCritical', 'free < 5%', 'critical'],
            ['NFSMountUnhealthy', 'mount check failed', 'critical'],
            ['CPUSaturation', 'usage > 90% for 15m', 'warning'],
            ['RAMSaturation', 'usage > 95% for 5m', 'warning'],
        ],
        col_widths=[2.0, 2.5, 1.5])
    
    add_subsection(doc, 'Security Alerts')
    add_table(doc,
        ['Alert Name', 'Condition', 'Severity'],
        [
            ['FailedLoginSpike', '> 10 failures in 5m', 'warning'],
            ['SuspiciousAPIPattern', 'anomaly detection trigger', 'warning'],
            ['UnauthorizedAccess', '401/403 spike', 'critical'],
        ],
        col_widths=[2.0, 2.5, 1.5])

# ============================================================================
# MAIN DOCUMENT GENERATION
# ============================================================================

def generate_document():
    """Generate the complete technical architecture document."""
    print("Creating document...")
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Generate sections
    print("Adding cover page...")
    add_cover_page(doc)
    add_page_break(doc)
    
    print("Adding table of contents...")
    add_toc(doc)
    add_page_break(doc)
    
    print("Adding Section 01 - Executive Summary...")
    add_section_01(doc)
    add_page_break(doc)
    
    print("Adding Section 02 - Platform Overview...")
    add_section_02(doc)
    add_page_break(doc)
    
    print("Adding Section 03 - System Architecture...")
    add_section_03(doc)
    add_page_break(doc)
    
    print("Adding Section 04 - Hardware Specifications...")
    add_section_04(doc)
    add_page_break(doc)
    
    print("Adding Section 05 - GPU Fractional Desktop...")
    add_section_05(doc)
    add_page_break(doc)
    
    print("Adding Section 06 - Infrastructure...")
    add_section_06(doc)
    add_page_break(doc)
    
    print("Adding Section 07 - Storage Architecture...")
    add_section_07(doc)
    add_page_break(doc)
    
    print("Adding Section 08 - Authentication...")
    add_section_08(doc)
    add_page_break(doc)
    
    print("Adding Section 09 - Monitoring...")
    add_section_09(doc)
    add_page_break(doc)
    
    print("Adding Section 10 - Application Architecture...")
    add_section_10(doc)
    add_page_break(doc)
    
    print("Adding Section 11 - Billing & Payment...")
    add_section_11(doc)
    add_page_break(doc)
    
    print("Adding Section 12 - Session Lifecycle...")
    add_section_12(doc)
    add_page_break(doc)
    
    print("Adding Section 13 - Isolation...")
    add_section_13(doc)
    add_page_break(doc)
    
    print("Adding Section 14 - Data Flow...")
    add_section_14(doc)
    add_page_break(doc)
    
    print("Adding Section 15 - Failure Analysis...")
    add_section_15(doc)
    add_page_break(doc)
    
    print("Adding Section 16 - Lessons Learned...")
    add_section_16(doc)
    add_page_break(doc)
    
    print("Adding Section 17 - Risk Assessment...")
    add_section_17(doc)
    add_page_break(doc)
    
    print("Adding Section 18 - Technology Stack...")
    add_section_18(doc)
    add_page_break(doc)
    
    print("Adding Appendix A - API Reference...")
    add_appendix_a(doc)
    add_page_break(doc)
    
    print("Adding Appendix B - Alert Rules...")
    add_appendix_b(doc)
    
    # Save document
    output_path = r'c:\Users\Punith\LaaS\LaaS_Technical_Architecture_Report.docx'
    print(f"\nSaving document to {output_path}...")
    doc.save(output_path)
    print("Document generated successfully!")
    return output_path

if __name__ == '__main__':
    generate_document()
