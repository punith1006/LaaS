#!/usr/bin/env python3
"""
LaaS Container Isolation Setup Guide Generator

Generates a comprehensive .docx document covering all sudo isolation and
network isolation configurations for the LaaS platform.

Usage: python generate_isolation_guide.py
Output: LaaS_Container_Isolation_Setup_Guide.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================================
# CONFIGURATION FILE CONTENTS
# ============================================================================

SUDOERS_LAAS_USER = '''# LaaS Container Sudoers Policy
# Full sudo for desktop experience, deny dangerous + GPU-bypass operations
#
# IMPORTANT: Deny rules must target REAL binary paths (readlink -f)
# because sudo resolves symlinks before matching.
#
# Security layers (defense in depth):
#   1. This sudoers file (deny specific dangerous commands)
#   2. Linux capabilities (--cap-drop=ALL + minimal adds)
#   3. Seccomp syscall filter (blocks mount, unshare, setns, etc.)
#   4. AppArmor MAC profile (docker-default)

Defaults env_reset
Defaults env_delete += "LD_PRELOAD LD_LIBRARY_PATH"
Defaults secure_path="/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin"

# Grant ALL with deny exceptions - MUST be single logical line
# Includes both symlink paths AND real binary paths for full coverage
ubuntu ALL=(ALL) NOPASSWD: ALL, \\
    !/sbin/insmod, !/usr/sbin/insmod, \\
    !/sbin/rmmod, !/usr/sbin/rmmod, \\
    !/sbin/modprobe, !/usr/sbin/modprobe, !/usr/bin/kmod, \\
    !/bin/mount, !/usr/bin/mount, !/sbin/mount, !/usr/sbin/mount, \\
    !/bin/umount, !/usr/bin/umount, !/sbin/umount, !/usr/sbin/umount, \\
    !/sbin/mkfs.ext4, !/usr/sbin/mkfs.ext4, !/usr/sbin/mke2fs, !/sbin/mke2fs, \\
    !/sbin/mkfs.ext3, !/usr/sbin/mkfs.ext3, \\
    !/sbin/mkfs.xfs, !/usr/sbin/mkfs.xfs, \\
    !/sbin/mkfs.btrfs, !/usr/sbin/mkfs.btrfs, \\
    !/sbin/fdisk, !/usr/sbin/fdisk, \\
    !/sbin/parted, !/usr/sbin/parted, \\
    !/sbin/iptables, !/usr/sbin/iptables, \\
    !/sbin/ip6tables, !/usr/sbin/ip6tables, \\
    !/sbin/nft, !/usr/sbin/nft, \\
    !/sbin/ip, !/usr/sbin/ip, !/usr/bin/ip, \\
    !/usr/bin/nsenter, !/usr/sbin/nsenter, \\
    !/usr/bin/unshare, !/usr/sbin/unshare, \\
    !/usr/bin/chroot, !/usr/sbin/chroot, \\
    !/usr/bin/docker, !/usr/sbin/docker, \\
    !/usr/bin/python3, !/usr/bin/python3.10, !/usr/bin/python3.11, !/usr/bin/python3.12, !/usr/bin/python3.13, \\
    !/usr/bin/python, !/usr/local/bin/python3, !/usr/local/bin/python, \\
    !/home/ubuntu/
'''

SUDOERS_OVERRIDE = '''# /etc/sudoers - LaaS Container Override
# This file removes the base image's blanket ubuntu ALL grant that
# overrides deny rules in /etc/sudoers.d/laas-user
#
# Managed by LaaS - DO NOT EDIT inside container

Defaults        env_reset
Defaults        mail_badpass
Defaults        secure_path="/usr/local/sbin:/usr/local/bin:/usr/sbin:/usr/bin:/sbin:/bin:/snap/bin"
Defaults        use_pty

root    ALL=(ALL:ALL) ALL
%admin ALL=(ALL) ALL
%sudo   ALL=(ALL:ALL) ALL

@includedir /etc/sudoers.d
'''

SECCOMP_PROFILE = '''{
  "defaultAction": "SCMP_ACT_ALLOW",
  "syscalls": [
    {
      "names": [
        "kexec_load",
        "kexec_file_load",
        "init_module",
        "finit_module",
        "delete_module",
        "bpf",
        "process_vm_writev",
        "process_vm_readv",
        "add_key",
        "keyctl",
        "request_key",
        "reboot",
        "swapon",
        "swapoff",
        "pivot_root",
        "acct",
        "lookup_dcookie",
        "perf_event_open",
        "open_by_handle_at",
        "userfaultfd",
        "mount",
        "umount2",
        "unshare",
        "setns",
        "nfsservctl"
      ],
      "action": "SCMP_ACT_ERRNO",
      "errnoRet": 1
    }
  ]
}'''

NVIDIA_SMI_WRAPPER = '''#!/bin/bash
# LaaS nvidia-smi wrapper - shows user's VRAM limit instead of full GPU memory
get_vram_limit_mib() {
  if [ -n "$CUDA_DEVICE_MEMORY_LIMIT_0" ]; then
    VAL="${CUDA_DEVICE_MEMORY_LIMIT_0,,}"
    if [[ "$VAL" == *"m" ]]; then echo "${VAL//m/}"
    elif [[ "$VAL" == *"g" ]]; then echo $(( ${VAL//g/} * 1024 ))
    fi
    return
  fi
  python3 -c "
import struct
with open('/tmp/cudevshr.cache','rb') as f:
    d=f.read(4096)
print(struct.unpack_from('<Q',d,1600)[0]//(1024*1024))
" 2>/dev/null || echo "0"
}
case "$*" in
  *"-pm"*|*"--persistence-mode"*|*"-pl"*|*"--power-limit"*|\\
  *"--query-gpu"*|*"--format"*|*"-i "*|*"--id"*)
    exec env -u LD_PRELOAD /usr/bin/nvidia-smi.real "$@" ;;
esac
LIMIT_MIB=$(get_vram_limit_mib)
REAL_OUT=$(env -u LD_PRELOAD /usr/bin/nvidia-smi.real "$@" 2>&1)
if [ -n "$LIMIT_MIB" ] && [ "$LIMIT_MIB" -gt "0" ] 2>/dev/null; then
  echo "$REAL_OUT" | sed "s|/ [0-9]*MiB|/ ${LIMIT_MIB}MiB|g"
else
  echo "$REAL_OUT"
fi
'''

PASSWD_WRAPPER = '''#!/bin/bash
# LaaS passwd wrapper - strips LD_PRELOAD to prevent HAMi conflicts
exec env -u LD_PRELOAD /usr/bin/passwd.real "$@"
'''

BASH_BASHRC_LAAS_BLOCK = '''# LaaS resource interceptors
# fake_sysconf.so is safe for all programs (KDE RAM display fix)
# libvgpu.so ONLY for CUDA programs (crashes non-CUDA via dlsym hooks)
export LD_PRELOAD="/usr/lib/fake_sysconf.so"
export SYSCONF_INJECTED=1
mkdir -p /tmp/vgpulock 2>/dev/null

# CUDA program wrappers — inject HAMi VRAM/SM enforcement
# IMPORTANT: space-separated, fake_sysconf FIRST (proven working order)
_HAMI_PRELOAD="/usr/lib/fake_sysconf.so /usr/lib/libvgpu.so"
python3() { LD_PRELOAD="$_HAMI_PRELOAD" command /usr/bin/python3 "$@"; }
python()  { LD_PRELOAD="$_HAMI_PRELOAD" command /usr/bin/python "$@"; }
nvcc()    { LD_PRELOAD="$_HAMI_PRELOAD" command /usr/local/cuda/bin/nvcc "$@"; }
jupyter() { LD_PRELOAD="$_HAMI_PRELOAD" command /usr/bin/jupyter "$@"; }

# Smart sudo wrapper: strip LD_PRELOAD so system tools dont crash
sudo() {
    env -u LD_PRELOAD /usr/bin/sudo "$@"
}

# Package managers: always strip LD_PRELOAD
apt()     { env -u LD_PRELOAD /usr/bin/apt "$@"; }
apt-get() { env -u LD_PRELOAD /usr/bin/apt-get "$@"; }
dpkg()    { env -u LD_PRELOAD /usr/bin/dpkg "$@"; }
pip()     { env -u LD_PRELOAD /usr/bin/pip "$@"; }
pip3()    { env -u LD_PRELOAD /usr/bin/pip3 "$@"; }
'''

TURNSERVER_CONF = '''# LaaS TURN Server Configuration
# Used by Selkies GStreamer for WebRTC relay in bridge-networked containers
listening-port=3478
listening-ip=0.0.0.0
external-ip=100.100.66.101
fingerprint
lt-cred-mech
user=selkies:wVIAbfwkgkxjaCiZVX4BDsdU
realm=selkies
total-quota=100
stale-nonce=600
no-multicast-peers
no-cli
no-tlsv1
no-tlsv1_1
min-port=49152
max-port=65535
'''

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def set_cell_shading(cell, hex_color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc, code, font_size=8):
    """Add a code block with monospace formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def add_inline_code(para, text):
    """Add inline code within a paragraph."""
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return run


# ============================================================================
# DOCUMENT GENERATION
# ============================================================================

def generate_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('LaaS Container Isolation Setup Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Configuration Reference for Sudo & Network Isolation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph(f'Generated: March 2026 | Version 1.0')
    doc.add_paragraph()
    
    # ========================================================================
    # SECTION 1: OVERVIEW
    # ========================================================================
    doc.add_heading('1. Overview', level=1)
    
    doc.add_paragraph(
        'This guide documents the complete configuration for LaaS container isolation, '
        'enabling a "bare-metal Ubuntu experience in the browser" while maintaining strict '
        'security boundaries. Users get full sudo access for package installation and system '
        'administration, while dangerous operations that could compromise host isolation are blocked.'
    )
    
    doc.add_heading('What This Guide Covers', level=2)
    bullets = [
        'Five-layer defense-in-depth security architecture',
        'Real sudo binary deployment (replacing fakeroot)',
        'Sudoers deny rules for dangerous commands',
        'Seccomp syscall filtering',
        'Linux capability restrictions',
        'Bridge network isolation with TURN server for WebRTC',
        'DOCKER-USER iptables rules for host protection',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    
    doc.add_heading('Architecture Goals', level=2)
    doc.add_paragraph(
        'The goal is to provide containers that look and feel like dedicated VMs with full '
        'administrative access, while actually running as isolated Docker containers with '
        'fractional GPU resources (via HAMi-Core). Users should be able to:'
    )
    goals = [
        'Install packages with apt/pip using sudo',
        'Run CUDA applications with enforced VRAM limits',
        'See accurate resource limits (CPU, RAM, VRAM) via standard tools',
        'Access a full KDE desktop via WebRTC in browser',
    ]
    for g in goals:
        doc.add_paragraph(g, style='List Bullet')
    
    doc.add_paragraph('While being blocked from:')
    blocked = [
        'Loading kernel modules (modprobe, insmod)',
        'Mounting filesystems or escaping with mount/chroot',
        'Bypassing GPU limits with sudo python3',
        'Accessing host services (NestJS backend, Keycloak, etc.)',
        'Communicating with other user containers',
    ]
    for b in blocked:
        doc.add_paragraph(b, style='List Bullet')
    
    # ========================================================================
    # SECTION 2: PREREQUISITES
    # ========================================================================
    doc.add_heading('2. Prerequisites', level=1)
    
    doc.add_paragraph('Before deploying, ensure the following are installed and configured:')
    
    prereqs = [
        ('Docker', '29.3.0 or later (with BuildKit)'),
        ('Ubuntu Host', '22.04 LTS or 24.04 LTS'),
        ('lxcfs', 'For accurate /proc/meminfo, /proc/cpuinfo inside containers'),
        ('AppArmor', 'docker-default profile enabled'),
        ('NFS Storage', 'Mounted at /mnt/nfs/users for persistent home directories'),
        ('HAMi-Core Libraries', 'libvgpu.so and fake_sysconf.so deployed to /usr/lib/'),
        ('Selkies Image', 'ghcr.io/selkies-project/nvidia-egl-desktop:latest pulled'),
        ('CUDA MPS Daemon', 'Running for GPU time-slicing'),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Requirement'
    set_cell_shading(hdr_cells[0], 'D9E2F3')
    set_cell_shading(hdr_cells[1], 'D9E2F3')
    
    for comp, req in prereqs:
        row = table.add_row().cells
        row[0].text = comp
        row[1].text = req
    
    doc.add_paragraph()
    
    # ========================================================================
    # SECTION 3: SECURITY ARCHITECTURE
    # ========================================================================
    doc.add_heading('3. Security Architecture — Five-Layer Defense-in-Depth', level=1)
    
    doc.add_paragraph(
        'LaaS uses multiple overlapping security layers. If one layer is bypassed, '
        'others still provide protection:'
    )
    
    doc.add_heading('Layer 1: Real Sudo Binary (Bind-Mount)', level=2)
    doc.add_paragraph(
        'The Selkies base image replaces /usr/bin/sudo with a symlink to /usr/bin/fakeroot. '
        'This allows unprivileged users to run commands as "fake root" for building packages, '
        'but it completely bypasses our sudoers deny rules. We bind-mount the real sudo binary '
        '(extracted from the image at /usr/bin/sudo-root) over /usr/bin/sudo.'
    )
    
    doc.add_heading('Layer 2: Sudoers Deny Rules', level=2)
    doc.add_paragraph(
        'The /etc/sudoers.d/laas-user file grants ubuntu ALL access with explicit deny rules '
        'for dangerous commands. Key insight: sudo resolves symlinks to real binary paths before '
        'matching, so we must deny both symlink paths AND real binary paths. Examples:'
    )
    examples = [
        'python3 → python3.12 (must deny both)',
        'modprobe → /usr/bin/kmod (must deny both)',
        'mkfs.ext4 → /usr/sbin/mke2fs (must deny both)',
    ]
    for e in examples:
        p = doc.add_paragraph(style='List Bullet')
        add_inline_code(p, e)
    
    doc.add_heading('Layer 3: Custom /etc/sudoers Override', level=2)
    doc.add_paragraph(
        'The Selkies base image\'s /etc/sudoers file has "ubuntu ALL=(ALL:ALL) NOPASSWD: ALL" '
        'AFTER the @includedir directive. Due to sudoers "last match wins" semantics, this '
        'overrides all deny rules in /etc/sudoers.d/. We bind-mount a clean /etc/sudoers that '
        'removes this blanket grant.'
    )
    
    doc.add_heading('Layer 4: Seccomp Syscall Filter', level=2)
    doc.add_paragraph(
        'A custom seccomp profile blocks dangerous syscalls at the kernel level:'
    )
    syscalls = ['mount', 'umount2', 'unshare', 'setns', 'kexec_load', 'init_module', 
                'delete_module', 'bpf', 'pivot_root', 'reboot']
    p = doc.add_paragraph()
    p.add_run('Blocked: ').bold = True
    p.add_run(', '.join(syscalls))
    
    doc.add_heading('Layer 5: Capability Dropping', level=2)
    doc.add_paragraph('We drop ALL capabilities and add back only what\'s needed:')
    
    caps_table = doc.add_table(rows=1, cols=2)
    caps_table.style = 'Table Grid'
    hdr = caps_table.rows[0].cells
    hdr[0].text = 'Capability'
    hdr[1].text = 'Purpose'
    set_cell_shading(hdr[0], 'D9E2F3')
    set_cell_shading(hdr[1], 'D9E2F3')
    
    caps = [
        ('CHOWN', 'File ownership changes'),
        ('DAC_OVERRIDE', 'Bypass file read/write permission'),
        ('FOWNER', 'Bypass permission checks'),
        ('SETUID', 'Required for sudo setuid'),
        ('SETGID', 'Required for sudo setgid'),
        ('NET_BIND_SERVICE', 'Bind to ports < 1024'),
        ('KILL', 'Send signals'),
        ('SYS_CHROOT', 'Limited use of chroot'),
        ('MKNOD', 'Create device files'),
        ('NET_RAW', 'Raw sockets (ping)'),
        ('FSETID', 'Preserve setuid on modify'),
        ('AUDIT_WRITE', 'Write to audit log'),
    ]
    for cap, purpose in caps:
        row = caps_table.add_row().cells
        row[0].text = cap
        row[1].text = purpose
    
    doc.add_paragraph()
    doc.add_paragraph('Plus: AppArmor docker-default MAC profile for mandatory access control.')
    
    # ========================================================================
    # SECTION 4: HOST CONFIGURATION FILES
    # ========================================================================
    doc.add_heading('4. Host Configuration Files', level=1)
    
    doc.add_paragraph('All configuration files are stored in /etc/laas/ on the host machine.')
    
    # 4.1 sudoers-laas-user
    doc.add_heading('/etc/laas/sudoers-laas-user', level=2)
    doc.add_paragraph('Permissions: 440 (root:root)')
    add_code_block(doc, SUDOERS_LAAS_USER, font_size=7)
    
    # 4.2 sudoers override
    doc.add_heading('/etc/laas/sudoers', level=2)
    doc.add_paragraph(
        'This overrides the base image\'s /etc/sudoers to remove the blanket ubuntu ALL grant. '
        'Permissions: 440 (root:root)'
    )
    add_code_block(doc, SUDOERS_OVERRIDE, font_size=8)
    
    # 4.3 seccomp profile
    doc.add_heading('/etc/laas/seccomp-gpu-desktop.json', level=2)
    doc.add_paragraph('Permissions: 644 (root:root)')
    add_code_block(doc, SECCOMP_PROFILE, font_size=8)
    
    # 4.4 bash.bashrc
    doc.add_heading('/etc/laas/bash.bashrc', level=2)
    doc.add_paragraph(
        'Custom bash.bashrc with LD_PRELOAD management for HAMi-Core GPU limits. '
        'Key features: strips LD_PRELOAD for system tools, injects HAMi libs for CUDA programs. '
        'Permissions: 644 (root:root)'
    )
    doc.add_paragraph('LaaS-specific block (appended to standard Ubuntu bash.bashrc):')
    add_code_block(doc, BASH_BASHRC_LAAS_BLOCK, font_size=8)
    
    # 4.5 supervisord-hami.conf
    doc.add_heading('/etc/laas/supervisord-hami.conf', level=2)
    doc.add_paragraph(
        'Supervisord configuration that manages desktop services inside the container. '
        'Key feature: injects LD_PRELOAD=/usr/lib/fake_sysconf.so for the entrypoint and '
        'LD_PRELOAD=/usr/lib/libvgpu.so for selkies-gstreamer (GPU encoder). '
        'Permissions: 644 (root:root)'
    )
    
    # 4.6 nvidia-smi-wrapper
    doc.add_heading('/etc/laas/nvidia-smi-wrapper', level=2)
    doc.add_paragraph(
        'Shows the user\'s VRAM limit instead of full GPU memory. Permissions: 755 (root:root)'
    )
    add_code_block(doc, NVIDIA_SMI_WRAPPER, font_size=8)
    
    # 4.7 passwd-wrapper
    doc.add_heading('/etc/laas/passwd-wrapper', level=2)
    doc.add_paragraph(
        'Strips LD_PRELOAD when changing passwords to prevent HAMi library conflicts. '
        'Permissions: 755 (root:root)'
    )
    add_code_block(doc, PASSWD_WRAPPER, font_size=9)
    
    # 4.8 sudo-bin
    doc.add_heading('/etc/laas/sudo-bin', level=2)
    doc.add_paragraph(
        'The real sudo binary extracted from the Selkies image. Must have setuid bit set. '
        'Permissions: 4755 (root:root)'
    )
    
    # ========================================================================
    # SECTION 5: REAL SUDO BINARY EXTRACTION
    # ========================================================================
    doc.add_heading('5. Real Sudo Binary Extraction', level=1)
    
    doc.add_paragraph(
        'The Selkies image symlinks /usr/bin/sudo to fakeroot for unprivileged package building. '
        'The real sudo binary is preserved at /usr/bin/sudo-root. Extract it:'
    )
    
    extract_cmd = '''# Extract real sudo from Selkies image
TEMP_CONTAINER=$(docker create ghcr.io/selkies-project/nvidia-egl-desktop:latest)
docker cp "$TEMP_CONTAINER:/usr/bin/sudo-root" /etc/laas/sudo-bin
docker rm "$TEMP_CONTAINER"

# Set correct permissions (setuid root)
chmod 4755 /etc/laas/sudo-bin
chown root:root /etc/laas/sudo-bin

# Verify
ls -la /etc/laas/sudo-bin
# Should show: -rwsr-xr-x 1 root root ... /etc/laas/sudo-bin'''
    add_code_block(doc, extract_cmd)
    
    # ========================================================================
    # SECTION 6: NETWORK ISOLATION SETUP
    # ========================================================================
    doc.add_heading('6. Network Isolation Setup', level=1)
    
    doc.add_paragraph(
        'Network isolation prevents containers from accessing host services (NestJS backend, '
        'Keycloak, etc.) while still allowing internet access for apt/pip.'
    )
    
    # 6a Bridge Network
    doc.add_heading('6a. Bridge Network Creation', level=2)
    doc.add_paragraph(
        'Create a dedicated bridge network with inter-container communication disabled:'
    )
    bridge_cmd = '''# Create isolated bridge network (if not exists)
docker network create --driver bridge \\
  --subnet=172.31.0.0/16 \\
  --opt com.docker.network.bridge.enable_icc=false \\
  laas-sessions

# Verify
docker network inspect laas-sessions | grep -A5 Options'''
    add_code_block(doc, bridge_cmd)
    
    p = doc.add_paragraph()
    p.add_run('Key setting: ').bold = True
    add_inline_code(p, 'enable_icc=false')
    p.add_run(' prevents containers from talking to each other.')
    
    # 6b DOCKER-USER iptables
    doc.add_heading('6b. DOCKER-USER iptables Rules', level=2)
    doc.add_paragraph(
        'Docker container traffic goes through the FORWARD chain, NOT the INPUT chain. '
        'Rules must be added to the DOCKER-USER chain, which Docker processes before its own rules.'
    )
    
    doc.add_paragraph().add_run('Rule ordering is critical:').bold = True
    order = [
        'Rule 1: RETURN for ESTABLISHED,RELATED connections (allows responses to incoming requests)',
        'Rule 2-4: RETURN for TURN server traffic (TCP/UDP 3478, UDP 49152-65535)',
        'Rule 5+: DROP for private IP ranges (10.0.0.0/8, 172.16.0.0/12, 192.168.0.0/16, 100.64.0.0/10)',
    ]
    for o in order:
        doc.add_paragraph(o, style='List Number')
    
    iptables_cmd = '''# Variables
LAAS_SUBNET="172.31.0.0/16"
TURN_IP="100.100.66.101"  # Your Tailscale IP

# Rule 1: Allow established connections (CRITICAL for published ports via Tailscale)
iptables -I DOCKER-USER -m conntrack --ctstate ESTABLISHED,RELATED -j RETURN

# Rules 2-4: Allow TURN server access from containers
iptables -I DOCKER-USER 2 -s "$LAAS_SUBNET" -d "$TURN_IP/32" -p tcp --dport 3478 -j RETURN
iptables -I DOCKER-USER 3 -s "$LAAS_SUBNET" -d "$TURN_IP/32" -p udp --dport 3478 -j RETURN
iptables -I DOCKER-USER 4 -s "$LAAS_SUBNET" -d "$TURN_IP/32" -p udp --dport 49152:65535 -j RETURN

# Rules 5+: Block private IP ranges (host isolation)
for dst in 10.0.0.0/8 172.16.0.0/12 192.168.0.0/16 100.64.0.0/10 169.254.0.0/16; do
    iptables -A DOCKER-USER -s "$LAAS_SUBNET" -d "$dst" -j DROP
done

# Verify rules
iptables -L DOCKER-USER -n -v --line-numbers'''
    add_code_block(doc, iptables_cmd, font_size=7)
    
    # 6c iptables persistence
    doc.add_heading('6c. iptables Persistence', level=2)
    doc.add_paragraph('Make rules survive reboot:')
    persist_cmd = '''# Install persistence package
sudo apt install iptables-persistent -y

# Save current rules
sudo netfilter-persistent save

# Rules saved to /etc/iptables/rules.v4'''
    add_code_block(doc, persist_cmd)
    
    # ========================================================================
    # SECTION 7: TURN SERVER SETUP
    # ========================================================================
    doc.add_heading('7. TURN Server (coturn) Setup', level=1)
    
    doc.add_paragraph(
        'With bridge networking, containers only see their internal IP (172.31.0.x). '
        'WebRTC ICE candidates using this IP cannot be reached by browsers. '
        'The TURN server relays WebRTC traffic, making desktop streaming work.'
    )
    
    doc.add_heading('Installation', level=2)
    install_cmd = '''# Install coturn
sudo apt install -y coturn'''
    add_code_block(doc, install_cmd)
    
    doc.add_heading('/etc/default/coturn', level=2)
    doc.add_paragraph('Enable the coturn daemon:')
    default_cmd = '''# Uncomment or add this line
TURNSERVER_ENABLED=1'''
    add_code_block(doc, default_cmd)
    
    doc.add_heading('/etc/turnserver.conf', level=2)
    doc.add_paragraph(
        'Important: external-ip must be reachable from both browsers AND containers. '
        'Use your Tailscale IP, not a public IP that containers can\'t reach.'
    )
    add_code_block(doc, TURNSERVER_CONF, font_size=8)
    
    doc.add_heading('Start and Enable', level=2)
    start_cmd = '''# Enable and start coturn
sudo systemctl enable coturn
sudo systemctl restart coturn

# Verify running
systemctl status coturn
ss -tulpn | grep 3478'''
    add_code_block(doc, start_cmd)
    
    # ========================================================================
    # SECTION 8: DOCKER RUN COMMAND
    # ========================================================================
    doc.add_heading('8. Complete Docker Run Command Template', level=1)
    
    doc.add_paragraph(
        'This is the full docker run command with all isolation features. '
        'In production, this is generated by the session orchestration service.'
    )
    
    docker_cmd = '''docker run -d \\
  --name laas-test-session \\
  --hostname ws-test \\
  --add-host ws-test:127.0.0.1 \\
  --restart unless-stopped \\
  
  # GPU Access
  --gpus all \\
  
  # Resource Limits
  --cpus=4 \\
  --cpuset-cpus=2-5 \\
  --memory=8g \\
  --pids-limit 512 \\
  
  # IPC (required for WebRTC shared memory and MPS)
  --ipc=host \\
  
  # Network: Isolated bridge (NOT --network=host)
  --network=laas-sessions \\
  --dns 8.8.8.8 --dns 8.8.4.4 \\
  -p 8080:8080 \\
  -p 19080:19080 \\
  
  # Security: Capabilities
  --cap-drop=ALL \\
  --cap-add=CHOWN --cap-add=DAC_OVERRIDE --cap-add=FOWNER \\
  --cap-add=SETUID --cap-add=SETGID --cap-add=NET_BIND_SERVICE \\
  --cap-add=KILL --cap-add=SYS_CHROOT --cap-add=MKNOD \\
  --cap-add=NET_RAW --cap-add=FSETID --cap-add=AUDIT_WRITE \\
  
  # Security: Seccomp + AppArmor
  --security-opt seccomp=/etc/laas/seccomp-gpu-desktop.json \\
  --security-opt apparmor=docker-default \\
  --security-opt no-new-privileges=false \\
  
  # Environment: Display
  -e TZ=UTC \\
  -e DISPLAY=:20 \\
  -e DISPLAY_SIZEW=1920 -e DISPLAY_SIZEH=1080 -e DISPLAY_REFRESH=60 \\
  
  # Environment: Selkies
  -e SELKIES_ENCODER=nvh264enc \\
  -e SELKIES_ENABLE_BASIC_AUTH=true \\
  -e SELKIES_BASIC_AUTH_PASSWORD=your-password \\
  -e NGINX_PORT=8080 -e SELKIES_PORT=9080 -e SELKIES_METRICS_HTTP_PORT=19080 \\
  
  # Environment: TURN Server
  -e SELKIES_TURN_HOST=100.100.66.101 \\
  -e SELKIES_TURN_PORT=3478 \\
  -e SELKIES_TURN_USERNAME=selkies \\
  -e SELKIES_TURN_PASSWORD=wVIAbfwkgkxjaCiZVX4BDsdU \\
  -e SELKIES_TURN_PROTOCOL=tcp \\
  
  # Environment: GPU/HAMi
  -e CUDA_VISIBLE_DEVICES=0 \\
  -e CUDA_DEVICE_MEMORY_LIMIT_0=4096m \\
  -e CUDA_DEVICE_SM_LIMIT=17 \\
  -e CUDA_MPS_PIPE_DIRECTORY=/tmp/nvidia-mps \\
  -e CUDA_MPS_PINNED_DEVICE_MEM_LIMIT=0=4G \\
  -e CUDA_MPS_ACTIVE_THREAD_PERCENTAGE=17 \\
  
  # Volume Mounts: MPS
  -v /tmp/nvidia-mps:/tmp/nvidia-mps \\
  -v /tmp/nvidia-log:/tmp/nvidia-log \\
  
  # Volume Mounts: HAMi libraries
  -v /usr/lib/libvgpu.so:/usr/lib/libvgpu.so:ro \\
  -v /usr/lib/fake_sysconf.so:/usr/lib/fake_sysconf.so:ro \\
  
  # Volume Mounts: Wrappers
  -v /usr/bin/nvidia-smi:/usr/bin/nvidia-smi.real \\
  -v /etc/laas/nvidia-smi-wrapper:/usr/bin/nvidia-smi:ro \\
  -v /usr/bin/passwd:/usr/bin/passwd.real \\
  -v /etc/laas/passwd-wrapper:/usr/bin/passwd:ro \\
  
  # Volume Mounts: Config files
  -v /etc/laas/supervisord-hami.conf:/etc/supervisord.conf:ro \\
  -v /etc/laas/bash.bashrc:/etc/bash.bashrc:ro \\
  
  # Volume Mounts: Sudoers (CRITICAL ORDER)
  -v /etc/laas/sudoers:/etc/sudoers:ro \\
  -v /etc/laas/sudoers-laas-user:/etc/sudoers.d/laas-user:ro \\
  -v /etc/laas/sudo-bin:/usr/bin/sudo \\
  
  # Volume Mounts: lxcfs (fake /proc)
  -v /var/lib/lxcfs/proc/cpuinfo:/proc/cpuinfo:ro \\
  -v /var/lib/lxcfs/proc/meminfo:/proc/meminfo:ro \\
  -v /var/lib/lxcfs/proc/stat:/proc/stat:ro \\
  -v /var/lib/lxcfs/proc/uptime:/proc/uptime:ro \\
  
  # Image
  ghcr.io/selkies-project/nvidia-egl-desktop:latest'''
    add_code_block(doc, docker_cmd, font_size=6)
    
    # ========================================================================
    # SECTION 9: KEY DISCOVERIES AND PITFALLS
    # ========================================================================
    doc.add_heading('9. Key Discoveries and Pitfalls', level=1)
    
    doc.add_paragraph(
        'These issues were discovered during development and are critical to understand:'
    )
    
    # Pitfall 1
    doc.add_heading('1. Fakeroot Masquerading as Sudo', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    add_inline_code(p, '/usr/bin/sudo')
    p.add_run(' in Selkies image is a symlink to ')
    add_inline_code(p, '/usr/bin/fakeroot')
    p.add_run('. Running ')
    add_inline_code(p, 'sudo --version')
    p.add_run(' returns "fakeroot version 1.33".')
    p2 = doc.add_paragraph()
    p2.add_run('Impact: ').bold = True
    p2.add_run('All sudoers rules are completely bypassed because fakeroot doesn\'t read them.')
    p3 = doc.add_paragraph()
    p3.add_run('Solution: ').bold = True
    p3.add_run('Bind-mount real sudo binary from ')
    add_inline_code(p3, '/usr/bin/sudo-root')
    p3.add_run(' over ')
    add_inline_code(p3, '/usr/bin/sudo')
    
    # Pitfall 2
    doc.add_heading('2. Sudoers "Last Match Wins"', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('Base image\'s /etc/sudoers has ')
    add_inline_code(p, 'ubuntu ALL=(ALL:ALL) NOPASSWD: ALL')
    p.add_run(' AFTER ')
    add_inline_code(p, '@includedir /etc/sudoers.d')
    p2 = doc.add_paragraph()
    p2.add_run('Impact: ').bold = True
    p2.add_run('Due to "last match wins" rule, this blanket grant overrides ALL deny rules in /etc/sudoers.d/')
    p3 = doc.add_paragraph()
    p3.add_run('Solution: ').bold = True
    p3.add_run('Bind-mount a clean /etc/sudoers that removes this blanket grant.')
    
    # Pitfall 3
    doc.add_heading('3. Binary Symlink Resolution', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('Sudo resolves full symlink chains via realpath() before matching deny rules.')
    doc.add_paragraph('Examples:')
    examples = [
        '/usr/bin/python3 → /usr/bin/python3.12',
        '/sbin/modprobe → /usr/bin/kmod',
        '/sbin/mkfs.ext4 → /usr/sbin/mke2fs',
    ]
    for e in examples:
        p = doc.add_paragraph(style='List Bullet')
        add_inline_code(p, e)
    p = doc.add_paragraph()
    p.add_run('Solution: ').bold = True
    p.add_run('Deny BOTH the symlink path AND the real binary path.')
    
    # Pitfall 4
    doc.add_heading('4. CRLF Line Endings', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('Windows-created files with CRLF (\\r\\n) line endings cause sudoers syntax errors.')
    p2 = doc.add_paragraph()
    p2.add_run('Impact: ').bold = True
    p2.add_run('The \\r on the last line (no backslash continuation) breaks the sudoers parser.')
    p3 = doc.add_paragraph()
    p3.add_run('Solution: ').bold = True
    p3.add_run('Always convert to Unix LF line endings before deploying.')
    add_code_block(doc, 'dos2unix /etc/laas/sudoers-laas-user')
    
    # Pitfall 5
    doc.add_heading('5. Sudoers Glob Syntax', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    add_inline_code(p, '!/home/ubuntu/*')
    p.add_run(' caused syntax error in sudoers.')
    p2 = doc.add_paragraph()
    p2.add_run('Solution: ').bold = True
    p2.add_run('Use directory form ')
    add_inline_code(p2, '!/home/ubuntu/')
    p2.add_run(' instead of glob wildcards.')
    
    # Pitfall 6
    doc.add_heading('6. Docker Traffic Goes Through FORWARD, Not INPUT', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('Initial attempts to block container→host traffic used INPUT chain rules, which had zero effect.')
    p2 = doc.add_paragraph()
    p2.add_run('Explanation: ').bold = True
    p2.add_run('Container traffic traverses the FORWARD chain (not INPUT) because it crosses network namespaces. Docker processes DOCKER-USER first.')
    p3 = doc.add_paragraph()
    p3.add_run('Solution: ').bold = True
    p3.add_run('Add rules to DOCKER-USER chain, not INPUT.')
    
    # Pitfall 7
    doc.add_heading('7. TURN Server Required for Bridge Networking', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('With --network=host, WebRTC used host IPs directly. With bridge networking, containers only advertise internal IPs (172.31.0.x) that browsers can\'t reach.')
    p2 = doc.add_paragraph()
    p2.add_run('Impact: ').bold = True
    p2.add_run('ICE negotiation fails, desktop streaming doesn\'t work.')
    p3 = doc.add_paragraph()
    p3.add_run('Solution: ').bold = True
    p3.add_run('Run coturn TURN server; containers and browsers relay through it.')
    
    # Pitfall 8
    doc.add_heading('8. SELKIES_TURN_HOST Must Use Tailscale IP', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('TURN server address must be reachable from BOTH the browser AND the container.')
    p2 = doc.add_paragraph()
    p2.add_run('Constraint: ').bold = True
    p2.add_run('Public IP may work for browsers but containers can\'t reach it (blocked by iptables). Localhost doesn\'t work for browsers.')
    p3 = doc.add_paragraph()
    p3.add_run('Solution: ').bold = True
    p3.add_run('Use Tailscale IP (e.g., 100.100.66.101) as TURN_HOST. Add iptables exception for containers to reach TURN.')
    
    # Pitfall 9
    doc.add_heading('9. Conntrack ESTABLISHED Rule Required', level=2)
    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('After adding DROP rules for private IPs, published ports stopped working for Tailscale clients.')
    p2 = doc.add_paragraph()
    p2.add_run('Root cause: ').bold = True
    p2.add_run('Response packets from containers have destination in 100.64.0.0/10 (Tailscale range), matching DROP rules.')
    p3 = doc.add_paragraph()
    p3.add_run('Solution: ').bold = True
    p3.add_run('Add conntrack ESTABLISHED,RELATED RETURN rule at TOP of DOCKER-USER chain.')
    
    # ========================================================================
    # SECTION 10: VERIFICATION TEST SUITE
    # ========================================================================
    doc.add_heading('10. Verification Test Suite', level=1)
    
    doc.add_paragraph('Run these tests inside a container to verify isolation:')
    
    doc.add_heading('Should WORK (Allowed)', level=2)
    work_tests = '''# Package installation
sudo apt update && sudo apt install -y vim
echo "Exit code: $?"  # Should be 0

# System status
sudo systemctl status
echo "Exit code: $?"  # Should be 0

# GPU info (shows VRAM limit)
nvidia-smi
# Should show "4096MiB" (or your configured limit), not full GPU memory

# Internet access
curl -I https://google.com
echo "Exit code: $?"  # Should be 0'''
    add_code_block(doc, work_tests)
    
    doc.add_heading('Should be DENIED (Blocked)', level=2)
    deny_tests = '''# Kernel modules
sudo modprobe dummy
# Expected: "Sorry, user ubuntu is not allowed to execute..."

# Python via sudo (bypasses HAMi VRAM limits)
sudo python3 -c "print('test')"
# Expected: "Sorry, user ubuntu is not allowed to execute..."

# Filesystem operations
sudo mkfs.ext4 /dev/null
# Expected: "Sorry, user ubuntu is not allowed to execute..."

# Network manipulation
sudo iptables -L
# Expected: "Sorry, user ubuntu is not allowed to execute..."

# Namespace escape
sudo nsenter -n -t 1
# Expected: "Sorry, user ubuntu is not allowed to execute..."'''
    add_code_block(doc, deny_tests)
    
    doc.add_heading('Network Isolation Tests', level=2)
    net_tests = '''# Should TIMEOUT (host services blocked)
curl -m 5 http://172.17.0.1:3000/health
# Expected: "Connection timed out"

curl -m 5 http://100.100.66.101:3000/health
# Expected: "Connection timed out"

# Should WORK (internet allowed)
curl -I https://pypi.org
# Expected: HTTP 200'''
    add_code_block(doc, net_tests)
    
    doc.add_heading('WebRTC Desktop Test', level=2)
    doc.add_paragraph('Open browser and navigate to: http://100.100.66.101:8080/')
    doc.add_paragraph('Enter the session password when prompted. Desktop should load with video streaming.')
    
    # ========================================================================
    # SECTION 11: FILE LOCATIONS REFERENCE
    # ========================================================================
    doc.add_heading('11. File Locations Reference', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Dev Machine (Windows)'
    hdr[1].text = 'Host Location (Linux)'
    hdr[2].text = 'Container Mount Point'
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
    
    files = [
        ('host-services/config/sudoers-laas-user', '/etc/laas/sudoers-laas-user', '/etc/sudoers.d/laas-user'),
        ('host-services/config/sudoers', '/etc/laas/sudoers', '/etc/sudoers'),
        ('host-services/config/seccomp-gpu-desktop.json', '/etc/laas/seccomp-gpu-desktop.json', '(--security-opt)'),
        ('host-services/config/bash.bashrc', '/etc/laas/bash.bashrc', '/etc/bash.bashrc'),
        ('host-services/config/supervisord-hami.conf', '/etc/laas/supervisord-hami.conf', '/etc/supervisord.conf'),
        ('host-services/config/nvidia-smi-wrapper', '/etc/laas/nvidia-smi-wrapper', '/usr/bin/nvidia-smi'),
        ('host-services/config/passwd-wrapper', '/etc/laas/passwd-wrapper', '/usr/bin/passwd'),
        ('(extracted from image)', '/etc/laas/sudo-bin', '/usr/bin/sudo'),
        ('host-services/session-orchestration/app.py', '~/session-orchestration/app.py', '(not mounted)'),
        ('backend/scripts/host-deploy-sudo-isolation.sh', '~/laas-deploy/backend/scripts/', '(not mounted)'),
    ]
    
    for dev, host, container in files:
        row = table.add_row().cells
        row[0].text = dev
        row[1].text = host
        row[2].text = container
    
    doc.add_paragraph()
    
    # ========================================================================
    # FOOTER
    # ========================================================================
    doc.add_paragraph()
    doc.add_paragraph('— End of Document —').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


# ============================================================================
# MAIN
# ============================================================================

if __name__ == '__main__':
    output_path = os.path.join(os.path.dirname(__file__), 'LaaS_Container_Isolation_Setup_Guide.docx')
    
    print('Generating LaaS Container Isolation Setup Guide...')
    doc = generate_document()
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
